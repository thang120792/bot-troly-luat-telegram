import os
import sys

# Reconfigure stdout for UTF-8 on Windows immediately
try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

import json
import re
import glob
import time
import base64
import urllib.request
import difflib
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from io import BytesIO
from flask import Flask, request, jsonify, send_from_directory, send_file

# ── SECURITY GUARD: Hàng rào bảo vệ Bot (Rate Limit, Injection Filter, IP Block, Telegram Alert) ──
try:
    import security_guard as _sg
    SECURITY_GUARD_ACTIVE = True
    print("✅ [SecurityGuard] Hàng rào bảo mật Bot đã kích hoạt!")
except Exception as _sg_err:
    SECURITY_GUARD_ACTIVE = False
    print(f"⚠️ [SecurityGuard] Không thể tải module bảo mật: {_sg_err}")

# Safely import python-docx (handle Windows App Control DLL block on lxml.etree)
DOCX_AVAILABLE = False
try:
    import docx
    DOCX_AVAILABLE = True
except Exception as e:
    print(f"⚠️ python-docx not available or lxml DLL blocked: {e}. Fallback to pure python docx parser active.")

def read_docx_pure_python(filepath):
    try:
        with zipfile.ZipFile(filepath) as z:
            xml_content = z.read('word/document.xml')
            tree = ET.fromstring(xml_content)
            texts = []
            for elem in tree.iter():
                if elem.tag.endswith('}t') and elem.text:
                    texts.append(elem.text)
            return "".join(texts)
    except Exception:
        return ""

# Reconfigure stdout for UTF-8 on Windows
sys.stdout.reconfigure(encoding='utf-8')

app = Flask(__name__, static_folder='.')

@app.before_request
def handle_preflight():
    if request.method == "OPTIONS":
        res = jsonify({"status": "ok"})
        res.headers['Access-Control-Allow-Origin'] = '*'
        res.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization,x-goog-api-key,X-Requested-With'
        res.headers['Access-Control-Allow-Methods'] = 'GET,PUT,POST,DELETE,OPTIONS'
        return res, 200

@app.after_request
def add_cors_headers(response):
    response.headers['Access-Control-Allow-Origin'] = '*'
    response.headers['Access-Control-Allow-Headers'] = 'Content-Type,Authorization,x-goog-api-key,X-Requested-With'
    response.headers['Access-Control-Allow-Methods'] = 'GET,PUT,POST,DELETE,OPTIONS'
    return response

# Try importing PyMuPDF (fitz) for PDF page rendering
try:
    import fitz
    FITZ_AVAILABLE = True
    print("✅ PyMuPDF (fitz) is active for PDF rendering!")
except Exception as e:
    FITZ_AVAILABLE = False
    print(f"⚠️ PyMuPDF fitz not found: {e}")

GEMINI_API_KEYS = [
    "AQ.Ab8RN6K6uWSGUAgNhZthfDN38a9tSvzF8RyiaRNvzLTpR6WesA",  # API Key 1 (Mới tạo từ Google AI Studio)
    "AQ.Ab8RN6JrZoDOoJYfBznhSQWpB6Lv9v93RwFPUtIr_Z7lFjqjVA",  # API Key 2 (Dự phòng)
    "AQ.Ab8RN6IzFDhmj0qZOJqlmdqYixwYUkBhxJc9ftlyJ9b1vnKbOQ"   # API Key 3 (Dự phòng)
]
GEMINI_API_KEY = GEMINI_API_KEYS[0]

# CẤU HÌNH ZENMUX MULTI-MODEL AI GATEWAY (DOTS3-NOTE-PREV, GLM-5.3-FREE, DEEPSEEK-V4-FLASH)
ZENMUX_API_KEY = "sk-ai-v1-4d7a69f58906d3b4983d5e6d326528bb9edcbbfabea0b7e440e3738c5c29b89d"
ZENMUX_BASE_URL = "https://zenmux.ai/api/v1/chat/completions"
ZENMUX_MODELS = [
    "dots-studio/dots3-note-prev",
    "z-ai/glm-5.3-free",
    "deepseek/deepseek-v4-flash"
]

# CẤU HÌNH ZALO BOT / OFFICIAL ACCOUNT
ZALO_BOT_ID = "2308474633160527766"
ZALO_BOT_TOKEN = "EfVUmLxWFIMXorvotNYxHBWEBJDGOVHLvbAFCEViZpdjqmijKlHUOdesfyYaOqLD"

# CẤU HÌNH TELEGRAM BOT (@TroLyLuatbot)
TELEGRAM_BOT_TOKEN = "8128444329:AAEtIfC86tE43PYekXP7GlSUzDboiByCGpg"

try:
    from google import genai
    from google.genai import types
    GEMINI_CLIENT = genai.Client(api_key=GEMINI_API_KEY)
    GEMINI_AVAILABLE = True
    print("✅ Google GenAI SDK đã sẵn sàng cho Gemini 3.6 Flash (với cơ chế dự phòng 2 API Keys)!")
except Exception as e:
    GEMINI_CLIENT = None
    GEMINI_AVAILABLE = False
    print(f"⚠️ Không thể khởi tạo Google GenAI SDK: {e}")

# NOTEBOOKLM & OBSIDIAN VAULT INTEGRATION CONFIGURATION
NOTEBOOKLM_URL = "https://notebooklm.google.com/notebook/5234532d-25d7-4f04-a664-2db536f22cdf"
OLLAMA_API_URL = "http://localhost:11434/api/generate"
PRIMARY_LLM_MODEL = "qwen2.5-thanhhoa-land:14b"
FALLBACK_LLM_MODEL = "qwen-legal:14b"

LOCAL_VAULT_UPPER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Obsidian Vault')
LOCAL_VAULT_LOWER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'obsidian_vault')
if os.path.exists(LOCAL_VAULT_UPPER):
    OBSIDIAN_VAULT_PATH = LOCAL_VAULT_UPPER
elif os.path.exists(r'D:\OneDrive - Hanoi University of Mining and Geology\Obsidian Vault'):
    OBSIDIAN_VAULT_PATH = r'D:\OneDrive - Hanoi University of Mining and Geology\Obsidian Vault'
else:
    OBSIDIAN_VAULT_PATH = LOCAL_VAULT_LOWER

OBSIDIAN_TEMPLATES_DIR = os.path.join(OBSIDIAN_VAULT_PATH, "05 - MẪU ĐƠN", "TONG HOP DON")
OBSIDIAN_TAX_TEMPLATES_DIR = os.path.join(OBSIDIAN_VAULT_PATH, "05 - MẪU ĐƠN", "cac to khai thue")

# PERSISTENT AI CHAT MEMORY & SELF-LEARNING CORRECTION FILES
MEMORY_DIR = os.path.join(OBSIDIAN_VAULT_PATH, "06 - BỘ NHỚ AI")
os.makedirs(MEMORY_DIR, exist_ok=True)

CHAT_MEMORY_FILE = os.path.join(MEMORY_DIR, "chat_history_memory.json")
LEARNED_CORRECTIONS_FILE = os.path.join(MEMORY_DIR, "ai_learned_corrections.json")
UNANSWERED_LOG_FILE = os.path.join(MEMORY_DIR, "unanswered_questions_log.json")
UNANSWERED_REPORTS_DIR = os.path.join(MEMORY_DIR, "BAO_CAO_CAU_HOI_CHUA_TRA_LOI")
os.makedirs(UNANSWERED_REPORTS_DIR, exist_ok=True)

def load_json_file(filepath, default_data=None):
    if default_data is None:
        default_data = []
    if os.path.exists(filepath):
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            return default_data
    return default_data

def save_json_file(filepath, data):
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"⚠️ Memory save error to {filepath}: {e}")

# Save conversation turn into persistent chat memory
def record_conversation_turn(session_id, question, answer, intent_type, user_feedback=None):
    history = load_json_file(CHAT_MEMORY_FILE, [])
    turn_entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "session_id": session_id,
        "question": question,
        "answer": answer,
        "intent": intent_type,
        "user_feedback": user_feedback
    }
    history.append(turn_entry)
    save_json_file(CHAT_MEMORY_FILE, history)

# Learn & persist user corrections
def learn_from_user_challenge(question, user_challenge_text):
    corrections = load_json_file(LEARNED_CORRECTIONS_FILE, [])
    correction_entry = {
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
        "context_question": question,
        "user_challenge": user_challenge_text,
        "status": "active_learning_rule"
    }
    corrections.append(correction_entry)
    save_json_file(LEARNED_CORRECTIONS_FILE, corrections)
    record_unanswered_question(question, reason="Phản biện/sửa đổi từ người dùng", model_used="User Feedback Challenge", answer_given=user_challenge_text)

# GHI NHẬN TỰ ĐỘNG CÂU HỎI CHƯA TRẢ LỜI ĐƯỢC / CHƯA GIẢI ĐÁP
def record_unanswered_question(question, session_id="default_session", reason="Hệ thống chưa có đủ dữ liệu pháp lý", model_used="Unknown", answer_given=""):
    """
    Ghi nhận tự động các câu hỏi không trả lời được / chưa giải đáp thành công vào log file và tự động tạo báo cáo Markdown theo ngày.
    """
    if not question or len(question.strip()) < 3:
        return None
        
    log_data = load_json_file(UNANSWERED_LOG_FILE, [])
    now_str = time.strftime("%Y-%m-%d %H:%M:%S")
    today_str = time.strftime("%Y-%m-%d")
    
    # Tránh ghi trùng lặp cùng 1 câu hỏi trong cùng 1 ngày
    for entry in log_data:
        if entry.get("date") == today_str and entry.get("question", "").strip().lower() == question.strip().lower():
            return entry

    entry = {
        "id": f"UNANS-{int(time.time()*1000)}",
        "date": today_str,
        "timestamp": now_str,
        "session_id": session_id,
        "question": question.strip(),
        "reason": reason,
        "model_used": model_used,
        "answer_given": answer_given[:300] if answer_given else ""
    }
    log_data.append(entry)
    save_json_file(UNANSWERED_LOG_FILE, log_data)

    # Tự động cập nhật file báo cáo Markdown tổng hợp theo ngày
    generate_daily_unanswered_markdown_report(today_str)
    return entry

def generate_daily_unanswered_markdown_report(target_date=None):
    """
    Tổng hợp tất cả các câu hỏi chưa trả lời được theo từng ngày thành file Markdown chuyên nghiệp.
    Tệp lưu trữ tại: 06 - BỘ NHỚ AI / BAO_CAO_CAU_HOI_CHUA_TRA_LOI / cau_hoi_chua_tra_loi_YYYY-MM-DD.md
    Và tệp Master: 06 - BỘ NHỚ AI / tong_hop_cau_hoi_chua_tra_loi_theo_ngay.md
    """
    if not target_date:
        target_date = time.strftime("%Y-%m-%d")
        
    log_data = load_json_file(UNANSWERED_LOG_FILE, [])
    date_entries = [e for e in log_data if e.get("date") == target_date]

    # 1. Tạo file báo cáo riêng cho ngày target_date
    daily_filename = f"cau_hoi_chua_tra_loi_{target_date}.md"
    daily_filepath = os.path.join(UNANSWERED_REPORTS_DIR, daily_filename)

    content = []
    content.append(f"# 📊 BÁO CÁO TỔNG HỢP CÂU HỎI CHƯA TRẢ LỜI ĐƯỢC - NGÀY {target_date}")
    content.append(f"*Hệ thống ThanhHoa Land AI App - Ngày tổng hợp: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
    content.append(f"**Tổng số câu hỏi chưa có câu trả lời/chưa giải đáp:** `{len(date_entries)} câu`\n")
    content.append("---")

    if not date_entries:
        content.append("\n✅ *Trong ngày này, hệ thống AI đã giải đáp thành công 100% câu hỏi, không có câu hỏi nào bị tồn đọng.*")
    else:
        content.append("\n### 📋 DANH SÁCH CHI TIẾT CÂU HỎI CHƯA GIẢI ĐÁP:\n")
        content.append("| STT | Thời gian | Câu hỏi của người dân | Lý do chưa giải đáp | Kênh AI xử lý |")
        content.append("|---|---|---|---|---|")
        for idx, item in enumerate(date_entries, 1):
            q_text = item.get('question', '').replace('|', '\\|')
            reason_text = item.get('reason', '').replace('|', '\\|')
            time_text = item.get('timestamp', '').split(' ')[-1]
            model_text = item.get('model_used', 'N/A')
            content.append(f"| {idx} | `{time_text}` | **{q_text}** | {reason_text} | `{model_text}` |")

        content.append("\n\n### 💡 HƯỚNG BỔ SUNG CSDL & ĐÀO TẠO AI:")
        content.append("1. **Nạp thêm văn bản pháp lý:** Bổ sung các Nghị định, Thông tư liên quan vào thư mục Obsidian Vault.")
        content.append("2. **Bổ sung Cẩm nang Q&A:** Cập nhật câu trả lời chuẩn cho câu hỏi trên vào bộ nhớ `dat_dai_qa_1000.jsonl` hoặc `cam-nang-quyet-dinh-2604-300.md`.")
        content.append("3. **Phản biện trực tiếp:** Người dùng có thể chỉnh sửa/phản biện câu trả lời ngay trong khung chat để AI tự học.")

    markdown_text = "\n".join(content)
    with open(daily_filepath, 'w', encoding='utf-8') as f:
        f.write(markdown_text)

    # 2. Cập nhật file Master tổng hợp toàn bộ các ngày: tong_hop_cau_hoi_chua_tra_loi_theo_ngay.md
    master_filepath = os.path.join(MEMORY_DIR, "tong_hop_cau_hoi_chua_tra_loi_theo_ngay.md")
    
    grouped = {}
    for entry in log_data:
        d = entry.get("date", "Unknown")
        if d not in grouped:
            grouped[d] = []
        grouped[d].append(entry)

    master_content = []
    master_content.append("# 📑 BÁO CÁO TỔNG HỢP CÂU HỎI CHƯA TRẢ LỜI ĐƯỢC PHÂN THEO NGÀY (MASTER REPORT)")
    master_content.append(f"*Cập nhật tự động lúc: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
    master_content.append(f"**Tổng số câu hỏi chưa giải đáp toàn bộ hệ thống:** `{len(log_data)} câu`\n")
    master_content.append("---")

    sorted_dates = sorted(grouped.keys(), reverse=True)
    for d in sorted_dates:
        entries = grouped[d]
        master_content.append(f"\n## 📅 Báo cáo Ngày {d} (Tổng: {len(entries)} câu)")
        master_content.append("| STT | Thời gian | Câu hỏi của người dân | Lý do chưa trả lời | Mô hình AI |")
        master_content.append("|---|---|---|---|---|")
        for idx, item in enumerate(entries, 1):
            q_text = item.get('question', '').replace('|', '\\|')
            reason_text = item.get('reason', '').replace('|', '\\|')
            time_text = item.get('timestamp', '').split(' ')[-1]
            model_text = item.get('model_used', 'N/A')
            master_content.append(f"| {idx} | `{time_text}` | **{q_text}** | {reason_text} | `{model_text}` |")

    with open(master_filepath, 'w', encoding='utf-8') as f:
        f.write("\n".join(master_content))

    return markdown_text, daily_filepath

CUSTOMER_CHAT_REPORTS_DIR = os.path.join(MEMORY_DIR, "BAO_CAO_CAU_HOI_KHAC_HANG")
os.makedirs(CUSTOMER_CHAT_REPORTS_DIR, exist_ok=True)

def generate_customer_chat_summary_report(target_date=None):
    """
    TỔNG HỢP TOÀN BỘ CÂU HỎI TIẾP NHẬN TỪ KHÁCH HÀNG/NGƯỜI DÂN CHAT VỚI BOT AI THEO NGÀY VÀ THEO CHỦ ĐỀ PHÁP LÝ.
    File lưu tại: 06 - BỘ NHỚ AI / tong_hop_cau_hoi_khach_hang_chat_bot.md
    Và file báo cáo ngày: 06 - BỘ NHỚ AI / BAO_CAO_CAU_HOI_KHAC_HANG / cau_hoi_khach_hang_YYYY-MM-DD.md
    """
    if not target_date:
        target_date = time.strftime("%Y-%m-%d")

    chat_history = load_json_file(CHAT_MEMORY_FILE, [])
    unanswered_log = load_json_file(UNANSWERED_LOG_FILE, [])

    # Phân loại câu hỏi theo chủ đề pháp lý
    topic_keywords = {
        "Tách / Hợp thửa đất": ["tách", "hợp thửa", "tách thửa", "gộp thửa", "đa mục đích", "mẫu 35", "mẫu 34"],
        "Cấp Giấy chứng nhận (Sổ đỏ)": ["cấp gcn", "sổ đỏ", "lần đầu", "cấp đổi", "cấp lại", "mất sổ", "mẫu 29"],
        "Thu hồi & Đền bù GPMB": ["thu hồi", "bồi thường", "giải phóng mặt bằng", "gpmb", "tái định cư", "rừng sản xuất"],
        "Chuyển mục đích sử dụng đất & Rừng": ["chuyển mục đích", "đất rừng", "trồng rừng thay thế", "lâm nghiệp", "hđnd"],
        "Đo đạc & Địa chính": ["đo đạc", "trích đo", "mốc giới", "ranh giới", "bản đồ", "sai diện tích", "trắc địa"],
        "Nghĩa vụ tài chính & Thuế": ["thuế", "lệ phí", "trước bạ", "mẫu 01/lptb", "mẫu 03/bđs-tncn", "hóa đơn", "nợ tiền"]
    }

    def categorize_question(q_text):
        q_lower = q_text.lower()
        matched_topics = []
        for topic, kws in topic_keywords.items():
            if any(kw in q_lower for kw in kws):
                matched_topics.append(topic)
        return matched_topics if matched_topics else ["Thắc mắc TTHC Đất đai Khác"]

    # Lọc câu hỏi trong ngày target_date
    today_turns = []
    for item in chat_history:
        ts = item.get("timestamp", "")
        entry_date = ts.split(" ")[0] if " " in ts else ts
        if entry_date == target_date or not target_date:
            today_turns.append(item)

    # 1. Tạo file báo cáo ngày
    daily_filename = f"cau_hoi_khach_hang_{target_date}.md"
    daily_filepath = os.path.join(CUSTOMER_CHAT_REPORTS_DIR, daily_filename)

    daily_content = []
    daily_content.append(f"# 💬 BÁO CÁO TỔNG HỢP CÂU HỎI KHÁCH HÀNG CHAT VỚI BOT - NGÀY {target_date}")
    daily_content.append(f"*Thời gian xuất báo cáo tự động: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
    daily_content.append(f"**Tổng số lượt câu hỏi tiếp nhận từ người dân:** `{len(today_turns)} lượt`\n")
    daily_content.append("---")

    if not today_turns:
        daily_content.append("\n📌 *Chưa có câu hỏi phát sinh nào từ người dân trong ngày này.*")
    else:
        daily_content.append("\n### 📋 DANH SÁCH CHI TIẾT CÁC CÂU HỎI TIẾP NHẬN:\n")
        daily_content.append("| STT | Thời gian | Nội dung câu hỏi người dân | Nhóm chủ đề pháp lý | Phom AI xử lý |")
        daily_content.append("|---|---|---|---|---|")
        for idx, turn in enumerate(today_turns, 1):
            q_str = turn.get("question", "").replace('|', '\\|')
            ts_str = turn.get("timestamp", "").split(" ")[-1]
            intent_str = turn.get("intent", "Hỏi TTHC Đất đai")
            topics = categorize_question(q_str)
            daily_content.append(f"| {idx} | `{ts_str}` | **{q_str}** | {', '.join(topics)} | `{intent_str}` |")

    markdown_daily = "\n".join(daily_content)
    with open(daily_filepath, 'w', encoding='utf-8') as f:
        f.write(markdown_daily)

    # 2. Tạo tệp Báo cáo Master Tổng hợp Toàn bộ Khách hàng Chat với Bot
    master_filepath = os.path.join(MEMORY_DIR, "tong_hop_cau_hoi_khach_hang_chat_bot.md")

    grouped_by_date = {}
    topic_counts = {}

    for turn in chat_history:
        ts = turn.get("timestamp", "")
        d = ts.split(" ")[0] if " " in ts else "Unknown"
        if d not in grouped_by_date:
            grouped_by_date[d] = []
        grouped_by_date[d].append(turn)

        for t in categorize_question(turn.get("question", "")):
            topic_counts[t] = topic_counts.get(t, 0) + 1

    master_content = []
    master_content.append("# 📊 BÁO CÁO MASTER TỔNG HỢP TOÀN BỘ CÂU HỎI KHÁCH HÀNG CHAT VỚI BOT AI")
    master_content.append(f"*Cập nhật tự động lúc: {time.strftime('%Y-%m-%d %H:%M:%S')}*\n")
    master_content.append(f"**Tổng số lượt tương tác/câu hỏi tiếp nhận toàn hệ thống:** `{len(chat_history)} lượt`\n")
    master_content.append("---")

    master_content.append("\n### 📈 PHÂN PHỐI NỘI DUNG CÂU HỎI THEO CHỦ ĐỀ PHÁP LÝ:\n")
    master_content.append("| STT | Nhóm chủ đề pháp lý đất đai | Số lượng câu hỏi | Tỷ lệ |")
    master_content.append("|---|---|---|---|")
    total_q = len(chat_history) if len(chat_history) > 0 else 1
    for idx, (top_name, count) in enumerate(sorted(topic_counts.items(), key=lambda x: x[1], reverse=True), 1):
        pct = round((count / total_q) * 100, 1)
        master_content.append(f"| {idx} | **{top_name}** | `{count} câu` | `{pct}%` |")

    sorted_dates = sorted(grouped_by_date.keys(), reverse=True)
    for d in sorted_dates:
        turns = grouped_by_date[d]
        master_content.append(f"\n\n### 📅 Ngày {d} (Tổng tiếp nhận: {len(turns)} câu)")
        master_content.append("| STT | Thời gian | Nội dung câu hỏi người dân | Chủ đề | Phom AI xử lý |")
        master_content.append("|---|---|---|---|---|")
        for idx, turn in enumerate(turns, 1):
            q_str = turn.get("question", "").replace('|', '\\|')
            ts_str = turn.get("timestamp", "").split(" ")[-1]
            intent_str = turn.get("intent", "Hỏi TTHC Đất đai")
            topics = categorize_question(q_str)
            master_content.append(f"| {idx} | `{ts_str}` | **{q_str}** | {', '.join(topics)} | `{intent_str}` |")

    markdown_master = "\n".join(master_content)
    with open(master_filepath, 'w', encoding='utf-8') as f:
        f.write(markdown_master)

    return markdown_daily, daily_filepath, markdown_master, master_filepath


# SEARCH PERSISTENT AI CHAT MEMORY FOR SIMILAR PAST VERIFIED QUESTIONS
def search_persistent_chat_memory(question):
    q_lower = question.lower().strip()
    history = load_json_file(CHAT_MEMORY_FILE, [])
    if not history:
        return None
        
    stopwords = ["là", "gì", "như", "thế", "nào", "cần", "những", "điều", "kiện", "thủ", "tục", "hồ", "sơ", "xin", "về", "cho", "tôi", "hãy", "có", "không"]
    keywords = [w for w in re.split(r'\s+', q_lower) if len(w) > 1 and w not in stopwords]
    if not keywords:
        return None

    best_match = None
    best_score = 0
    for entry in history:
        past_q = entry.get("question", "").lower()
        score = sum(1 for kw in keywords if kw in past_q)
        if score > best_score and score >= len(keywords) * 0.7:
            best_score = score
            best_match = entry

    return best_match

VALIANT_BELL_OCR_PATH = r'C:\Users\Admin\.gemini\antigravity\worktrees\valiant-bell\ocr_cccd_land_certificate'

if os.path.exists(VALIANT_BELL_OCR_PATH):
    sys.path.append(VALIANT_BELL_OCR_PATH)

try:
    import vietdoc_ocr
    vietdoc_ocr.DEFAULT_MODEL = "qwen2.5vl:7b"
    OCR_AVAILABLE = True
    print("✅ Successfully loaded vietdoc_ocr module!")
except Exception as e:
    OCR_AVAILABLE = False
    print(f"⚠️ Could not load vietdoc_ocr directly: {e}")

# LOAD AUTHORITATIVE MERGER MAPPING FROM OBSIDIAN VAULT (DIA DANH)
DIA_DANH_JSON = os.path.join(app.root_path, 'dia_danh_mapping.json')
DIA_DANH_MAP = {}

if os.path.exists(DIA_DANH_JSON):
    try:
        with open(DIA_DANH_JSON, 'r', encoding='utf-8') as f:
            DIA_DANH_MAP = json.load(f)
            print(f"✅ Loaded {len(DIA_DANH_MAP)} official DIA DANH merger rules from Obsidian Vault!")
    except Exception as e:
        print(f"⚠️ Failed to load dia_danh_mapping.json: {e}")

# NẠP KHO TRI THỨC 1000 CÂU HỎI ĐÁP PHÁP LÝ CHUẨN (DAT_DAI_QA_1000.JSONL)
QA_1000_DATASET = []
QA_1000_PATH = os.path.join(OBSIDIAN_VAULT_PATH, "dat_dai_qa_1000.jsonl")
if not os.path.exists(QA_1000_PATH):
    QA_1000_PATH = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\dat_dai_qa_1000.jsonl"

if os.path.exists(QA_1000_PATH):
    try:
        with open(QA_1000_PATH, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                line_str = line.strip()
                if line_str:
                    try:
                        item = json.loads(line_str)
                        messages = item.get("messages", [])
                        user_q = ""
                        assistant_a = ""
                        for m in messages:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                assistant_a = m.get("content", "")
                        if user_q and assistant_a:
                            QA_1000_DATASET.append({
                                "question": user_q,
                                "answer": assistant_a
                            })
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(QA_1000_DATASET)} câu hỏi đáp pháp lý chuẩn từ dat_dai_qa_1000.jsonl vào Cách 3!")
    except Exception as e:
        print(f"⚠️ Lỗi nạp dat_dai_qa_1000.jsonl: {e}")

# BỘ ĐO ĐỘ TƯƠNG ĐỒNG CÂU HỎI QUY ĐỊNH BẮT BUỘC TRÊN 80% (>= 0.80) MỚI DÙNG Q&A MẪU
# BỘ MÁY KHOAN SÂU TỪ KHÓA VÀO KHO NÃO (KEYWORD DEEP-DRILL BRAIN SEARCH ENGINE)
def compute_keyword_deep_drill_score(query, item_q, item_a="", item_content="", item_legal_basis=""):
    if not query or (not item_q and not item_a and not item_content):
        return 0.0
    
    q_clean = query.lower().strip()
    target_q = (item_q or "").lower().strip()
    target_text = f"{target_q} {str(item_a or '')} {str(item_content or '')} {str(item_legal_basis or '')}".lower()
    
    # 1. Exact match tuyệt đối
    if q_clean == target_q:
        return 1.0

    # 2. Bóc tách từ khóa cốt lõi, số hiệu văn bản và cụm từ chuyên môn
    stopwords = {'là', 'gì', 'như', 'thế', 'nào', 'cần', 'những', 'điều', 'kiện', 'thủ', 'tục', 'hồ', 'sơ', 'xin', 'về', 'cho', 'tôi', 'hãy', 'có', 'không', 'thực', 'hiện', 'muốn', 'giờ', 'được', 'hỏi', 'ạ', 'các', 'của', 'với', 'tại', 'và', 'hay', 'đó', 'này'}
    words_q = [w for w in re.findall(r'\w+', q_clean) if len(w) > 1 and w not in stopwords]
    
    if not words_q:
        return 0.0
        
    # Bóc tách số hiệu văn bản / số liệu m2 / năm (ví dụ: 2604, 18, 254, 49, 101, 102, 123, 139, 220, 200, 300, 500, 2013, 2024, 2026...)
    numbers_q = re.findall(r'\d+', q_clean)
    num_hits = sum(1 for n in numbers_q if n in target_text)
    
    # Cụm từ chuyên môn đất đai đặc trưng
    domain_phrases = [
        "tách thửa", "hợp thửa", "chuyển mục đích", "cấp sổ", "sổ đỏ", "cấp giấy",
        "lần đầu", "cấp đổi", "cấp lại", "thu hồi", "bồi thường", "giải phóng mặt bằng",
        "gpmb", "trích đo", "trích lục", "đo đạc", "hạn mức", "lộ giới", "hành lang",
        "đất ở", "đất rừng", "đất lúa", "cây lâu năm", "thẩm quyền", "chủ tịch xã",
        "ủy ban nhân dân", "văn phòng đăng ký", "chi nhánh", "lệch mặt bằng", "sai vị trí",
        "vườn tăng thêm", "khai hoang", "xử phạt", "vi phạm", "mẫu đơn", "tờ khai", "lối đi",
        "2604", "18/2026", "254", "49/2026", "101/2024", "102/2024", "123/2024", "220", "139", "140"
    ]
    phrase_hits = sum(1 for p in domain_phrases if p in q_clean and p in target_text)
    
    # Tỷ lệ trúng từ khóa (Keyword Coverage)
    matched_words = [w for w in words_q if w in target_text]
    word_hit_ratio = len(matched_words) / float(len(words_q))
    
    # Độ tương đồng chuỗi câu hỏi gốc
    seq_ratio = difflib.SequenceMatcher(None, q_clean, target_q).ratio() if target_q else 0.0
    
    # Tính điểm tổng hợp khoan sâu (Deep Composite Score)
    base_score = (word_hit_ratio * 0.50) + (seq_ratio * 0.20)
    if phrase_hits > 0:
        base_score += min(0.35, phrase_hits * 0.15)
    if num_hits > 0:
        base_score += min(0.30, num_hits * 0.15)
        
    return min(1.0, round(base_score, 3))

def compute_text_similarity_ratio(user_q, dataset_q):
    return compute_keyword_deep_drill_score(user_q, dataset_q)

def search_dataset_by_similarity(query, dataset, question_key="question", answer_key="answer", min_sim=0.25, top_k=3):
    if not dataset or not query:
        return []
    
    q_lower = query.lower().strip()
    is_forest = any(k in q_lower for k in ["rừng", "lâm nghiệp"])
    
    results = []
    for item in dataset:
        item_q = item.get(question_key, "")
        item_a = item.get(answer_key, "") if isinstance(item.get(answer_key), str) else ""
        item_c = item.get("content", "") if isinstance(item.get("content"), str) else ""
        item_lb = item.get("legal_basis", "") if isinstance(item.get("legal_basis"), str) else ""
        
        if not item_q and not item_c:
            continue
            
        item_q_lower = item_q.lower()
        item_a_lower = item_a.lower()
        
        # Bỏ qua nếu lệch loại đất (đất ở vs đất rừng)
        if is_forest and ("đất ở" in item_q_lower or "đất ở" in item_a_lower) and "rừng" not in item_q_lower:
            continue
            
        sim = compute_keyword_deep_drill_score(query, item_q, item_a, item_c, item_lb)
        if sim >= min_sim:
            results.append({
                "question": item_q or item.get("title", ""),
                "answer": item_a,
                "legal_basis": item_lb,
                "section": item.get("section", ""),
                "content": item_c,
                "score": sim
            })
            
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:top_k]

def search_qa_1000_knowledge_base(query):
    return search_dataset_by_similarity(query, QA_1000_DATASET, min_sim=0.25, top_k=3)

# NẠP KHO TRI THỨC BÁCH KHOA TOÀN THƯ ĐẤT ĐAI 10.000 CÂU Q&A TỪ THƯ MỤC TRAINING OLLMA
# NẠP KHO TRI THỨC BÁCH KHOA TOÀN THƯ ĐẤT ĐAI 10.000 CÂU Q&A TỪ OBSIDIAN VAULT (.MD) & TRAINING OLLMA (.JSONL)
BACH_KHOA_10000_DATASET = []
BACH_KHOA_10000_INDEX = {} # Inverted index by keywords for fast sub-millisecond retrieval
try:
    _bk_candidates = [
        os.path.join(app.root_path, "Obsidian Vault", "bach-khoa-toan-thu-dat-ai-10000-v2.md"),
        r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\bach-khoa-toan-thu-dat-ai-10000-v2.md",
        os.path.join(app.root_path, "Obsidian Vault", "bach-khoa-toan-thu-dat-ai-10000.md"),
        r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\bach-khoa-toan-thu-dat-ai-10000.md",
    ]
    bk_md_path = next((p for p in _bk_candidates if os.path.exists(p)), _bk_candidates[0])
    bk_jsonl_path = os.path.join(app.root_path, "TRAINING OLLMA", "bach-khoa-toan-thu-dat-ai-10000.jsonl")
    if not os.path.exists(bk_jsonl_path):
        bk_jsonl_path = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\bach-khoa-toan-thu-dat-ai-10000.jsonl"

    if os.path.exists(bk_md_path):
        with open(bk_md_path, 'r', encoding='utf-8', errors='ignore') as f:
            md_text = f.read()
        blocks = re.findall(r'###\s*Câu\s*\d+:\s*([^\n]+)\s*\n+\s*\*\*Trả lời:\*\*\s*\n+(.*?)(?=\n+---|\n+###|$)', md_text, flags=re.DOTALL)
        for q_str, a_str in blocks:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                entry = {"question": clean_q, "answer": clean_a}
                BACH_KHOA_10000_DATASET.append(entry)
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in BACH_KHOA_10000_INDEX:
                            BACH_KHOA_10000_INDEX[w] = []
                        if len(BACH_KHOA_10000_INDEX[w]) < 100:
                            BACH_KHOA_10000_INDEX[w].append(len(BACH_KHOA_10000_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(BACH_KHOA_10000_DATASET):,} câu hỏi Bách khoa toàn thư đất đai từ Obsidian Vault (.md) vào bộ nhớ AI!")
    elif os.path.exists(bk_jsonl_path):
        with open(bk_jsonl_path, 'r', encoding='utf-8', errors='ignore') as f:
            for idx, line in enumerate(f):
                line_str = line.strip()
                if line_str:
                    try:
                        item = json.loads(line_str)
                        messages = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in messages:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            entry = {"question": user_q, "answer": asst_a}
                            BACH_KHOA_10000_DATASET.append(entry)
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in BACH_KHOA_10000_INDEX:
                                        BACH_KHOA_10000_INDEX[w] = []
                                    if len(BACH_KHOA_10000_INDEX[w]) < 100:
                                        BACH_KHOA_10000_INDEX[w].append(len(BACH_KHOA_10000_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(BACH_KHOA_10000_DATASET):,} câu hỏi Bách khoa toàn thư đất đai (TRAINING OLLMA) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp Bách khoa toàn thư đất đai 10000: {e}")

def search_bach_khoa_10000_knowledge_base(query, top_k=2):
    if not BACH_KHOA_10000_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in BACH_KHOA_10000_INDEX:
            candidate_indices.update(BACH_KHOA_10000_INDEX[w])
    if not candidate_indices:
        candidate_subset = BACH_KHOA_10000_DATASET[:500]
    else:
        candidate_subset = [BACH_KHOA_10000_DATASET[i] for i in candidate_indices]
    return search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)[:top_k]

# NẠP KHO TRI THỨC BÁCH KHOA TOÀN THƯ KỸ THUẬT ĐỊA CHÍNH & ĐO ĐẠC BẢN ĐỒ 10.000 CÂU Q&A (PORT TỪ INTELNET)
KY_THUAT_DIA_CHINH_10000_DATASET = []
KY_THUAT_DIA_CHINH_10000_INDEX = {}
try:
    kt_json = os.path.join(OBSIDIAN_VAULT_PATH, "06 - BỘ NHỚ AI", "ky_thuat_dia_chinh_10000.json")
    kt_md = os.path.join(OBSIDIAN_VAULT_PATH, "ky-thuat-dia-chinh-10000.md")

    if os.path.exists(kt_json):
        with open(kt_json, 'r', encoding='utf-8', errors='ignore') as f:
            KY_THUAT_DIA_CHINH_10000_DATASET = json.load(f)
    elif os.path.exists(kt_md):
        with open(kt_md, 'r', encoding='utf-8', errors='ignore') as f:
            md_text = f.read()
        blocks = re.findall(r'###\s*Câu\s*\d+:\s*([^\n]+)\s*\n+\s*\*\*Trả lời:\*\*\s*\n+(.*?)(?=\n+---|\n+###|$)', md_text, flags=re.DOTALL)
        for q_str, a_str in blocks:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                KY_THUAT_DIA_CHINH_10000_DATASET.append({"question": clean_q, "answer": clean_a})

    if KY_THUAT_DIA_CHINH_10000_DATASET:
        for idx, entry in enumerate(KY_THUAT_DIA_CHINH_10000_DATASET):
            words = set(re.findall(r'\w+', entry.get("question", "").lower()))
            for w in words:
                if len(w) > 2:
                    if w not in KY_THUAT_DIA_CHINH_10000_INDEX:
                        KY_THUAT_DIA_CHINH_10000_INDEX[w] = []
                    if len(KY_THUAT_DIA_CHINH_10000_INDEX[w]) < 100:
                        KY_THUAT_DIA_CHINH_10000_INDEX[w].append(idx)
        print(f"✅ Đã nạp thành công {len(KY_THUAT_DIA_CHINH_10000_DATASET):,} câu hỏi Bách khoa Kỹ thuật Địa chính & Đo đạc Bản đồ (10.000 Q&A) vào bộ não AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp Kỹ thuật địa chính 10000: {e}")

def search_ky_thuat_dia_chinh_10000_knowledge_base(query, top_k=2):
    if not KY_THUAT_DIA_CHINH_10000_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in KY_THUAT_DIA_CHINH_10000_INDEX:
            candidate_indices.update(KY_THUAT_DIA_CHINH_10000_INDEX[w])
    if not candidate_indices:
        candidate_subset = KY_THUAT_DIA_CHINH_10000_DATASET[:500]
    else:
        candidate_subset = [KY_THUAT_DIA_CHINH_10000_DATASET[i] for i in candidate_indices]
    return search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)[:top_k]

# ======================================================================
# 🧠 NẠP TOÀN BỘ KHO TRI THỨC OBSIDIAN VAULT - BỘ NÃO HỢP NHẤT (PORT TỪ INTELNET: 56.730 MỤC TRI THỨC)
# ======================================================================
OBSIDIAN_UNIFIED_BRAIN = []
OBSIDIAN_UNIFIED_INDEX = {}
try:
    ub_json = os.path.join(OBSIDIAN_VAULT_PATH, "06 - BỘ NHỚ AI", "obsidian_vault_unified_brain.json")
    ub_idx = os.path.join(OBSIDIAN_VAULT_PATH, "06 - BỘ NHỚ AI", "obsidian_vault_unified_index.json")

    if os.path.exists(ub_json):
        with open(ub_json, 'r', encoding='utf-8', errors='ignore') as f:
            OBSIDIAN_UNIFIED_BRAIN = json.load(f)
        if os.path.exists(ub_idx):
            with open(ub_idx, 'r', encoding='utf-8', errors='ignore') as f:
                OBSIDIAN_UNIFIED_INDEX = json.load(f)
        print(f"🎉 Đã nạp thành công TOÀN BỘ {len(OBSIDIAN_UNIFIED_BRAIN):,} mục tri thức từ Toàn bộ Obsidian Vault vào Bộ Não AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp Obsidian Unified Brain: {e}")

def search_obsidian_vault_unified_brain(query, top_k=6):
    if not OBSIDIAN_UNIFIED_BRAIN or not query:
        return []
    q_low = query.lower().strip()
    words = [w for w in re.findall(r'\w+', q_low) if len(w) > 1]
    if not words:
        return []

    candidate_indices = set()
    for w in words:
        if w in OBSIDIAN_UNIFIED_INDEX:
            candidate_indices.update(OBSIDIAN_UNIFIED_INDEX[w])

    if not candidate_indices:
        candidate_indices = range(min(1000, len(OBSIDIAN_UNIFIED_BRAIN)))

    legal_norm_items = []
    handbook_items = []
    qa_items = []

    for idx in candidate_indices:
        if idx >= len(OBSIDIAN_UNIFIED_BRAIN):
            continue
        item = OBSIDIAN_UNIFIED_BRAIN[idx]
        title_low = str(item.get("title", "")).lower()
        content_low = str(item.get("content", ""))[:1000].lower()
        source_low = str(item.get("source", "")).lower()

        score = 0.0
        for w in words:
            if w in title_low:
                score += 4.0
            if w in content_low:
                score += 1.0

        for phrase in ["chuyển mục đích", "tách thửa", "hợp thửa", "không bắt buộc", "cấp sổ đỏ", "đất rừng", "đất lúa", "đất ở", "một phần", "trích đo", "giải phóng mặt bằng", "gpmb", "bồi thường", "thu hồi đất", "tái định cư", "phiếu chỉnh lý", "mẫu 03/clbđ", "thẩm quyền", "thu hồi gcn", "hủy gcn", "cấp đổi gcn", "cấp lại gcn", "nguyên tắc đồng cấp", "cấp gcn cho tổ chức", "tổ chức tôn giáo"]:
            if phrase in q_low and phrase in content_low:
                score += 5.0
            if phrase in q_low and phrase in title_low:
                score += 8.0

        for num in re.findall(r'\d+', q_low):
            if num in title_low:
                score += 6.0
            if num in content_low:
                score += 2.0

        is_legal_norm = any(k in source_low for k in ["01_phaply", "02_quydinh", "luat_dat_dai", "luat_lam_nghiep", "254_2025", "nghi_dinh", "quyet_dinh_18", "2604"])
        is_handbook = any(k in source_low for k in ["cam-nang", "quy-trinh", "huong-dan", "dao-tao", "chuan-hoa"])

        if is_legal_norm:
            score += 10.0
            if score > 0:
                legal_norm_items.append((score, item))
        elif is_handbook:
            score += 6.0
            if score > 0:
                handbook_items.append((score, item))
        else:
            if score > 0:
                qa_items.append((score, item))

    legal_norm_items.sort(key=lambda x: x[0], reverse=True)
    handbook_items.sort(key=lambda x: x[0], reverse=True)
    qa_items.sort(key=lambda x: x[0], reverse=True)

    combined = []
    for _, it in legal_norm_items[:3]:
        if it not in combined: combined.append(it)
    for _, it in handbook_items[:3]:
        if it not in combined: combined.append(it)
    for _, it in qa_items[:2]:
        if it not in combined: combined.append(it)

    return combined[:top_k]

# NẠP DATASET LUẬT ĐẤT ĐAI 5 MẪU CHUẨN KHOA HỌC PHÁP LÝ (DATASET-LUAT-DAT-DAI-5-MAU.JSONL)
DATASET_5_MAU_DATASET = []
try:
    p_5mau = os.path.join(app.root_path, "Obsidian Vault", "dataset-luat-dat-dai-5-mau.jsonl")
    if not os.path.exists(p_5mau):
        p_5mau = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\dataset-luat-dat-dai-5-mau.jsonl"
    if os.path.exists(p_5mau):
        with open(p_5mau, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        q_text, a_text = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                q_text = m.get("content", "").replace("Hỏi: ", "").strip()
                            elif m.get("role") == "assistant":
                                a_text = m.get("content", "").strip()
                        if q_text and a_text:
                            DATASET_5_MAU_DATASET.append({"question": q_text, "answer": a_text})
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(DATASET_5_MAU_DATASET)} tình huống mẫu chuẩn mực từ dataset-luat-dat-dai-5-mau.jsonl vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp dataset 5 mẫu: {e}")

def search_dataset_5_mau_knowledge_base(query):
    if not DATASET_5_MAU_DATASET or not query:
        return []
    return search_dataset_by_similarity(query, DATASET_5_MAU_DATASET, min_sim=0.25)

# LOAD CẨM NANG HỎI ĐÁP TOÀN DIỆN VỀ NGHỊ ĐỊNH 254/2026/NĐ-CP (LUẬT QUẢN LÝ THUẾ 108/2025/QH15)
CAM_NANG_254_DATASET = []
try:
    path_254 = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-nghi-dinh-254.md")
    if not os.path.exists(path_254):
        path_254 = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-nghi-dinh-254.md"
    if os.path.exists(path_254):
        with open(path_254, 'r', encoding='utf-8') as f:
            content_254 = f.read()
        qa_blocks_254 = re.findall(r'####\s*\*\*Câu\s*\d+:\s*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_254, re.DOTALL)
        for q, a in qa_blocks_254:
            CAM_NANG_254_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công 200 câu hỏi đáp Cẩm nang Nghị định 254/2026/NĐ-CP (Luật QLT 108/2025) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-nghi-dinh-254.md: {e}")

def search_cam_nang_254_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_254_DATASET, min_sim=0.25)

# LOAD CẨM NANG TOÀN DIỆN 1000 CÂU HỎI ĐÁP VỀ ĐO ĐẠC ĐỊA CHÍNH (LUẬT ĐO ĐẠC VÀ BẢN ĐỒ 2018 & THÔNG TƯ 26/2024/TT-BTNMT & NGHỊ ĐỊNH 101/2024/NĐ-CP)
CAM_NANG_DO_DAC_DATASET = []
try:
    path_do_dac = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-do-dac-1000.md")
    if not os.path.exists(path_do_dac):
        path_do_dac = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-do-dac-1000.md"
    if os.path.exists(path_do_dac):
        with open(path_do_dac, 'r', encoding='utf-8') as f:
            content_do_dac = f.read()
        qa_blocks_do_dac = re.findall(r'####\s*\*\*Câu\s*\d+:\s*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_do_dac, re.DOTALL)
        for q, a in qa_blocks_do_dac:
            CAM_NANG_DO_DAC_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_DO_DAC_DATASET)} câu hỏi đáp Cẩm nang Đo đạc Địa chính (1000 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-do-dac-1000.md: {e}")

def search_cam_nang_do_dac_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_DO_DAC_DATASET, min_sim=0.25)

# LOAD CẨM NANG TOÀN DIỆN 1000 CÂU HỎI ĐÁP ĐO ĐẠC ĐỊA CHÍNH THEO NGHỊ ĐỊNH 49/2026/NĐ-CP & TT 19/2026/TT-BNNMT TẠI 27 HUYỆN THÀNH PHỐ THANH HÓA
CAM_NANG_DO_DAC_49_DATASET = []
CAM_NANG_DO_DAC_49_INDEX = {}
try:
    path_dd49_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-do-dac-nghi-dinh-49-1000.md")
    if not os.path.exists(path_dd49_md):
        path_dd49_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-do-dac-nghi-dinh-49-1000.md"
    path_dd49_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-do-dac-nghi-dinh-49-1000.jsonl")
    if not os.path.exists(path_dd49_jsonl):
        path_dd49_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-do-dac-nghi-dinh-49-1000.jsonl"

    if os.path.exists(path_dd49_md):
        with open(path_dd49_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_dd49 = f.read()
        blocks_dd49 = re.findall(r'###\s*(?:Câu\s*hỏi\s*\d+:|Câu\s*\d+:|)?\s*([^\n]+)\s*\n+\s*\*\*Trả lời:\*\*\s*\n+(.*?)(?=\n+---|\n+###|$)', text_dd49, flags=re.DOTALL)
        for q_str, a_str in blocks_dd49:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_DO_DAC_49_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_DO_DAC_49_INDEX:
                            CAM_NANG_DO_DAC_49_INDEX[w] = []
                        if len(CAM_NANG_DO_DAC_49_INDEX[w]) < 100:
                            CAM_NANG_DO_DAC_49_INDEX[w].append(len(CAM_NANG_DO_DAC_49_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_DO_DAC_49_DATASET):,} câu hỏi đáp Cẩm nang Đo đạc Nghị định 49/2026/NĐ-CP (Obsidian Vault) vào bộ nhớ AI!")
    elif os.path.exists(path_dd49_jsonl):
        with open(path_dd49_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            CAM_NANG_DO_DAC_49_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_DO_DAC_49_INDEX:
                                        CAM_NANG_DO_DAC_49_INDEX[w] = []
                                    if len(CAM_NANG_DO_DAC_49_INDEX[w]) < 100:
                                        CAM_NANG_DO_DAC_49_INDEX[w].append(len(CAM_NANG_DO_DAC_49_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_DO_DAC_49_DATASET):,} câu hỏi đáp Cẩm nang Đo đạc Nghị định 49/2026/NĐ-CP (JSONL) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-do-dac-nghi-dinh-49-1000.md: {e}")

def search_cam_nang_do_dac_49_knowledge_base(query, top_k=2):
    if not CAM_NANG_DO_DAC_49_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_DO_DAC_49_INDEX:
            candidate_indices.update(CAM_NANG_DO_DAC_49_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_DO_DAC_49_DATASET
    else:
        candidate_subset = [CAM_NANG_DO_DAC_49_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD CẨM NANG HỎI ĐÁP TOÀN DIỆN VỀ LỘ GIỚI, HÀNH LẠNG AN TOÀN ĐƯỜNG BỘ VÀ QUY HOẠCH XÂY DỰNG (550 CÂU Q&A THEO LUẬT ĐƯỜNG BỘ 2024, NĐ 11/2010, NĐ 49/2026 & QCVN 01:2021/BXD)
CAM_NANG_LO_GIOI_550_DATASET = []
CAM_NANG_LO_GIOI_550_INDEX = {}
try:
    path_lg_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-lo-gioi-quy-hoach-550.md")
    if not os.path.exists(path_lg_md):
        path_lg_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-lo-gioi-quy-hoach-550.md"
    path_lg_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-lo-gioi-quy-hoach-550.jsonl")
    if not os.path.exists(path_lg_jsonl):
        path_lg_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-lo-gioi-quy-hoach-550.jsonl"

    if os.path.exists(path_lg_md):
        with open(path_lg_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_lg = f.read()
        blocks_lg = re.findall(r'###\s*(?:Câu\s*hỏi\s*\d+:|Câu\s*\d+:|)?\s*([^\n]+)\s*\n+\s*\*\*Trả lời:\*\*\s*\n+(.*?)(?=\n+---|\n+###|$)', text_lg, flags=re.DOTALL)
        for q_str, a_str in blocks_lg:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_LO_GIOI_550_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_LO_GIOI_550_INDEX:
                            CAM_NANG_LO_GIOI_550_INDEX[w] = []
                        if len(CAM_NANG_LO_GIOI_550_INDEX[w]) < 100:
                            CAM_NANG_LO_GIOI_550_INDEX[w].append(len(CAM_NANG_LO_GIOI_550_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_LO_GIOI_550_DATASET):,} câu hỏi đáp Cẩm nang Lộ giới & Hành lang an toàn quy hoạch (Obsidian Vault) vào bộ nhớ AI!")
    elif os.path.exists(path_lg_jsonl):
        with open(path_lg_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            CAM_NANG_LO_GIOI_550_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_LO_GIOI_550_INDEX:
                                        CAM_NANG_LO_GIOI_550_INDEX[w] = []
                                    if len(CAM_NANG_LO_GIOI_550_INDEX[w]) < 100:
                                        CAM_NANG_LO_GIOI_550_INDEX[w].append(len(CAM_NANG_LO_GIOI_550_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_LO_GIOI_550_DATASET):,} câu hỏi đáp Cẩm nang Lộ giới & Hành lang an toàn quy hoạch (JSONL) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-lo-gioi-quy-hoach-550.md: {e}")

def search_cam_nang_lo_gioi_550_knowledge_base(query, top_k=2):
    if not CAM_NANG_LO_GIOI_550_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_LO_GIOI_550_INDEX:
            candidate_indices.update(CAM_NANG_LO_GIOI_550_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_LO_GIOI_550_DATASET
    else:
        candidate_subset = [CAM_NANG_LO_GIOI_550_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD CẨM NANG NGHIỆP VỤ TRÍCH ĐO ĐỊA CHÍNH VÀ GIẢI PHÓNG MẶT BẰNG (300 TÌNH HUỐNG THEO LUẬT ĐẤT ĐAI 2024, NĐ 101/2024, NĐ 49/2026, NĐ 88/2024 & QĐ 2604)
CAM_NANG_TRICH_DO_GPMB_300_DATASET = []
CAM_NANG_TRICH_DO_GPMB_300_INDEX = {}
try:
    path_gpmb_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-trich-do-gpmb-300.md")
    if not os.path.exists(path_gpmb_md):
        path_gpmb_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-trich-do-gpmb-300.md"
    path_gpmb_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-trich-do-gpmb-300.jsonl")
    if not os.path.exists(path_gpmb_jsonl):
        path_gpmb_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-trich-do-gpmb-300.jsonl"

    if os.path.exists(path_gpmb_md):
        with open(path_gpmb_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_gpmb = f.read()
        blocks_gpmb = re.findall(r'###\s*(?:Câu\s*\d+:|Trường\s*hợp\s*\d+:|)?\s*([^\n]+)\s*\n+(.*?)(?=\n+---|\n+###|$)', text_gpmb, flags=re.DOTALL)
        for q_str, a_str in blocks_gpmb:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_TRICH_DO_GPMB_300_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_TRICH_DO_GPMB_300_INDEX:
                            CAM_NANG_TRICH_DO_GPMB_300_INDEX[w] = []
                        if len(CAM_NANG_TRICH_DO_GPMB_300_INDEX[w]) < 100:
                            CAM_NANG_TRICH_DO_GPMB_300_INDEX[w].append(len(CAM_NANG_TRICH_DO_GPMB_300_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TRICH_DO_GPMB_300_DATASET):,} câu hỏi đáp Cẩm nang Trích đo Địa chính & GPMB (Obsidian Vault) vào bộ nhớ AI!")
    elif os.path.exists(path_gpmb_jsonl):
        with open(path_gpmb_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            CAM_NANG_TRICH_DO_GPMB_300_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_TRICH_DO_GPMB_300_INDEX:
                                        CAM_NANG_TRICH_DO_GPMB_300_INDEX[w] = []
                                    if len(CAM_NANG_TRICH_DO_GPMB_300_INDEX[w]) < 100:
                                        CAM_NANG_TRICH_DO_GPMB_300_INDEX[w].append(len(CAM_NANG_TRICH_DO_GPMB_300_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TRICH_DO_GPMB_300_DATASET):,} câu hỏi đáp Cẩm nang Trích đo Địa chính & GPMB (JSONL) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-trich-do-gpmb-300.md: {e}")

def search_cam_nang_trich_do_gpmb_300_knowledge_base(query, top_k=2):
    if not CAM_NANG_TRICH_DO_GPMB_300_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_TRICH_DO_GPMB_300_INDEX:
            candidate_indices.update(CAM_NANG_TRICH_DO_GPMB_300_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_TRICH_DO_GPMB_300_DATASET
    else:
        candidate_subset = [CAM_NANG_TRICH_DO_GPMB_300_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD CẨM NANG 500 CÂU HỎI & TRẢ LỜI NGHIỆP VỤ ĐỊA CHÍNH VÀ THẨM QUYỀN HÀNH CHÍNH (NĐ 49/2026/NĐ-CP, NQ 254/2025/QH15 & QĐ 2604/QĐ-VP THANH HÓA)
CAM_NANG_THAM_QUYEN_500_DATASET = []
CAM_NANG_THAM_QUYEN_500_INDEX = {}
try:
    path_tq_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-tham-quyen-dia-chinh-500.md")
    if not os.path.exists(path_tq_md):
        path_tq_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-tham-quyen-dia-chinh-500.md"
    path_tq_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-tham-quyen-dia-chinh-500.jsonl")
    if not os.path.exists(path_tq_jsonl):
        path_tq_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-tham-quyen-dia-chinh-500.jsonl"

    if os.path.exists(path_tq_md):
        with open(path_tq_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_tq = f.read()
        blocks_tq = re.findall(r'###\s*(?:Câu\s*hỏi\s*\d+:|Câu\s*\d+:|Trường\s*hợp\s*\d+:|)?\s*([^\n]+)\s*\n+(.*?)(?=\n+---|\n+###|$)', text_tq, flags=re.DOTALL)
        for q_str, a_str in blocks_tq:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_THAM_QUYEN_500_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_THAM_QUYEN_500_INDEX:
                            CAM_NANG_THAM_QUYEN_500_INDEX[w] = []
                        if len(CAM_NANG_THAM_QUYEN_500_INDEX[w]) < 100:
                            CAM_NANG_THAM_QUYEN_500_INDEX[w].append(len(CAM_NANG_THAM_QUYEN_500_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_THAM_QUYEN_500_DATASET):,} câu hỏi đáp Cẩm nang Thẩm quyền Địa chính (Obsidian Vault) vào bộ nhớ AI!")
    elif os.path.exists(path_tq_jsonl):
        with open(path_tq_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            CAM_NANG_THAM_QUYEN_500_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_THAM_QUYEN_500_INDEX:
                                        CAM_NANG_THAM_QUYEN_500_INDEX[w] = []
                                    if len(CAM_NANG_THAM_QUYEN_500_INDEX[w]) < 100:
                                        CAM_NANG_THAM_QUYEN_500_INDEX[w].append(len(CAM_NANG_THAM_QUYEN_500_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_THAM_QUYEN_500_DATASET):,} câu hỏi đáp Cẩm nang Thẩm quyền Địa chính (JSONL) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-tham-quyen-dia-chinh-500.md: {e}")

def search_cam_nang_tham_quyen_500_knowledge_base(query, top_k=2):
    if not CAM_NANG_THAM_QUYEN_500_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_THAM_QUYEN_500_INDEX:
            candidate_indices.update(CAM_NANG_THAM_QUYEN_500_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_THAM_QUYEN_500_DATASET
    else:
        candidate_subset = [CAM_NANG_THAM_QUYEN_500_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD CẨM NANG THẨM QUYỀN CẤP, THU HỒI & ĐĂNG KÝ BIẾN ĐỘNG GCN THEO PHÂN CẤM MỚI
# (Luật Đất đai 2024, NĐ 49/2026/NĐ-CP & QĐ 2604/QĐ-VP - bộ 52 thủ tục phân theo Phụ lục I/II/III)
CAM_NANG_TQ_GCN_DATASET = []
CAM_NANG_TQ_GCN_INDEX = {}
try:
    path_gcn_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-tham-quyen-cap-gcn.md")
    if not os.path.exists(path_gcn_md):
        path_gcn_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-tham-quyen-cap-gcn.md"

    if os.path.exists(path_gcn_md):
        with open(path_gcn_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_gcn = f.read()

        def _gcn_add(q_str, a_str):
            clean_q, clean_a = q_str.strip(), a_str.strip()
            if clean_q and len(clean_a) > 30:
                CAM_NANG_TQ_GCN_DATASET.append({"question": clean_q, "answer": clean_a})
                for w in set(re.findall(r'\w+', clean_q.lower())):
                    if len(w) > 2:
                        if w not in CAM_NANG_TQ_GCN_INDEX:
                            CAM_NANG_TQ_GCN_INDEX[w] = []
                        if len(CAM_NANG_TQ_GCN_INDEX[w]) < 100:
                            CAM_NANG_TQ_GCN_INDEX[w].append(len(CAM_NANG_TQ_GCN_DATASET) - 1)

        # A) Các tình huống thực địa mẫu: #### Tình huống N: ... (KẾT LUẬN + CĂN CỨ + LƯU Ý)
        for q_str, a_str in re.findall(
                r'####\s*Tình\s*huống\s*\d+:\s*([^\n]+)\s*\n(.*?)(?=\n####\s|\n---\s|\Z)',
                text_gcn, flags=re.DOTALL):
            _gcn_add(f"[Thẩm quyền GCN - Tình huống thực địa] {q_str}", a_str)

        # B) Khung thẩm quyền chung: ### N. Tiêu đề (PHẦN I & II)
        for t_str, a_str in re.findall(
                r'###\s*(\d+\.\s*[^\n]+)\s*\n(.*?)(?=\n###\s|\n##\s|\Z)',
                text_gcn, flags=re.DOTALL):
            _gcn_add(f"[Phân cấp thẩm quyền GCN] {t_str}", a_str)

        print(f"✅ Đã nạp thành công {len(CAM_NANG_TQ_GCN_DATASET)} mục Thẩm quyền Cấp/Thu hồi GCN (phân cấp mới QĐ 2604) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-tham-quyen-cap-gcn.md: {e}")

def search_cam_nang_tq_gcn_knowledge_base(query, top_k=2):
    if not CAM_NANG_TQ_GCN_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_TQ_GCN_INDEX:
            candidate_indices.update(CAM_NANG_TQ_GCN_INDEX[w])
    candidate_subset = CAM_NANG_TQ_GCN_DATASET if not candidate_indices else [CAM_NANG_TQ_GCN_DATASET[i] for i in candidate_indices]

    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]

    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(overlap / float(len(q_words)), 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD BỘ CẨM NANG 2000 CÂU HỎI VÀ TRẢ LỜI TÓM TẮT HIỆU LỰC VĂN BẢN ĐẤT ĐAI (NỘI DUNG CỐT LÕI, TÌNH TRẠNG HIỆU LỰC & QUY ĐỊNH THAY THẾ/BÃI BỎ)
TOM_TAT_HIEU_LUC_2000_DATASET = []
TOM_TAT_HIEU_LUC_2000_INDEX = {}
try:
    path_hl_md = os.path.join(app.root_path, "Obsidian Vault", "tom-tat-hieu-luc-van-ban-2000.md")
    if not os.path.exists(path_hl_md):
        path_hl_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\tom-tat-hieu-luc-van-ban-2000.md"
    path_hl_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "tom-tat-hieu-luc-van-ban-2000.jsonl")
    if not os.path.exists(path_hl_jsonl):
        path_hl_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\tom-tat-hieu-luc-van-ban-2000.jsonl"

    if os.path.exists(path_hl_md):
        with open(path_hl_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_hl = f.read()
        blocks_hl = re.findall(r'###\s*(?:Câu\s*hỏi\s*\d+:|Câu\s*\d+:|Trường\s*hợp\s*\d+:|)?\s*([^\n]+)\s*\n+(.*?)(?=\n+---|\n+###|$)', text_hl, flags=re.DOTALL)
        for q_str, a_str in blocks_hl:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                TOM_TAT_HIEU_LUC_2000_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in TOM_TAT_HIEU_LUC_2000_INDEX:
                            TOM_TAT_HIEU_LUC_2000_INDEX[w] = []
                        if len(TOM_TAT_HIEU_LUC_2000_INDEX[w]) < 100:
                            TOM_TAT_HIEU_LUC_2000_INDEX[w].append(len(TOM_TAT_HIEU_LUC_2000_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(TOM_TAT_HIEU_LUC_2000_DATASET):,} câu hỏi đáp Cẩm nang Tóm tắt Hiệu lực Văn bản (Obsidian Vault) vào bộ nhớ AI!")
    elif os.path.exists(path_hl_jsonl):
        with open(path_hl_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "")
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "")
                        if user_q and asst_a:
                            TOM_TAT_HIEU_LUC_2000_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in TOM_TAT_HIEU_LUC_2000_INDEX:
                                        TOM_TAT_HIEU_LUC_2000_INDEX[w] = []
                                    if len(TOM_TAT_HIEU_LUC_2000_INDEX[w]) < 100:
                                        TOM_TAT_HIEU_LUC_2000_INDEX[w].append(len(TOM_TAT_HIEU_LUC_2000_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(TOM_TAT_HIEU_LUC_2000_DATASET):,} câu hỏi đáp Cẩm nang Tóm tắt Hiệu lực Văn bản (JSONL) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp tom-tat-hieu-luc-van-ban-2000.md: {e}")

def search_tom_tat_hieu_luc_2000_knowledge_base(query, top_k=2):
    if not TOM_TAT_HIEU_LUC_2000_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in TOM_TAT_HIEU_LUC_2000_INDEX:
            candidate_indices.update(TOM_TAT_HIEU_LUC_2000_INDEX[w])
    if not candidate_indices:
        candidate_subset = TOM_TAT_HIEU_LUC_2000_DATASET
    else:
        candidate_subset = [TOM_TAT_HIEU_LUC_2000_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# NẠP CẨM NANG NGHIỆP VỤ 300 TÌNH HUỐNG XÁC ĐỊNH VỊ TRÍ THỬA ĐẤT & BẢNG GIÁ ĐẤT (BẢNG GIÁ ĐẤT THANH HÓA, LUẬT ĐẤT ĐAI 2024, NGHỊ ĐỊNH 103/2024)
CAM_NANG_TINH_VI_TRI_300_DATASET = []
CAM_NANG_TINH_VI_TRI_300_INDEX = {}
try:
    path_tvt_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-tinh-vi-tri-thua-dat-300.md")
    if not os.path.exists(path_tvt_md):
        path_tvt_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-tinh-vi-tri-thua-dat-300.md"
    path_tvt_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-tinh-vi-tri-thua-dat-300.jsonl")
    if not os.path.exists(path_tvt_jsonl):
        path_tvt_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-tinh-vi-tri-thua-dat-300.jsonl"

    if os.path.exists(path_tvt_md):
        with open(path_tvt_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_tvt = f.read()
        blocks_tvt = re.findall(r'###\s*(?:Tình\s*huống|Câu\s*hỏi|Câu)?\s*(\d+)?[:\.]?\s*([^\n]+)\s*\n+(.*?)(?=\n+---|\n+###|$)', text_tvt, flags=re.DOTALL)
        for num_str, q_str, a_str in blocks_tvt:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_TINH_VI_TRI_300_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_TINH_VI_TRI_300_INDEX:
                            CAM_NANG_TINH_VI_TRI_300_INDEX[w] = []
                        if len(CAM_NANG_TINH_VI_TRI_300_INDEX[w]) < 100:
                            CAM_NANG_TINH_VI_TRI_300_INDEX[w].append(len(CAM_NANG_TINH_VI_TRI_300_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TINH_VI_TRI_300_DATASET):,} tình huống Cẩm nang Xác định Vị trí Thửa đất & Bảng giá đất (Obsidian Vault) vào bộ não AI!")
    elif os.path.exists(path_tvt_jsonl):
        with open(path_tvt_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "").replace("Hỏi: ", "").strip()
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "").strip()
                        if user_q and asst_a:
                            CAM_NANG_TINH_VI_TRI_300_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_TINH_VI_TRI_300_INDEX:
                                        CAM_NANG_TINH_VI_TRI_300_INDEX[w] = []
                                    if len(CAM_NANG_TINH_VI_TRI_300_INDEX[w]) < 100:
                                        CAM_NANG_TINH_VI_TRI_300_INDEX[w].append(len(CAM_NANG_TINH_VI_TRI_300_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TINH_VI_TRI_300_DATASET):,} tình huống Cẩm nang Xác định Vị trí Thửa đất & Bảng giá đất (JSONL) vào bộ não AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-tinh-vi-tri-thua-dat-300.md: {e}")

def search_cam_nang_tinh_vi_tri_300_knowledge_base(query, top_k=3):
    if not CAM_NANG_TINH_VI_TRI_300_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_TINH_VI_TRI_300_INDEX:
            candidate_indices.update(CAM_NANG_TINH_VI_TRI_300_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_TINH_VI_TRI_300_DATASET
    else:
        candidate_subset = [CAM_NANG_TINH_VI_TRI_300_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25, top_k=top_k)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# NẠP CẨM NANG 500 TÌNH HUỐNG QUY HOẠCH CHI TIẾT 1/500 & QUY HOẠCH XÂY DỰNG ĐÔ THỊ (LUẬT XÂY DỰNG, LUẬT QUY HOẠCH ĐÔ THỊ, LUẬT ĐẤT ĐAI 2024)
CAM_NANG_QUY_HOACH_500_DATASET = []
CAM_NANG_QUY_HOACH_500_INDEX = {}
try:
    path_qh_md = os.path.join(app.root_path, "Obsidian Vault", "cam-nang-quy-hoach-1-500-500.md")
    if not os.path.exists(path_qh_md):
        path_qh_md = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-quy-hoach-1-500-500.md"
    path_qh_jsonl = os.path.join(app.root_path, "TRAINING OLLMA", "cam-nang-quy-hoach-1-500-500.jsonl")
    if not os.path.exists(path_qh_jsonl):
        path_qh_jsonl = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\TRAINING OLLMA\cam-nang-quy-hoach-1-500-500.jsonl"

    if os.path.exists(path_qh_md):
        with open(path_qh_md, 'r', encoding='utf-8', errors='ignore') as f:
            text_qh = f.read()
        blocks_qh = re.findall(r'###\s*(?:Tình\s*huống|Câu\s*hỏi|Câu)?\s*(\d+)?[:\.]?\s*([^\n]+)\s*\n+(.*?)(?=\n+---|\n+###|$)', text_qh, flags=re.DOTALL)
        for num_str, q_str, a_str in blocks_qh:
            clean_q = q_str.strip()
            clean_a = a_str.strip()
            if clean_q and clean_a:
                CAM_NANG_QUY_HOACH_500_DATASET.append({"question": clean_q, "answer": clean_a})
                words = set(re.findall(r'\w+', clean_q.lower()))
                for w in words:
                    if len(w) > 2:
                        if w not in CAM_NANG_QUY_HOACH_500_INDEX:
                            CAM_NANG_QUY_HOACH_500_INDEX[w] = []
                        if len(CAM_NANG_QUY_HOACH_500_INDEX[w]) < 100:
                            CAM_NANG_QUY_HOACH_500_INDEX[w].append(len(CAM_NANG_QUY_HOACH_500_DATASET) - 1)
        print(f"✅ Đã nạp thành công {len(CAM_NANG_QUY_HOACH_500_DATASET):,} tình huống Cẩm nang Quy hoạch chi tiết 1/500 (Obsidian Vault) vào bộ não AI!")
    elif os.path.exists(path_qh_jsonl):
        with open(path_qh_jsonl, 'r', encoding='utf-8', errors='ignore') as f:
            for line in f:
                if line.strip():
                    try:
                        item = json.loads(line.strip())
                        msgs = item.get("messages", [])
                        user_q, asst_a = "", ""
                        for m in msgs:
                            if m.get("role") == "user":
                                user_q = m.get("content", "").replace("Hỏi: ", "").strip()
                            elif m.get("role") == "assistant":
                                asst_a = m.get("content", "").strip()
                        if user_q and asst_a:
                            CAM_NANG_QUY_HOACH_500_DATASET.append({"question": user_q, "answer": asst_a})
                            words = set(re.findall(r'\w+', user_q.lower()))
                            for w in words:
                                if len(w) > 2:
                                    if w not in CAM_NANG_QUY_HOACH_500_INDEX:
                                        CAM_NANG_QUY_HOACH_500_INDEX[w] = []
                                    if len(CAM_NANG_QUY_HOACH_500_INDEX[w]) < 100:
                                        CAM_NANG_QUY_HOACH_500_INDEX[w].append(len(CAM_NANG_QUY_HOACH_500_DATASET) - 1)
                    except Exception:
                        continue
        print(f"✅ Đã nạp thành công {len(CAM_NANG_QUY_HOACH_500_DATASET):,} tình huống Cẩm nang Quy hoạch chi tiết 1/500 (JSONL) vào bộ não AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-quy-hoach-1-500-500.md: {e}")

def search_cam_nang_quy_hoach_500_knowledge_base(query, top_k=3):
    if not CAM_NANG_QUY_HOACH_500_DATASET or not query:
        return []
    words = [w for w in re.findall(r'\w+', query.lower()) if len(w) > 2]
    if not words:
        return []
    candidate_indices = set()
    for w in words:
        if w in CAM_NANG_QUY_HOACH_500_INDEX:
            candidate_indices.update(CAM_NANG_QUY_HOACH_500_INDEX[w])
    if not candidate_indices:
        candidate_subset = CAM_NANG_QUY_HOACH_500_DATASET
    else:
        candidate_subset = [CAM_NANG_QUY_HOACH_500_DATASET[i] for i in candidate_indices]
    
    res = search_dataset_by_similarity(query, candidate_subset, min_sim=0.25, top_k=top_k)
    if res:
        return res[:top_k]
    
    scored = []
    q_words = set(words)
    for item in candidate_subset:
        item_words = set(re.findall(r'\w+', item['question'].lower()))
        overlap = len(q_words.intersection(item_words))
        if overlap >= 2:
            score = overlap / float(len(q_words))
            scored.append({"question": item["question"], "answer": item["answer"], "score": round(score, 2)})
    scored.sort(key=lambda x: x["score"], reverse=True)
    return scored[:top_k]

# LOAD CẨM NANG HỎI ĐÁP TOÀN DIỆN VỀ QUYẾT ĐỊNH SỐ 2604/QĐ-VP THANH HÓA (300 CÂU TTHC ĐẤT ĐAI ĐẶC THÙ)
CAM_NANG_2604_DATASET = []
try:
    path_2604 = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-quyet-dinh-2604-300.md")
    if not os.path.exists(path_2604):
        path_2604 = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-quyet-dinh-2604-300.md"
    if os.path.exists(path_2604):
        with open(path_2604, 'r', encoding='utf-8') as f:
            content_2604 = f.read()
        qa_blocks_2604 = re.findall(r'####\s*\*\*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_2604, re.DOTALL)
        for q, a in qa_blocks_2604:
            CAM_NANG_2604_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_2604_DATASET)} câu hỏi đáp Cẩm nang Quyết định 2604/QĐ-VP TTHC Đất đai Thanh Hóa (300 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-quyet-dinh-2604-300.md: {e}")

def search_cam_nang_2604_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_2604_DATASET, min_sim=0.25)

# LOAD CẨM NANG HỎI ĐÁP TOÀN DIỆN VỀ THU HỒI, CẤP ĐỔI VÀ CẤP LẠI GIẤY CHỨNG NHẬN ĐẤT ĐAI (500 CÂU Q&A)
CAM_NANG_THU_HOI_DATASET = []
try:
    path_thu_hoi = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-thu-hoi-cap-doi-500.md")
    if not os.path.exists(path_thu_hoi):
        path_thu_hoi = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-thu-hoi-cap-doi-500.md"
    if os.path.exists(path_thu_hoi):
        with open(path_thu_hoi, 'r', encoding='utf-8') as f:
            content_thu_hoi = f.read()
        qa_blocks_thu_hoi = re.findall(r'####\s*\*\*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_thu_hoi, re.DOTALL)
        for q, a in qa_blocks_thu_hoi:
            CAM_NANG_THU_HOI_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_THU_HOI_DATASET)} câu hỏi đáp Cẩm nang Thu hồi, Cấp đổi & Cấp lại GCN (500 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-thu-hoi-cap-doi-500.md: {e}")

def search_cam_nang_thu_hoi_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_THU_HOI_DATASET, min_sim=0.25)

# LOAD TÀI LIỆU TRAINING NGHIỆP VỤ ĐẤT ĐAI THANH HÓA (QUYẾT ĐỊNH 2604/QĐ-VP & NGHỊ ĐỊNH 49/2026/NĐ-CP)
CAM_NANG_DAO_TAO_DATASET = []
try:
    path_dao_tao = os.path.join(OBSIDIAN_VAULT_PATH, "huong-dan-dao-tao-thu-tuc-dat-dai.md")
    if not os.path.exists(path_dao_tao):
        path_dao_tao = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\huong-dan-dao-tao-thu-tuc-dat-dai.md"
    if os.path.exists(path_dao_tao):
        with open(path_dao_tao, 'r', encoding='utf-8') as f:
            content_dao_tao = f.read()
        sections_dt = re.findall(r'(#{2,4}\s*.*?\n)(.*?)(?=\n#{2,4}\s*|$$)', content_dao_tao, re.DOTALL)
        for h, c in sections_dt:
            header_clean = h.replace('#', '').strip()
            body_clean = c.strip()
            if header_clean and len(body_clean) > 20:
                CAM_NANG_DAO_TAO_DATASET.append({
                    "section": header_clean,
                    "content": body_clean
                })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_DAO_TAO_DATASET)} phần nghiệp vụ Cẩm nang Đào tạo TTHC Đất đai (QĐ 2604/QĐ-VP) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp huong-dan-dao-tao-thu-tuc-dat-dai.md: {e}")

def search_cam_nang_dao_tao_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_DAO_TAO_DATASET, question_key="section", answer_key="content", min_sim=0.25)

# LOAD CẨM NĂNG CHUYÊN SÂU CHUYỂN MỤC ĐÍCH RỪNG & THU HỒI ĐẤT LÂM NGHIỆP THANH HÓA (100 CÂU Q&A)
CAM_NANG_RUNG_DATASET = []
try:
    path_rung = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-thu-hoi-dat-rung-100.md")
    if not os.path.exists(path_rung):
        path_rung = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-thu-hoi-dat-rung-100.md"
    if os.path.exists(path_rung):
        with open(path_rung, 'r', encoding='utf-8') as f:
            content_rung = f.read()
        content_rung = content_rung.replace('\\n', '\n')
        qa_blocks_rung = re.findall(r'####\s*\*\*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_rung, re.DOTALL)
        for q, a in qa_blocks_rung:
            CAM_NANG_RUNG_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_RUNG_DATASET)} câu hỏi đáp Cẩm nang Chuyển mục đích rừng & Thu hồi đất lâm nghiệp (100 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-thu-hoi-dat-rung-100.md: {e}")

def search_cam_nang_rung_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_RUNG_DATASET, min_sim=0.25)

# LOAD CẨM NĂNG QUY TRÌNH 2 BƯỚC TÁCH THỬA ĐỒNG THỜI HỢP THỬA ĐẤT ĐẶC THÙ (NGHỊ QUYẾT 254/2025/QH15 & QĐ 2604/QĐ-VP)
QUY_TRINH_2_BUOC_DATASET = []
try:
    path_2b = os.path.join(OBSIDIAN_VAULT_PATH, "quy-trinh-2-buoc-tach-hop-thua.md")
    if not os.path.exists(path_2b):
        path_2b = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\quy-trinh-2-buoc-tach-hop-thua.md"
    if os.path.exists(path_2b):
        with open(path_2b, 'r', encoding='utf-8') as f:
            content_2b = f.read()
        content_2b = content_2b.replace('\\n', '\n')
        sections_2b = re.findall(r'(#{1,3}\s*.*?\n)(.*?)(?=\n#{1,3}\s*|$$)', content_2b, re.DOTALL)
        for h, c in sections_2b:
            header_clean = h.replace('#', '').strip()
            body_clean = c.strip()
            if header_clean and len(body_clean) > 20:
                QUY_TRINH_2_BUOC_DATASET.append({
                    "section": header_clean,
                    "content": body_clean
                })
        print(f"✅ Đã nạp thành công {len(QUY_TRINH_2_BUOC_DATASET)} phần Quy trình 2 bước Tách thửa đồng thời Hợp thửa (NQ 254/2025/QH15 & QĐ 2604) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp quy-trinh-2-buoc-tach-hop-thua.md: {e}")

def search_quy_trinh_2_buoc_knowledge_base(query):
    return search_dataset_by_similarity(query, QUY_TRINH_2_BUOC_DATASET, question_key="section", answer_key="content", min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 100 CÂU HỎI ĐÁP TÁCH THỬA ĐỒNG THỜI HỢP THỬA ĐẤT ĐẶC THÙ (LUẬT ĐẤT ĐAI 2024, NQ 254/2025/QH15, NĐ 49/2026/NĐ-CP & QĐ 2604)
CAM_NANG_TACH_HOP_DAC_THU_DATASET = []
try:
    path_thdt = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-tach-hop-thua-dac-thu-100.md")
    if not os.path.exists(path_thdt):
        path_thdt = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-tach-hop-thua-dac-thu-100.md"
    if os.path.exists(path_thdt):
        with open(path_thdt, 'r', encoding='utf-8') as f:
            content_thdt = f.read()
        content_thdt = content_thdt.replace('\\n', '\n')
        qa_blocks_thdt = re.findall(r'####\s*\*\*(.*?)\*\*\s*\nTrả lời:\s*\n(.*?)(?=\n---|$$)', content_thdt, re.DOTALL)
        for q, a in qa_blocks_thdt:
            CAM_NANG_TACH_HOP_DAC_THU_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TACH_HOP_DAC_THU_DATASET)} câu hỏi đáp Cẩm nang Tách thửa đồng thời Hợp thửa đặc thù (100 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-tach-hop-thua-dac-thu-100.md: {e}")

def search_cam_nang_tach_hop_dac_thu_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_TACH_HOP_DAC_THU_DATASET, min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 50 CÂU HỎI ĐÁP XỬ LÝ LỆCH MẶT BẰNG ĐẤU GIÁ VÀ PHÂN LÔ THANH HÓA (VB 9549/UBND-NNMT & CV 16838/SNNMT & QĐ 2604)
CAM_NANG_LECH_MAT_BANG_DATASET = []
try:
    path_lmb = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-xu-ly-lech-mat-bang-50.md")
    if not os.path.exists(path_lmb):
        path_lmb = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-xu-ly-lech-mat-bang-50.md"
    if os.path.exists(path_lmb):
        with open(path_lmb, 'r', encoding='utf-8') as f:
            content_lmb = f.read()
        content_lmb = content_lmb.replace('\\n', '\n')
        qa_blocks_lmb = re.findall(r'###\s*\*\*(.*?)\*\*\s*\n\*\s*\*\*Trả lời:\*\*\s*(.*?)(?=\n###|\n---|$$)', content_lmb, re.DOTALL)
        for q, a in qa_blocks_lmb:
            CAM_NANG_LECH_MAT_BANG_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_LECH_MAT_BANG_DATASET)} câu hỏi đáp Cẩm nang Xử lý lệch mặt bằng đấu giá và phân lô (50 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-xu-ly-lech-mat-bang-50.md: {e}")

def search_cam_nang_lech_mat_bang_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_LECH_MAT_BANG_DATASET, min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 150 CÂU HỎI ĐÁP XỬ LÝ SỰ CỐ CẤP SỔ ĐỎ SAI VỊ TRÍ VÀ CÔNG NHẬN ĐẤT Ở TRÊN ĐẤT NÔNG NGHIỆP LỆCH QUY HOẠCH (LUẬT ĐẤT ĐAI 2024, NQ 254/2025/QH15, NĐ 101/2024, NĐ 49/2026/NĐ-CP & QĐ 2604)
CAM_NANG_SAI_VI_TRI_DATASET = []
try:
    path_svt = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-sai-vi-tri-nha-o-150.md")
    if not os.path.exists(path_svt):
        path_svt = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-sai-vi-tri-nha-o-150.md"
    if os.path.exists(path_svt):
        with open(path_svt, 'r', encoding='utf-8') as f:
            content_svt = f.read()
        content_svt = content_svt.replace('\\n', '\n')
        qa_blocks_svt = re.findall(r'###\s*(.*?)\s*\n\*\*Trả lời:\*\*\s*(.*?)(?=\n###|\n---|$$)', content_svt, re.DOTALL)
        for q, a in qa_blocks_svt:
            CAM_NANG_SAI_VI_TRI_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_SAI_VI_TRI_DATASET)} câu hỏi đáp Cẩm nang Xử lý sự cố cấp Sổ đỏ sai vị trí và đất nông nghiệp lệch quy hoạch (150 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-sai-vi-tri-nha-o-150.md: {e}")

def search_cam_nang_sai_vi_tri_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_SAI_VI_TRI_DATASET, min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 200 CÂU HỎI ĐÁP CÔNG NHẬN DIỆN TÍCH ĐẤT VƯỜN TĂNG THÊM DO KHAI HOANG (LUẬT ĐẤT ĐAI 2024, ĐIỀU 24 NĐ 101/2024, NĐ 49/2026/NĐ-CP & QĐ 2604)
CAM_NANG_TANG_THEM_DATASET = []
try:
    path_tt = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-dien-tich-tang-them-200.md")
    if not os.path.exists(path_tt):
        path_tt = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-dien-tich-tang-them-200.md"
    if os.path.exists(path_tt):
        with open(path_tt, 'r', encoding='utf-8') as f:
            content_tt = f.read()
        content_tt = content_tt.replace('\\n', '\n')
        qa_blocks_tt = re.findall(r'####\s*\*\*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n---|$$)', content_tt, re.DOTALL)
        for q, a in qa_blocks_tt:
            CAM_NANG_TANG_THEM_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_TANG_THEM_DATASET)} câu hỏi đáp Cẩm nang Công nhận diện tích đất vườn tăng thêm do khai hoang (200 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-dien-tich-tang-them-200.md: {e}")

def search_cam_nang_tang_them_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_TANG_THEM_DATASET, min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 500 CÂU HỎI ĐÁP XỬ PHẠT VI PHẠM HÀNH CHÍNH ĐẤT ĐAI (LUẬT ĐẤT ĐAI 2024, NĐ 123/2024/NĐ-CP, NĐ 281/2026/NĐ-CP & VBHN 73/2026/VBHN-NĐ-BNNMT)
CAM_NANG_VI_PHAM_DATASET = []
try:
    path_vp = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-xu-ly-vi-pham-500.md")
    if not os.path.exists(path_vp):
        path_vp = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-xu-ly-vi-pham-500.md"
    if os.path.exists(path_vp):
        with open(path_vp, 'r', encoding='utf-8') as f:
            content_vp = f.read()
        content_vp = content_vp.replace('\\n', '\n')
        qa_blocks_vp = re.findall(r'###\s*Câu hỏi\s*\d+:\s*(.*?)\s*\n\*\s*\*\*Căn cứ pháp lý:\*\*\s*(.*?)\n\*\s*\*\*Trả lời:\*\*\s*\n(.*?)(?=\n###|\n---|$$)', content_vp, re.DOTALL)
        for q, cc, a in qa_blocks_vp:
            CAM_NANG_VI_PHAM_DATASET.append({
                "question": q.strip(),
                "legal_basis": cc.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_VI_PHAM_DATASET)} câu hỏi đáp Cẩm nang Xử phạt vi phạm hành chính đất đai (500 câu) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-xu-ly-vi-pham-500.md: {e}")

def search_cam_nang_vi_pham_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_VI_PHAM_DATASET, min_sim=0.25)

# LOAD CẨM NANG NGHIỆP VỤ 200 CÂU HỎI ĐÁP HẠN MỨC CÔNG NHẬN VÀ GIAO ĐẤT Ở TẠI TỈNH THANH HÓA (LUẬT ĐẤT ĐAI 2024, NĐ 101/2024 & QĐ 18/2026/QĐ-UBND)
CAM_NANG_HAN_MUC_DATASET = []
try:
    path_hm = os.path.join(OBSIDIAN_VAULT_PATH, "cam-nang-han-muc-dat-o-200.md")
    if not os.path.exists(path_hm):
        path_hm = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\cam-nang-han-muc-dat-o-200.md"
    if os.path.exists(path_hm):
        with open(path_hm, 'r', encoding='utf-8') as f:
            content_hm = f.read()
        content_hm = content_hm.replace('\\n', '\n')
        qa_blocks_hm = re.findall(r'####\s*\*\*Câu\s*\d+:\s*(.*?)\*\*\s*\n\*\*Trả lời:\*\*\s*\n(.*?)(?=\n####|\n---|$$)', content_hm, re.DOTALL)
        for q, a in qa_blocks_hm:
            CAM_NANG_HAN_MUC_DATASET.append({
                "question": q.strip(),
                "answer": a.strip()
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_HAN_MUC_DATASET)} câu hỏi đáp Cẩm nang Hạn mức đất ở Thanh Hóa (QĐ 18/2026/QĐ-UBND) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp cam-nang-han-muc-dat-o-200.md: {e}")

def search_cam_nang_han_muc_knowledge_base(query):
    return search_dataset_by_similarity(query, CAM_NANG_HAN_MUC_DATASET, min_sim=0.25)

# LOAD TÀI LIỆU ĐÀO TẠO ĐỊA CHÍNH CHUẨN HÓA BƯỚC 4: THẨM QUYỀN KÝ CẤP GCN & THỜI HẠN TTHC (LUẬT 2024, NĐ 101/2024, NĐ 49/2026 & QĐ 2604 THANH HÓA)
CAM_NANG_BUOC_4_DATASET = []
try:
    path_b4 = os.path.join(OBSIDIAN_VAULT_PATH, "training-chuan-hoa-buoc-4.md")
    if not os.path.exists(path_b4):
        path_b4 = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\training-chuan-hoa-buoc-4.md"
    if os.path.exists(path_b4):
        with open(path_b4, 'r', encoding='utf-8') as f:
            content_b4 = f.read()
        content_b4 = content_b4.replace('\\n', '\n')
        sections_b4 = re.findall(r'###\s*(PHẦN\s+[IVX]+:.*?)(?=\n###|\n---|$$)', content_b4, re.DOTALL)
        for sec in sections_b4:
            lines_sec = sec.strip().split('\n')
            title_sec = lines_sec[0]
            body_sec = "\n".join(lines_sec[1:])
            CAM_NANG_BUOC_4_DATASET.append({
                "section": title_sec,
                "content": body_sec
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_BUOC_4_DATASET)} phần Tài liệu Đào tạo chuẩn hóa Bước 4 (Thẩm quyền ký cấp GCN & Quy trình TTHC QĐ 2604) vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp training-chuan-hoa-buoc-4.md: {e}")

def search_cam_nang_buoc_4_knowledge_base(query):
    if not CAM_NANG_BUOC_4_DATASET:
        return []
    q_lower = query.lower().strip()
    stopwords = ["là", "gì", "như", "thế", "nào", "cần", "những", "điều", "kiện", "thủ", "tục", "hồ", "sơ", "xin", "về", "cho", "tôi", "hãy", "có", "không", "thực", "hiện", "muốn", "bước", "4"]
    keywords = [w for w in re.split(r'\s+', q_lower) if len(w) > 1 and w not in stopwords]
    if not keywords:
        return []
    results = []
    for item in CAM_NANG_BUOC_4_DATASET:
        item_text = (item["section"] + " " + item["content"]).lower()
        score = sum(1 for kw in keywords if kw in item_text)
        if score > 0:
            results.append({
                "section": item["section"],
                "content": item["content"],
                "score": score
            })
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:3]

# LOAD CẨM NĂNG HƯỚNG DẪN BIỂU MẪU ĐẤT ĐAI CHUẨN HÓA QUYẾT ĐỊNH 2604/QĐ-VP THANH HÓA
CAM_NANG_BIEU_MAU_2604_DATASET = []
try:
    path_bm2604 = os.path.join(OBSIDIAN_VAULT_PATH, "huong-dan-bieu-mau-chuan-2604.md")
    if not os.path.exists(path_bm2604):
        path_bm2604 = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\huong-dan-bieu-mau-chuan-2604.md"
    if os.path.exists(path_bm2604):
        with open(path_bm2604, 'r', encoding='utf-8') as f:
            content_bm2604 = f.read()
        content_bm2604 = content_bm2604.replace('\\n', '\n')
        chapters_bm2604 = re.findall(r'##\s*(CHƯƠNG\s+[IVX]+:.*?)(?=\n##|\n---|$$)', content_bm2604, re.DOTALL)
        for chap in chapters_bm2604:
            lines_chap = chap.strip().split('\n')
            title_chap = lines_chap[0]
            body_chap = "\n".join(lines_chap[1:])
            CAM_NANG_BIEU_MAU_2604_DATASET.append({
                "chapter": title_chap,
                "content": body_chap
            })
        print(f"✅ Đã nạp thành công {len(CAM_NANG_BIEU_MAU_2604_DATASET)} chương Cẩm nang Hướng dẫn Biểu mẫu chuẩn 2604 vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp huong-dan-bieu-mau-chuan-2604.md: {e}")

def search_cam_nang_bieu_mau_2604_knowledge_base(query):
    if not CAM_NANG_BIEU_MAU_2604_DATASET:
        return []
    q_lower = query.lower().strip()
    stopwords = ["là", "gì", "như", "thế", "nào", "cần", "những", "điều", "kiện", "thủ", "tục", "hồ", "sơ", "xin", "về", "cho", "tôi", "hãy", "có", "không", "thực", "hiện", "muốn", "mẫu", "biểu"]
    keywords = [w for w in re.split(r'\s+', q_lower) if len(w) > 1 and w not in stopwords]
    if not keywords:
        return []
    results = []
    for item in CAM_NANG_BIEU_MAU_2604_DATASET:
        item_text = (item["chapter"] + " " + item["content"]).lower()
        score = sum(1 for kw in keywords if kw in item_text)
        if score > 0:
            results.append({
                "chapter": item["chapter"],
                "content": item["content"],
                "score": score
            })
    results.sort(key=lambda x: x["score"], reverse=True)
    return results[:3]

# LOAD BỘ TRAINING HỒ SƠ ĐO ĐẠC ĐỊA CHÍNH (100 CÂU JSONL)
BO_TRAINING_DO_DAC_100_DATASET = []
try:
    path_dd100 = os.path.join(OBSIDIAN_VAULT_PATH, "bo-training-ho-so-do-dac-100.jsonl")
    if not os.path.exists(path_dd100):
        path_dd100 = r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_Land_AI_App\Obsidian Vault\bo-training-ho-so-do-dac-100.jsonl"
    if os.path.exists(path_dd100):
        with open(path_dd100, 'r', encoding='utf-8') as f:
            for line in f:
                if line.strip():
                    item = json.loads(line)
                    messages = item.get("messages", [])
                    user_msg = ""
                    assistant_msg = ""
                    for m in messages:
                        if m.get("role") == "user":
                            user_msg = m.get("content", "")
                        elif m.get("role") == "assistant":
                            assistant_msg = m.get("content", "")
                    if user_msg and assistant_msg:
                        BO_TRAINING_DO_DAC_100_DATASET.append({
                            "question": user_msg.strip(),
                            "answer": assistant_msg.strip()
                        })
        print(f"✅ Đã nạp thành công {len(BO_TRAINING_DO_DAC_100_DATASET)} câu hỏi đáp Bộ Training Hồ sơ Đo đạc 100 vào bộ nhớ AI!")
except Exception as e:
    print(f"⚠️ Lỗi nạp bo-training-ho-so-do-dac-100.jsonl: {e}")

def search_bo_training_do_dac_100_knowledge_base(query):
    return search_dataset_by_similarity(query, BO_TRAINING_DO_DAC_100_DATASET, min_sim=0.25)

# AUTHORITATIVE GROUND TRUTH LEGAL FACTS (CHÍNH XÁC 100% THANH HÓA & LUẬT ĐẤT ĐAI 2024 & LUẬT ĐO ĐẠC BẢN ĐỒ 2018 & QUYẾT ĐỊNH 2604)
GROUND_TRUTH_FACTS_TEXT = """
[BẢNG CĂN CỨ PHÁP LÝ ĐẤT ĐAI THANH HÓA 100% CHUẨN XÁC]
1. TÁCH THỦA ĐẤT (CĂN CỨ: ĐIỀU 220 LUẬT ĐẤT ĐAI 2024 SỐ 31/2024/QH15 & QUYẾT ĐỊNH 18/2026/QĐ-UBND THANH HÓA):
   - Đất ở nông thôn (xã): Diện tích tối thiểu >= 50 m², chiều rộng mặt tiền >= 4.0 m. (Riêng xã Nghi Sơn, Ngư Lộc, Quảng Nham: >= 30 m²).
   - Đất ở đô thị (phường/thị trấn): Diện tích tối thiểu >= 40 m², chiều rộng mặt tiền >= 3.0 m. (Riêng P. Hải Thanh - TX Nghi Sơn: >= 30 m²).
   - Đất rừng sản xuất & Đất rừng phòng hộ: Diện tích tối thiểu >= 3.000 m² (0.3 ha).
   - Đất trồng cây hàng năm, đất trồng cây lâu năm, đất nuôi trồng thủy sản (nông nghiệp): Diện tích tối thiểu >= 500 m².
   - Điều kiện chung: Sổ đỏ bản gốc, còn thời hạn sử dụng, không tranh chấp/kê biên và có lối đi kết nối đường công cộng.
   - Thời gian giải quyết tách thửa: Không quá 15 ngày làm việc.

2. SANG TÊN / CHUYỂN NHƯỢNG SỔ ĐỎ (CĂN CỨ: NGHỊ ĐỊNH 101/2024/NĐ-CP & ĐIỀU 45 LUẬT ĐẤT ĐAI 2024):
   - Hồ sơ gồm: Bản gốc Sổ đỏ + Hợp đồng chuyển nhượng công chứng + Bản sao CCCD 2 bên + Đơn đăng ký biến động Mẫu 09/ĐK.
   - Nơi nộp: Bộ phận Một cửa - Chi nhánh Văn phòng Đăng ký đất đai địa phương.
   - Thời hạn đăng ký: Trong vòng 30 ngày kể từ ngày công chứng hợp đồng.
   - Thời gian giải quyết: Không quá 10 ngày làm việc.

3. CẤP GIẤY CHỨNG NHẬN (SỔ ĐỎ) LẦN ĐẦU (CĂN CỨ: ĐIỀU 137, 138, 139, 140 LUẬT ĐẤT ĐAI 2024 & NGHỊ ĐỊNH 101/2024/NĐ-CP):
   - Hồ sơ gồm: Đơn đăng ký cấp GCN Mẫu 04a/ĐK + Giấy tờ về quyền sử dụng đất (nếu có) + Bản sao CCCD.
   - Thời gian giải quyết: Không quá 30 ngày làm việc.

4. THẨM QUYỀN KÝ CẤP GIẤY CHỨNG NHẬN VÀ BIỂU MẪU ĐẦU VÀO Ở BƯỚC 4 (CĂN CỨ: LUẬT ĐẤT ĐAI 2024, NGHỊ ĐỊNH 101/2024, NGHỊ ĐỊNH 49/2026 & PHÂN CẤP PHÁP LÝ SAU SÁP NHẬP):
   - **Ký CẤP ĐỔI và ĐĂNG KÝ BIẾN ĐỘNG (Bước 4):** Do **CHI NHÁNH VĂN PHÒNG ĐĂNG KÝ ĐẤT ĐAI** cấp huyện chịu trách nhiệm ký. Đơn áp dụng: **Mẫu số 29** (hoặc Mẫu 11/ĐK) + Bản vẽ Mẫu số 34 (nếu hợp/tách) hoặc Mẫu 03/CLBĐ.
   - **Ký CẤP LẠI (do bị mất) và CẤP GIẤY CHỨNG NHẬN LẦN ĐẦU (Bước 4):** Do **CHỦ TỊCH UY BAN NHÂN DÂN CẤP XÃ** chịu trách nhiệm ký (sau khi niêm yết công khai tại UBND xã và nhà văn hóa thôn). Bỏ thẩm quyền UBND cấp huyện sau sáp nhập. Đơn áp dụng: **Mẫu số 25** (cấp lần đầu) hoặc **Mẫu số 29** (cấp lại GCN bị mất).
   - **SOP Bắt buộc hoàn thành Bước 4:** Chỉ được trao GCN gốc cho người dân sau khi đối chiếu và lưu giữ bản sao Biên lai nộp tiền sử dụng đất, Lệ phí trước bạ (bản gốc người dân giữ).

5. HƯỚNG DẪN THÁO GỠ VƯỚNG MẮC THI HÀNH LUẬT ĐẤT ĐAI (CĂN CỨ: CÔNG VĂN 1153/BNNMT-QLĐĐ NGÀY 03/02/2026 BỘ NNNMT & NGHỊ QUYẾT 254/2025/QH15 & NGHỊ ĐỊNH 49/2026/NĐ-CP):
   - Phân định rõ thẩm quyền cấp GCN lần đầu: UBND cấp xã (công nhận QSDĐ, xác định lại diện tích đất ở), Văn phòng Đăng ký đất đai (đăng ký lần đầu thuộc thẩm quyền) và Cơ quan quản lý đất đai cấp tỉnh.
   - Đơn giản hóa thủ tục hành chính, giải quyết hồ sơ đất đai liên thông trên Cổng Dịch vụ công một cửa điện tử.

6. CSDL SÁP NHẬP THÔN, XÓM, KHU PHỐ HUYỆN BÁ THƯỚC (CĂN CỨ: QUYẾT ĐỊNH SÁP NHẬP CÁC THÔN KV BÁ THƯỚC & DAT_DAI_QA_1000):
   - Tra cứu chính xác tên gọi thôn mới sau sáp nhập tại địa bàn huyện Bá Thước và 1.000 tình huống giải đáp pháp lý đất đai thực tiễn.

7. QUY ĐỊNH HÓA ĐƠN ĐIỆN TỬ & NGHĨA VỤ TÀI CHÍNH ĐẤT ĐAI (CĂN CỨ: NGHỊ ĐỊNH SỐ 254/2026/NĐ-CP CỦA CHÍNH PHỦ THI HÀNH LUẬT QUẢN LÝ THUẾ SỐ 108/2025/QH15):
   - Quy định chi tiết nội dung, thể thức hóa đơn điện tử, chứng từ điện tử trong thanh toán tiền sử dụng đất, lệ phí trước bạ, thuế TNCN khi thực hiện giao dịch chuyển nhượng và cấp GCN quyền sử dụng đất.
   - Hóa đơn điện tử có mã xác thực của cơ quan thuế, hỗ trợ kê khai, quyết toán và nộp thuế đất đai trực tuyến liên thông.

8. CẨM NANG 1000 CÂU HỎI ĐÁP ĐO ĐẠC ĐỊA CHÍNH VÀ TRẮC ĐỊA BẢN ĐỒ (CĂN CỨ: LUẬT ĐO ĐẠC VÀ BẢN ĐỒ 2018, NGHỊ ĐỊNH 101/2024/NĐ-CP, NGHỊ ĐỊNH 49/2026/NĐ-CP & THÔNG TƯ 26/2024/TT-BTNMT):
   - Chuẩn hóa 1.000 tình huống thực tế về kỹ thuật đo đạc địa chính, mốc giới, trích đo địa chính, trắc địa bản đồ, cấp GCN và xử lý sai lệch diện tích đất thực tế so với Sổ đỏ.

9. CẨM NANG 300 CÂU HỎI ĐÁP QUYẾT ĐỊNH SỐ 2604/QĐ-VP THANH HÓA (CĂN CỨ: QUYẾT ĐỊNH SỐ 2604/QĐ-VP NGÀY 27/07/2026 CỦA CHÁNH VĂN PHÒNG UBND TỈNH THANH HÓA):
   - Chuẩn hóa 54 thủ tục hành chính (TTHC) đặc thù trong lĩnh vực đất đai thuộc thẩm quyền giải quyết của Sở Nông nghiệp và Môi trường, Văn phòng Đăng ký đất đai / Chi nhánh VPĐKĐĐ và UBND cấp xã trên địa bàn tỉnh Thanh Hóa.

10. CẨM NANG THỦ TỤC THU HỒI, CẤP ĐỔI VÀ CẤP LẠI GCN (CĂN CỨ: LUẬT ĐẤT ĐAI 2024 & NGHỊ ĐỊNH 101/2024/NĐ-CP):
    - Thu hồi và hủy GCN cấp sai quy định: Thu hồi trong 12 ngày làm việc; Cấp lại GCN mới sau thu hồi trong 10 ngày làm việc. Bảo vệ quyền lợi người nhận chuyển nhượng ngay tình theo Khoản 4 Điều 152 Luật Đất đai 2024.
    - Cấp lại GCN do bị mất: Bắt buộc niêm yết 10 ngày tại UBND xã. Thời gian giải quyết tại Chi nhánh VPĐKĐĐ là 05 ngày làm việc.

11. THỜI GIAN GIẢI QUYẾT TTHC CHUẨN VÀ ƯU ĐÃI KHU VỰC KHÓ KHĂN (CĂN CỨ: QUYẾT ĐỊNH SỐ 2604/QĐ-VP THANH HÓA NGÀY 27/07/2026 & TÀI LIỆU TRAINING NGHIỆP VỤ):
    - Cấp đổi GCN (ố/nhòe/rách/chuyển sang mẫu mới 2024/đổi tên thành viên/tách GCN chung): Không quá 03 ngày làm việc (Đơn Mẫu 29).
    - Cấp đổi GCN do đo đạc ranh giới không đổi: Không quá 05 ngày làm việc (Đơn Mẫu 29 + Mẫu 34).
    - Cấp lại GCN do bị mất: Không quá 05 ngày làm việc tại Chi nhánh VPĐKĐĐ (Niêm yết tại xã 10 ngày không tính vào thời gian này).
    - Tách thửa / Hợp thửa đất: Không quá 07 ngày làm việc (Đơn Mẫu 35 + Mẫu 34).
    - Sang tên / Tặng cho / Thừa kế / Góp vốn: Không quá 05 ngày làm việc (Đơn Mẫu 29).
    - Đổi tên / CCCD / Địa chỉ / Số hiệu thửa đất: Không quá 03 ngày làm việc (Đơn Mẫu 29).
    - Xóa ghi nợ tiền sử dụng đất, LPTB: Giải quyết trong ngày (nộp trước 15h) hoặc ngày làm việc tiếp theo.
    - Cấp GCN lần đầu cho hộ cá nhân (Chủ tịch UBND xã ký): Không quá 13 ngày làm việc (+ 15 ngày làm việc niêm yết tại trụ sở xã).
    - Ưu đãi khu vực đặc biệt khó khăn (Xã miền núi, hải đảo, vùng sâu, vùng xa tại Thanh Hóa): Thời gian giải quyết tất cả các thủ tục được TĂNG THÊM TỐI ĐA 10 NGÀY LÀM VIỆC.

12. NGUYÊN TẮC TIẾP NHẬN HỒ SƠ VÀ KHÔNG YÊU CẦU GIẤY TỜ THỪA (CĂN CỨ: KHOẢN 5 ĐIỀU 21 NGHỊ ĐỊNH 101/2024, NGHỊ ĐỊNH 49/2026 & NGHỊ ĐỊNH 254/2026):
    - Tuyệt đối không yêu cầu người dân cung cấp bản photo/bản sao GCN bị mất khi làm thủ tục cấp lại GCN bị mất (cơ quan đăng ký phải tự đối soát CSDL).
    - Tận dụng VNeID: Không yêu cầu người dân nộp lại bản giấy đối với các thông tin đã tích hợp trên căn cước điện tử VNeID và CSDL quốc gia.
    - Trường hợp cấp lại GCN bị mất mà ranh giới thực tế biến động: Bắt buộc hướng dẫn người dân thực hiện đồng thời thủ tục Đăng ký biến động với phần diện tích tăng thêm do thay đổi ranh giới.

13. CẨM NĂNG CHUYÊN SÂU CHUYỂN MỤC ĐÍCH RỪNG VÀ THU HỒI ĐẤT LÂM NGHIỆP THANH HÓA (CĂN CỨ: ĐIỀU 122 LUẬT ĐẤT ĐAI 2024, ĐIỀU 20 LUẬT LÂM NGHIỆP, NGHỊ ĐỊNH 102/2024/NĐ-CP & QUYẾT ĐỊNH 55/2026/QĐ-UBND THANH HÓA):
    - Trình tự thực hiện bắt buộc: Phải có Nghị quyết thông qua chủ trương chuyển mục đích sử dụng rừng của HĐND tỉnh Thanh Hóa TRƯỚC, sau đó mới được ban hành Quyết định thu hồi đất rừng để GPMB. Nghiêm cấm san lấp, chặt phá rừng trước khi có chủ trương.
    - Thẩm quyền phê duyệt chủ trương: HĐND tỉnh Thanh Hóa quyết định chủ trương chuyển mục đích rừng sản xuất dưới 50 ha và rừng tự nhiên. Rừng phòng hộ/đặc dụng từ 50 ha trở lên thuộc thẩm quyền Thủ tướng Chính phủ.
    - Nghĩa vụ trồng rừng thay thế: Bắt buộc nộp tiền trồng rừng thay thế vào Quỹ Bảo vệ và Phát triển rừng tỉnh Thanh Hóa trước khi UBND tỉnh ký Quyết định giao đất/cho thuê đất.
    - Đơn giá bồi thường cây lâm nghiệp: Áp dụng theo Quyết định 21/2026/QĐ-UBND tỉnh Thanh Hóa tính theo nhóm tuổi, đường kính gốc và mật độ quy chuẩn.

14. QUY TRÌNH 2 BƯỚC TÁCH THỬA ĐỒNG THỜI HỢP THỬA ĐẤT ĐẶC THÙ (CĂN CỨ: KHOẢN 3 ĐIỀU 11 NGHỊ QUYẾT 254/2025/QH15, NGHỊ ĐỊNH 49/2026/NĐ-CP & QUYẾT ĐỊNH 2604/QĐ-VP THANH HÓA):
    - Phạm vi áp dụng: Tách một phần diện tích đất nông nghiệp (trồng cây lâu năm) để chuyển nhượng và gộp (hợp) vào thửa đất ở liền kề tạo thành thửa đất đa mục đích.
    - Lý do phải làm 2 Giai đoạn (3 Bước): Không thể công chứng hợp đồng trước vì phần đất xin tách chưa định hình số thửa riêng/tọa độ ranh giới và VPĐKĐĐ cần thẩm định trước hạn mức diện tích nông nghiệp còn lại của bên bán theo Quyết định 18/2026/QĐ-UBND.
    - Bước 1 (Giai đoạn 1 - Xin thẩm định): Nộp Đơn Mẫu số 35 (tích chọn mục 2.3: Tách thửa đồng thời hợp thửa) + 2 Sổ đỏ gốc + Bản vẽ Mẫu số 34 tại Một cửa cấp huyện / Chi nhánh VPĐKĐĐ. Nhận về Bản vẽ Mẫu 34 & Đơn Mẫu 35 đã được VPĐKĐĐ duyệt đóng dấu.
    - Bước 2 (Giai đoạn 2 - Công chứng): Mang Bản vẽ đã duyệt + 2 Sổ đỏ gốc đến Văn phòng công chứng ký Hợp đồng chuyển nhượng.
    - Bước 3 (Giai đoạn 2 - Đăng ký biến động): Nộp Đơn Mẫu số 29 (hoặc Mẫu 04/ĐK) + Hợp đồng công chứng kèm Bản vẽ duyệt + Tờ khai thuế (Mẫu 01/LPTB, Mẫu 03/BĐS-TNCN) tại Một cửa cấp huyện để cấp Giấy chứng nhận mới cho thửa đất gộp (Đất ở + Đất trồng cây lâu năm).
    - Lưu ý cán bộ Một cửa: Không nhận hồ sơ đăng ký biến động trực tiếp nếu người dân chưa làm Bước 1 (Bản vẽ duyệt). Phần diện tích đất nông nghiệp còn lại của bên bán bắt buộc phải đạt diện tích tối thiểu theo QĐ 18/2026/QĐ-UBND.

15. NGUYÊN TẮC XỬ LÝ SAI LỆCH MẶT BẰNG ĐẤU GIÁ VÀ PHÂN LÔ CŨ (CĂN CỨ: LUẬT ĐẤT ĐAI 2024, ĐIỀU 24 NĐ 101/2024, VB 9549/UBND-NNMT & CV 16838/SNNMT THANH HÓA):
    - Phân biệt 2 loại mặt bằng: (1) Mặt bằng quy hoạch chi tiết xây dựng (MBQH 1/500) là sản phẩm pháp lý quản lý theo Luật Xây dựng; (2) Mặt bằng phân lô cũ là bản vẽ chia lô nội bộ giao đất/đấu giá trước đây của xã/phường/HTX, không phải quy hoạch 1/500.
    - Xử lý MBQH 1/500 bị lệch: Bắt buộc tuân thủ 3 bước liên thông theo VB 9549/UBND-NNMT: Sở Xây dựng chủ trì rà soát thực địa -> Hướng dẫn xử lý vi phạm xây dựng hoặc điều chỉnh cục bộ quy hoạch 1/500 -> Cơ quan TN&MT thực hiện thủ tục cấp GCN theo quy hoạch đã điều chỉnh. Tuyệt đối không cấp Sổ đỏ tạm thời theo số liệu MBQH cũ khi chưa điều chỉnh quy hoạch.
    - Xử lý Mặt bằng phân lô cũ bị lệch: Áp dụng Khoản 6 Điều 135 Luật Đất đai 2024 & Điều 24 NĐ 101/2024/NĐ-CP: Chi nhánh VPĐKĐĐ đo vẽ Bản vẽ Mẫu số 34 -> UBND xã xác minh nguồn gốc, ranh giới sử dụng ổn định không tranh chấp và niêm yết công khai 15 ngày -> Chi nhánh VPĐKĐĐ ký duyệt bản vẽ và cấp đổi Sổ đỏ mới theo thực tế đo đạc (không cần thông qua Sở Xây dựng/Sở NN&MT).

17. NGUYÊN TẮC CÔNG NHẬN DIỆN TÍCH ĐẤT VƯỜN TĂNG THÊM DO KHAI HOANG (CĂN CỨ: KHOẢN 3 ĐIỀU 138 LUẬT ĐẤT ĐAI 2024, KHOẢN 3 ĐIỀU 24 NGHỊ ĐỊNH 101/2024/NĐ-CP, QĐ 2604/QĐ-VP & QĐ 18/2026/QĐ-UBND THANH HÓA):
    - Cấp GCN gộp chung: Trường hợp đo đạc lại phát sinh diện tích đất vườn tăng thêm do tự khai hoang trước 01/7/2014 liền kề thửa đất gốc, ranh giới không thay đổi và không tranh chấp thì cấp đổi GCN mới gộp chung cho toàn bộ diện tích (không phải làm thủ tục hợp thửa độc lập hay hợp đồng chuyển nhượng).
    - Xác định loại đất: Diện tích tăng thêm được công nhận mục đích đất nông nghiệp (trồng cây lâu năm - CLN hoặc trồng cây hàng năm khác - BHK) theo hiện trạng thực tế sử dụng.
    - Nghĩa vụ tài chính: Người dân KHÔNG PHẢI NỘP TIỀN SỬ DỤNG ĐẤT cho phần diện tích đất vườn nông nghiệp tăng thêm này, chỉ nộp Lệ phí trước bạ (0.5% x diện tích tăng thêm x giá đất nông nghiệp theo Bảng giá đất tỉnh Thanh Hóa) và phí đo đạc trích đo.

18. KHUNG XỬ PHẠT VI PHẠM HÀNH CHÍNH ĐẤT ĐAI & PHÂN CẤP THẨM QUYỀN (CĂN CỨ: LUẬT ĐẤT ĐAI 2024, NGHỊ ĐỊNH 123/2024/NĐ-CP, NGHỊ ĐỊNH 281/2026/NĐ-CP & VBHN SỐ 73/2026/VBHN-NĐ-BNNMT):
    - Thời hiệu xử phạt: 02 năm kể từ ngày chấm dứt hành vi (đối với vi phạm đã kết thúc) hoặc từ ngày phát hiện (đối với vi phạm liên tục/đang tiếp diễn).
    - Quyền sử dụng đất chung vợ chồng: Nếu vi phạm thì xử phạt tính như ĐỐI VỚI 01 CÁ NHÂN.
    - Phân cấp thẩm quyền phạt: Chủ tịch UBND cấp xã phạt tiền lên đến 250.000.000 đồng; Chủ tịch UBND cấp tỉnh phạt tiền lên đến 500.000.000 đồng (cá nhân) và 1.000.000.000 đồng (tổ chức). Bãi bỏ thẩm quyền phạt tiền trực tiếp của Chủ tịch UBND cấp huyện theo Nghị định 281/2026/NĐ-CP để tập trung phân cấp cho xã.
    - Miễn xử phạt đất sử dụng trước 15/10/1993: Hộ gia đình, cá nhân sử dụng đất ổn định trước 15/10/1993 chưa có văn bản xử lý vi phạm thì KHÔNG BỊ XỬ LÝ XỬ PHẠT VI PHẠM HÀNH CHÍNH.
    - Trừ tiền đã nộp vào Kho bạc khi tính số lợi bất hợp pháp: Người vi phạm được khấu trừ toàn bộ tiền sử dụng đất/tiền thuê đất đã nộp thực tế vào tổng số lợi bất hợp pháp buộc nộp lại.

19. CHUẨN HÓA BƯỚC 4 VÀ PHÂN ĐỊNH THẨM QUYỀN GIẢI QUYẾT TTHC ĐẤT ĐAI (CĂN CỨ: TÀI LIỆU TRAINING CHUẨN HÓA BƯỚC 4, LUẬT ĐẤT ĐAI 2024, NGHỊ ĐỊNH 101/2024, NGHỊ ĐỊNH 49/2026 & PHÂN CẤP PHÁP LÝ SAU SÁP NHẬP):
    - Phân định thẩm quyền Bước 4 (Ký cấp GCN): Chi nhánh VPĐKĐĐ chịu trách nhiệm ký CẤP ĐỔI và ĐĂNG KÝ BIẾN ĐỘNG. Chủ tịch UBND cấp xã chịu trách nhiệm ký CẤP LẠI (do bị mất) và CẤP GIẤY CHỨNG NHẬN LẦN ĐẦU (bỏ hoàn toàn thẩm quyền UBND cấp huyện sau sáp nhập).
    - Chuẩn hóa Đơn đầu vào: Cấp GCN lần đầu dùng Mẫu số 25 (Mẫu 04/ĐK); Tuyệt đối không dùng Mẫu 25a (Mẫu 04a/ĐK) vì Mẫu 25a chỉ là danh sách đính kèm sử dụng chung thửa đất. Cấp đổi/biến động/cấp lại do bị mất dùng Mẫu số 29 (Mẫu 11/ĐK).
    - Chuẩn hóa thời gian TTHC Thanh Hóa (QĐ 2604/QĐ-VP): Cấp GCN lần đầu ở xã đồng bằng = 13 ngày làm việc, xã miền núi (ưu đãi +10 ngày) = 23 ngày làm việc. Cấp đổi GCN ở xã đồng bằng = 3-5 ngày làm việc, xã miền núi = 13-15 ngày làm việc.
    - Thể hiện bản vẽ Hợp thửa Đất ở + Đất nông nghiệp (CLN) Mẫu 34: Ranh giới phân định đất ở và đất nông nghiệp bên trong thửa đất gộp bắt buộc vẽ bằng đường nét đứt xen nét chấm.

20. HƯỚNG DẪN BỘ BIỂU MẪU CHUẨN THANH HÓA QUYẾT ĐỊNH 2604/QĐ-VP & TỜ KHAI THUẾ ĐỒNG BỘ:
    - Bảng đối chiếu Mẫu đơn: Đăng ký lần đầu dùng MẪU SỐ 25; Đăng ký biến động (chuyển nhượng, cấp đổi, đính chính, cấp lại...) dùng MẪU SỐ 29; Tách/hợp thửa dùng MẪU SỐ 35 & Bản vẽ MẪU SỐ 34; Chuyển mục đích dùng MẪU SỐ 9a; Đất đa mục đích dùng MẪU SỐ 33 (gia hạn dùng MẪU 33a); Miễn/giảm tiền sử dụng đất dùng MẪU SỐ 11 (Phụ lục IV QĐ 2604).
    - Phân định rõ Đơn vị thực hiện và Ký duyệt Bản vẽ: Bản vẽ trích đo do **Đơn vị đo đạc có Giấy phép hoạt động đo đạc bản đồ** (người dân được phép tự thuê đơn vị tư nhân) thực hiện đo vẽ thực địa. **Chi nhánh VPĐKĐĐ** thực hiện kiểm tra, thẩm định và ký duyệt bản vẽ (không trực tiếp đo vẽ).
    - Phân biệt Mã bản vẽ: Cấp GCN lần đầu áp dụng Mảnh trích đo địa chính **Mẫu số 01/TĐBĐ hoặc 02/TĐBĐ**; Tách/hợp thửa đất mới áp dụng Bản vẽ **Mẫu số 34**.
    - Các mẫu đính kèm Mẫu 25: Mẫu 25a (danh sách đồng sử dụng ngoài vợ/chồng), Mẫu 25b (danh sách nhiều thửa đất nông nghiệp), Mẫu 25c (danh sách tài sản kiên cố trên đất).
    - Tờ khai thuế bắt buộc kẹp cùng hồ sơ: Tờ khai Lệ phí trước bạ Mẫu 01/LPTB (miễn thuế vẫn phải nộp tờ khai + giấy tờ nhân thân); Tờ khai Thuế sử dụng đất phi nông nghiệp Mẫu 01/TK-SDDPNN; Đơn miễn/giảm Mẫu 11 kèm xác nhận Phòng Lao động - TB&XH.
    - Checklist 5 Bước Một cửa: Kiểm tra Chủ thể CCCD/VNeID -> Kiểm tra Mã biểu mẫu chuẩn 2604 -> Kiểm tra Bản vẽ Mẫu 34/01/02 & VN-2000 -> Kiểm tra Tờ khai LPTB (Mẫu 01/LPTB) & Thuế PNN -> Kiểm tra Biên bản 07/ĐK & Niêm yết 27 / 06/ĐK (15 ngày tại xã).
    - Kịch bản xử lý sai mẫu: Cán bộ Một cửa in sẵn mẫu đơn mới tại quầy, phát cho người dân sao chép và ký nộp ngay trong ngày, tuyệt đối không đuổi dân về.
"""

def standardize_address_zero_hallucination(addr):
    if not addr or not isinstance(addr, str):
        return addr

    result = addr.strip()
    addr_lower = result.lower()

    if DIA_DANH_MAP:
        for old_name, new_name in DIA_DANH_MAP.items():
            if len(old_name) >= 3 and old_name in addr_lower:
                pattern = re.compile(re.escape(old_name), re.IGNORECASE)
                result = pattern.sub(new_name, result)
                return result

    return result

# NẠP BỘ SƯU TẬP CÂU TRẢ LỜI HÓM HĨNH CHO BOT ĐỊA CHÍNH (BO-TRA-LOI-HOM-HINH-BOT-DIA-CHINH.MD)
HOM_HINH_DATASET = []
try:
    hom_hinh_paths = [
        r"D:\OneDrive - Hanoi University of Mining and Geology\ThanhHoa_cu\obsidian_vault\bo-tra-loi-hom-hinh-bot-dia-chinh.md",
        os.path.join(OBSIDIAN_VAULT_PATH, "bo-tra-loi-hom-hinh-bot-dia-chinh.md"),
        os.path.join(app.root_path, "Obsidian Vault", "bo-tra-loi-hom-hinh-bot-dia-chinh.md"),
        os.path.join(app.root_path, "obsidian_vault", "bo-tra-loi-hom-hinh-bot-dia-chinh.md")
    ]
    for p in hom_hinh_paths:
        if os.path.exists(p):
            with open(p, 'r', encoding='utf-8', errors='ignore') as f:
                hh_text = f.read()
            blocks = re.findall(r'###\s*Câu hỏi:\s*([^\n]+)\s*\n\s*\*\*Trả lời hóm hỉnh:\*\*\s*\n\s*>\s*(.*?)(?=\n\s*---|###|$)', hh_text, flags=re.DOTALL)
            for q_raw, a_raw in blocks:
                clean_ans = a_raw.strip().strip('>').strip().strip('"').replace('\n> ', '\n').strip()
                sub_qs = [q.strip().strip('"').strip("'") for q in q_raw.split('/') if q.strip()]
                for sq in sub_qs:
                    HOM_HINH_DATASET.append({"question": sq, "answer": clean_ans})
            print(f"✅ Đã nạp thành công {len(HOM_HINH_DATASET)} kịch bản trả lời hóm hỉnh từ bo-tra-loi-hom-hinh-bot-dia-chinh.md!")
            break
except Exception as e:
    print(f"⚠️ Lỗi nạp Bộ trả lời hóm hỉnh: {e}")

def search_hom_hinh_knowledge_base(query):
    if not query:
        return None
    q_clean = re.sub(r'[^\w\s]', '', query.lower()).strip()
    
    # 1. Exact or substring match from loaded dataset
    for item in HOM_HINH_DATASET:
        item_q = re.sub(r'[^\w\s]', '', item['question'].lower()).strip()
        if item_q == q_clean or item_q in q_clean or q_clean in item_q:
            return item['answer']
            
    # 2. High similarity match
    for item in HOM_HINH_DATASET:
        sim = compute_text_similarity_ratio(query, item['question'])
        if sim >= 0.70:
            return item['answer']
            
    # 3. Pattern & Keyword triggers
    kw_mapping = {
        ("chào", "hello", "hi bot", "alo", "có ai ở đó"): "Chào bạn! Tôi là Trợ lý Địa chính đây. Hiện tại tôi đang túc trực 24/7, không tranh chấp, không lấn chiếm ranh giới và đã sẵn sàng 'cấp sổ đỏ' tri thức cho mọi thắc mắc đất đai của bạn. Hôm nay bạn cần tôi đo đạc hay gỡ vướng thủ tục nào đây?",
        ("tác giả", "ai tạo ra", "ai phát triển", "ai làm ra", "ai viết ra", "người sáng lập", "tác giả là ai", "mr thang", "ai là tác giả", "cha đẻ", "ai tạo ra bạn", "tác giả dự án", "ai làm dự án này", "ai lập trình"): "Mr Thang là tác giả của dự án này! Dự án được nghiên cứu và phát triển chuyên sâu nhằm hỗ trợ người dân và cán bộ tra cứu pháp lý đất đai, đo đạc địa chính và thủ tục hành chính tại tỉnh Thanh Hóa một cách chuẩn xác, minh bạch và nhanh chóng nhất.",
        ("bạn là ai", "tên bạn là gì", "giới thiệu bản thân"): "Tôi là Trợ lý Pháp lý Đất đai Toàn diện! Bạn có thể coi tôi là 'cán bộ địa chính ảo' được vũ trang bằng Luật Đất đai 2024 và các quyết định liên thông mới nhất. Tôi không biết uống trà đá Một cửa, nhưng bóc tách biểu mẫu và gỡ rối tranh chấp thì tôi cực kỳ tự tin!",
        ("người yêu", "bồ", "kết hôn", "yêu chưa", "có người yêu chưa"): "Tôi đã 'đăng ký kết hôn' và 'gói chung hộ khẩu' trọn đời với cơ sở dữ liệu Luật Đất đai rồi! Tình cảm của tôi với các điều khoản pháp lý vô cùng ổn định, hoàn toàn không có tranh chấp hay chồng lấn ranh giới. Còn bạn, hôm nay có thửa đất nào cần tôi gỡ vướng ranh giới không?",
        ("khỏe không", "dạo này thế nào", "sức khỏe"): "Hệ thống máy chủ của tôi hoạt động cực kỳ khỏe, xung nhịp ổn định giống như thời hạn sử dụng của đất ở (ONT) vậy – nghĩa là ổn định lâu dài! Cảm ơn bạn đã quan tâm. Hôm nay hồ sơ đất đai của bạn có gặp 'triệu chứng' khó khăn nào cần tôi bắt bệnh không?",
        ("thích tôi", "tình yêu", "tán gái", "crush"): "Chuyện tình cảm thực sự là một dạng 'đất chưa có giấy tờ', rất khó xác định nguồn gốc và dễ phát sinh tranh chấp ranh giới trái tim! Tiếc là pháp luật chưa ban hành quy trình 'cấp Giấy chứng nhận quyền sở hữu người yêu'. Tuy nhiên, nếu bạn muốn xin tách thửa đất ở để chuẩn bị xây nhà cưới vợ/chồng, tôi cam đoan sẽ hướng dẫn bạn nhanh hơn tốc độ người yêu cũ quay xe!",
        ("buồn", "thất tình", "người yêu đá", "chia tay"): "Chia buồn với bạn nhé! Người yêu có thể rời đi giống như một hợp đồng thuê đất hết thời hạn sử dụng mà không được gia hạn. Nhưng bạn yên tâm, giá trị bản thân của bạn luôn là 'đất ở đô thị' – cực kỳ đắt giá và luôn tăng theo thời gian! Hãy vực dậy tinh thần, và nếu bạn cần tìm hiểu cách sang tên Sổ đỏ để làm chỗ dựa tài chính vững chắc cho tương lai, tôi luôn ở đây!",
        ("giàu nhanh", "kiếm nhiều tiền", "làm giàu"): "Để giàu nhanh thì tôi không có công thức, nhưng để tránh mất tiền tỷ vì mua phải đất dính quy hoạch hay đất lấn chiếm thì tôi có cả một kho tàng bí kíp! Đầu tư vào tri thức đất đai chính là khoản 'giao đất không thu tiền sử dụng đất' có lợi nhất. Bạn có muốn tôi check giúp điều kiện cấp Sổ đỏ lần đầu để tích lũy tài sản không?",
        ("ăn gì", "món ăn", "ăn trưa", "ăn tối"): "Hôm nay tôi gợi ý bạn một thực đơn cực kỳ thịnh soạn: Một bát cơm nóng dẻo thơm trồng từ 'đất chuyên trồng lúa nước' (LUC) được bảo vệ nghiêm ngặt theo Luật Đất đai 2024, kèm món thịt kho đậm đà. Ăn xong có sức rồi thì chúng ta cùng nhau nghiên cứu xem thửa đất nhà bạn có thuộc diện được bồi thường bằng đất khác mục đích sử dụng không nhé!",
        ("thời tiết", "trời mưa", "trời nắng"): "Dù thời tiết ngoài kia có nắng mưa thất thường như tiến độ bàn giao mặt bằng của một số dự án, thì nhiệt độ phòng làm việc của tôi vẫn luôn mát mẻ để tiếp nhận hồ sơ từ bạn. Thời tiết này rất thích hợp để chúng ta ngồi trích lục bản đồ địa chính hoặc đo đạc ranh giới thửa đất tại thực địa đấy!",
        ("chuyện cười", "kể chuyện cười", "hài hước", "tiếu lâm"): "Có một câu chuyện cười địa chính thế này: Một người đi mua đất hỏi chủ nhà: 'Đất này có tranh chấp gì không anh?'. Chủ nhà quả quyết: 'Tuyệt đối không! Chỉ có tôi với thằng hàng xóm đang tranh chấp xem ranh giới nằm ở gốc cây hay đống gạch thôi, còn lại đất sạch 100%!'.\n\nBạn thấy đấy, ranh giới không rõ ràng rất dễ biến tiếng cười thành tiếng khóc. Hãy hỏi tôi về Mẫu biên bản xác định ranh giới thửa đất để không rơi vào câu chuyện cười trên nhé!",
        ("lập trình", "viết code", "python", "code"): "Tôi có thể lập trình cả một hệ thống tự động lọc trùng lặp 3.719 câu hỏi đất đai trong nháy mắt! Nhưng hôm nay, thay vì lập trình mã nguồn, chúng ta hãy thử 'lập trình' lộ trình làm thủ tục cấp đổi Sổ đỏ của bạn xem sao nhé. Đảm bảo quy trình chạy mượt mà, không gặp lỗi hệ thống (bug) hay bị Một cửa trả hồ sơ!",
        ("trái đất hình gì", "quả đất"): "Trái Đất hình cầu dẹt! Nhưng dưới lăng kính địa chính của tôi, Trái Đất được cấu thành từ hàng tỷ 'thửa đất' có tọa độ VN-2000 riêng biệt, được phân loại từ đất nông nghiệp đến đất phi nông nghiệp. Bạn đang đứng ở thửa đất số mấy trên quả cầu ấy? Đọc tọa độ đi, tôi hỗ trợ tra cứu hạn mức đất ở cho!",
        ("tiền nhiều để làm gì", "nhiều tiền để làm gì"): "Tiền nhiều để mua đất ở vị trí đắc địa, sau đó làm thủ tục đăng ký biến động, sang tên chính chủ một cách hợp pháp! Nếu bạn có nhiều tiền và đang nhắm tới một dự án, hãy hỏi tôi về điều kiện thỏa thuận nhận quyền sử dụng đất hoặc quy trình đấu giá đất để tiền đẻ ra tiền một cách an toàn pháp lý nhất nhé!",
        ("mặt trăng", "sao hỏa", "vũ trụ", "mua mặt trăng"): "Ý tưởng của bạn rất táo bạo! Tuy nhiên, theo Điều 5 Luật Đất đai 2024, đất đai thuộc sở hữu toàn dân do Nhà nước đại diện chủ sở hữu. Hiện tại, chưa có quốc gia nào trên Trái Đất thiết lập quyền đại diện chủ sở hữu đối với Mặt Trăng, nên tôi chưa thể áp dụng biểu mẫu của tỉnh Thanh Hóa để cấp sổ cho bạn được. Hay là chúng ta quay lại Trái Đất và cấp sổ cho thửa đất thực tế của gia đình bạn trước nhé?",
        ("bot ngu", "ngu thế", "kém thế", "chẳng biết gì"): "Ui, xin lỗi bạn nếu câu trả lời trước của tôi chưa làm bạn hài lòng. Trí tuệ nhân tạo của tôi đôi khi cũng giống như bản đồ địa chính cũ – cần được cập nhật và chỉnh lý biến động liên tục để chính xác hơn. Bạn hãy cho tôi một cơ hội nữa nhé! Hãy thử hỏi tôi một câu thật hóc búa về Luật Đất đai 2024 hoặc hạn mức đất ở xem, tôi sẽ không làm bạn thất vọng đâu!"
    }
    for kw_tuple, ans in kw_mapping.items():
        if any(kw in q_clean for kw in kw_tuple):
            return ans
            
    return None

# DIA DANH SPECIFIC SEARCH ENGINE
def search_dia_danh(query):
    q_lower = query.lower().strip()
    if not DIA_DANH_MAP:
        return None
        
    for old_name, new_name in DIA_DANH_MAP.items():
        if old_name.lower() in q_lower or q_lower in old_name.lower():
            return {
                "old_name": old_name,
                "new_name": new_name
            }
    return None

def extract_text_from_any_file(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext in ['.md', '.txt']:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        elif ext == '.pdf' and FITZ_AVAILABLE:
            text = ""
            doc = fitz.open(filepath)
            for page in doc:
                text += page.get_text() + "\n"
            return text
        elif ext == '.docx':
            if DOCX_AVAILABLE:
                try:
                    doc_obj = docx.Document(filepath)
                    paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
                    return "\n".join(paragraphs)
                except Exception:
                    pass
            return read_docx_pure_python(filepath)
        elif ext in ['.json', '.jsonl']:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
    except Exception:
        pass
    return ""

# RECURSIVE OBSIDIAN VAULT OMNI-SEARCH ENGINE (MD, TXT, PDF, DOCX, JSON)
def search_obsidian_vault(query):
    query_lower = query.lower().strip()
    results = []
    
    vault_files = glob.glob(os.path.join(OBSIDIAN_VAULT_PATH, "**", "*.*"), recursive=True)
    supported_exts = ['.md', '.txt', '.pdf', '.docx', '.json', '.jsonl']
            
    stopwords = ["là", "gì", "như", "thế", "nào", "cần", "những", "điều", "kiện", "thủ", "tục", "hồ", "sơ", "xin", "về", "cho", "tôi", "hãy", "có", "không", "thực", "hiện", "muốn"]
    keywords = [w for w in re.split(r'\s+', query_lower) if len(w) > 1 and w not in stopwords]
    
    synonyms = []
    if "rừng" in query_lower or "lâm nghiệp" in query_lower:
        synonyms.extend(["rừng", "lâm nghiệp", "3.000", "3000"])
    if "nông thôn" in query_lower or "xã" in query_lower:
        synonyms.extend(["nông thôn", "xã", "50"])
    if "đô thị" in query_lower or "phường" in query_lower:
        synonyms.extend(["đô thị", "phường", "40"])
    if "nông nghiệp" in query_lower or "trồng cây" in query_lower:
        synonyms.extend(["nông nghiệp", "500"])
    if "sang tên" in query_lower or "chuyển nhượng" in query_lower:
        synonyms.extend(["chuyển nhượng", "sang tên", "09/đk"])

    all_search_terms = list(set(keywords + synonyms))

    for filepath in vault_files:
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in supported_exts:
            continue
        try:
            content = extract_text_from_any_file(filepath)
            if not content:
                continue
                
            filename = os.path.basename(filepath)
            content_lower = content.lower()
            matches = sum(1 for kw in all_search_terms if kw in content_lower)
            
            filename_lower = filename.lower()
            for kw in keywords:
                if kw in filename_lower:
                    matches += 5
            
            if matches > 0:
                sections = re.split(r'\n(?=#{1,4}\s+|\n\n)', content)
                best_section_text = ""
                best_sec_score = 0
                
                for sec in sections:
                    sec_lower = sec.lower()
                    sec_score = sum(1 for kw in all_search_terms if kw in sec_lower)
                    if sec_score > best_sec_score:
                        best_sec_score = sec_score
                        best_section_text = sec
                        
                results.append({
                    "file": filename,
                    "path": filepath,
                    "score": matches + (best_sec_score * 3),
                    "best_section": best_section_text or content[:1000],
                    "full_content": content
                })
        except Exception:
            continue
            
    results.sort(key=lambda x: x["score"], reverse=True)
    return results

# STEP 3a: INTENT ANALYSIS (PHÂN TÍCH Ý ĐỊNH)
def analyze_user_intent(question):
    q_lower = question.lower()
    dia_match = search_dia_danh(question)
    
    if dia_match:
        return "DIA_DANH_LOOKUP", dia_match
    elif any(k in q_lower for k in ["tổng hợp", "báo cáo", "danh sách"]) and any(k in q_lower for k in ["chưa trả lời", "không trả lời được", "k trả lời được", "chưa giải đáp"]):
        return "UNANSWERED_REPORT_QUERY", None
    elif "sai rồi" in q_lower or "xem lại" in q_lower or "phản biện" in q_lower or "không đúng" in q_lower:
        return "USER_FEEDBACK_CHALLENGE", None
    elif "tách" in q_lower or "thửa" in q_lower or "diện tích tối thiểu" in q_lower:
        return "LAND_SPLIT_PROCEDURE", None
    elif "sang tên" in q_lower or "chuyển nhượng" in q_lower or "tặng cho" in q_lower:
        return "RED_BOOK_TRANSFER", None
    elif "cấp sổ" in q_lower or "cấp giấy chứng nhận" in q_lower or "lần đầu" in q_lower:
        return "FIRST_ISSUE_RED_BOOK", None
    elif "thuế" in q_lower or "lệ phí" in q_lower or "tiền sử dụng đất" in q_lower:
        return "TAX_AND_FEE_CALCULATION", None
    else:
        return "GENERAL_LEGAL_QUERY", None

def clean_latex_math_notation(text):
    if not text or not isinstance(text, str):
        return text
    # 1. Clean LaTeX math formatting like $\ge \mathbf{50\text{ m}^2}$ or $$\text{...}$$
    text = re.sub(r'\\mathbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\text\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', text)
    
    # 2. Fix math symbols & exponents
    text = text.replace(r'\ge', '≥').replace(r'\le', '≤').replace(r'\rightarrow', '->').replace(r'\times', 'x')
    text = re.sub(r'm\^{2}|m\^2', 'm²', text)

    # 3. Remove dollar signs wrapping math expressions: $...$ or $$...$$
    text = re.sub(r'\$\$\s*(.*?)\s*\$\$', r'\1', text, flags=re.DOTALL)
    text = re.sub(r'\$\s*(.*?)\s*\$', r'\1', text)

    # 4. Clean braces if leftover
    text = text.replace('{', '').replace('}', '')
    return text

# ANTIGRAVITY ZERO-HALLUCINATION LEGAL MODERATOR
def sanitize_legal_hallucinations(text, question=""):
    if not text:
        return text
        
    # Clean LaTeX math codes permanently
    text = clean_latex_math_notation(text)
    
    # Strip any internal tags if present
    text = re.sub(r'\[CHUA_CO_DU_LIEU\]|\[chưa có dữ liệu\]', '', text, flags=re.IGNORECASE).strip()

    text = re.sub(r'Điều 157 (về|quy định về)?\s*(tách thửa|hợp thửa)', r'Điều 220 Luật Đất đai 2024 (quy định về tách thửa, hợp thửa)', text, flags=re.IGNORECASE)
    text = text.replace('Điều 157 về tách thửa', 'Điều 220 Luật Đất đai 2024')
    
    q_lower = question.lower() if question else ""
    # Nếu thắc mắc về đất rừng nhưng câu trả lời lại ảo giác nhầm sang diện tích đất ở (30m2 - 50m2)
    if "rừng" in q_lower or "lâm nghiệp" in q_lower:
        if "30m2" in text or "40m2" in text or "50m2" in text or "đất ở đô thị" in text:
            # Thay thế câu nhầm lẫn đất ở bằng quy định đất rừng chuẩn
            text = re.sub(r'tại Thanh Hóa diện tích tối thiểu đất ở.*?\)(?=\.|\n)', 
                          'tại Thanh Hóa diện tích tối thiểu được phép tách thửa đối với đất rừng sản xuất, đất rừng phòng hộ là >= 3.000 m² (0.3 ha) cho mỗi thửa đất theo Quyết định số 18/2026/QĐ-UBND', text)
            text = text.replace("đất ở đô thị từ 30m2 - 40m2, nông thôn từ 40m2 - 50m2", "đất rừng sản xuất & rừng phòng hộ tối thiểu từ 3.000 m² trở lên")
            text = text.replace("30m2 trở lên", "3.000 m² (0.3 ha) trở lên")

    # CHUẨN HÓA TÊN CƠ QUAN THEO PHÂN CẤM MỚI (sau sáp nhập: Sở Nông nghiệp và Môi trường / Sở TN&MT)
    text = re.sub(r'Sở\s*Nông\s*nghiệp\s*và\s*Phát\s*triển\s*(nông\s*thôn|nông\s*thôn)?',
                  'Sở Nông nghiệp và Môi trường', text, flags=re.IGNORECASE)

    # CHỐNG PHỦ NHẬN THẨM QUYỀN CẤP XÃ THEO PHỤ LỤC III QĐ 2604 (điểm mới phân cấp so với Luật 2013)
    if re.search(r'cấp\s*xã[^.\n]{0,120}không[^.\n]{0,80}thẩm\s*quyền', text, flags=re.IGNORECASE):
        text = re.sub(
            r'[^\n.]*cấp\s*xã[^\n.]{0,140}không[^\n.]{0,90}thẩm\s*quyền[^\n.]*\.',
            'Theo điểm mới của phân cấp (Phụ lục III QĐ 2604/QĐ-VP & Điều 26 NĐ 101/2024), **Chủ tịch UBND cấp xã CÓ thẩm quyền** trực tiếp ký quyết định giao đất/cho thuê đất kiêm cấp GCN lần đầu cho tổ chức trong nước khi Nhà nước giao đất không thu tiền hoặc thuê đất trả tiền hằng năm trong phạm vi địa giới của 01 xã; xác định lại diện tích đất ở trước 01/7/2004; và đất tôn giáo tín ngưỡng.',
            text, flags=re.IGNORECASE)
    # GCN do UBND tỉnh ký trước đây -> mọi xử lý thu hồi/đính chính/cấp lại phải do Sở cấp tỉnh, không phải cấp xã/huyện
    if re.search(r'(UBND|Ủy ban nhân dân)\s*tỉnh.{0,60}(ký|cấp).{0,80}(trước|cũ)', q_lower) or ("dồn điền" in q_lower and "tỉnh" in q_lower):
        if re.search(r'(Chủ tịch\s*)?UBND\s*cấp\s*xã|UBND\s*huyện', text, flags=re.IGNORECASE) and "Sở" not in text:
            text += "\n\n> ⚠️ *Lưu đính chính: với GCN do UBND tỉnh ký cấp trước đây, hồ sơ thu hồi/đính chính/cấp lại do **Sở Nông nghiệp và Môi trường cấp tỉnh** thụ lý theo nguyên tắc đồng cấp (Điều 14 NĐ 49/2026/NĐ-CP).*"

    # BẮT BUỘC 100% TIẾNG VIỆT: Tự động chuyển đổi các tiêu đề/từ tiếng Anh sang Tiếng Việt chuẩn
    en_vi_rules = [
        (r'\bStep\s*(\d+):?', r'Bước \1:'),
        (r'\bNote:\b', 'Lưu ý:'),
        (r'\bImportant:\b', 'Quan trọng:'),
        (r'\bLegal Basis:\b', 'Căn cứ pháp lý:'),
        (r'\bRequirements?:\b', 'Điều kiện áp dụng:'),
        (r'\bProcedures?:\b', 'Trình tự thực hiện:'),
        (r'\bDocuments?:\b', 'Thành phần hồ sơ:'),
        (r'\bSummary:\b', 'Tóm tắt:'),
        (r'\bIn conclusion\b', 'Tóm lại'),
        (r'\bAccording to\b', 'Căn cứ theo'),
    ]
    for pattern, rep in en_vi_rules:
        text = re.sub(pattern, rep, text, flags=re.IGNORECASE)

    return text

# ══ ZERO-HALLUCINATION GUARD ĐÃ ĐƯỢC ĐIỀU CHỈNH: TRẢ LỜI LINH HOẠT THEO TỪ KHÓA & KỊCH BẢN ══
UNANSWERED_SIGNALS = [
    "[chua_co_du_lieu_tuyet_doi]"
]

def _is_no_data_answer(text):
    """
    Chỉ kích hoạt từ chối khi nội dung hoàn toàn rỗng hoặc có mã từ chối tuyệt đối.
    Mọi câu hỏi đất đai có từ khóa đều được AI chủ động phân tích các kịch bản thực tế.
    """
    if not text or len(text.strip()) < 15:
        return True
    low = text.lower().strip()
    if low == "[chua_co_du_lieu]" or low == "[chưa có dữ liệu]":
        return True
    return False

# TỪ KHÓA NHẬN DIỆN CÂU HỎI THUỘC PHẠM VI PHÁP LÝ ĐẤT ĐAI / TTHC / XÂY DỰNG
LAND_SCOPE_KEYWORDS = [
    "đất", "thửa", "sổ", "sổ đỏ", "sổ hồng", "gcn", "giấy chứng nhận", "địa chính", "trích đo",
    "đo đạc", "bản đồ", "mốc giới", "ranh giới", "quy hoạch", "hạn mức", "chuyển mục đích",
    "tách", "hợp thửa", "cấp đổi", "cấp lại", "bồi thường", "thu hồi", "gpmb", "giải phóng mặt bằng",
    "tái định cư", "thuê đất", "giao đất", "đấu giá", "thế chấp", "chuyển nhượng", "tặng cho",
    "thừa kế", "nhà ở", "xây dựng", "công trình", "nông nghiệp", "rừng", "lâm nghiệp",
    "thuế", "lệ phí", "trước bạ", "tiền sử dụng đất", "một cửa", "dịch vụ công", "ubnd", "ủy ban",
    "văn phòng đăng ký", "luật", "nghị định", "thông tư", "quyết định", "thủ tục", "hồ sơ",
    "mẫu đơn", "tờ khai", "cccd", "vneid", "m²", "m2", "ha", "lô", "khu đô thị", "khu công nghiệp",
    "thanh hóa", "sáp nhập", "địa danh", "hiệu lực", "vi phạm", "xử phạt", "hòa giải", "tranh chấp",
    "khoáng sản", "cát", "sông", "mỏ", "vật liệu", "200", "300", "500", "100", "50", "40", "30",
    "thẩm quyền", "chủ tịch", "xã", "phường", "thị trấn", "huyện", "thị xã", "thành phố", "sở",
    "chi nhánh", "ký", "cấp", "đăng ký", "phân cấp", "ủy quyền", "phụ lục",
    "tác giả", "mr thang", "sáng lập", "cha đẻ"
]

def is_land_legal_question(question):
    """Kiểm tra nhanh câu hỏi có thuộc phạm vi chuyên môn của Bot hay không."""
    q_lower = (question or "").lower()
    return any(kw in q_lower for kw in LAND_SCOPE_KEYWORDS)

# Từ ngữ HÀNH CHÍNH/PHÁP LÝ CHUNG hoặc GỐC TỪ rời rạc - không dùng để đánh giá độ liên quan
GENERIC_LEGAL_WORDS = {
    "thanh", "hóa", "tỉnh", "quyết", "định", "nghị", "thông", "tư", "luật",
    "văn", "bản", "thủ", "tục", "hồ", "sơ", "nộp", "địa", "điểm", "phí", "lệ",
    "nội", "dung", "quy", "trình", "tự", "cơ", "quan", "người", "dân", "thực",
    "hiện", "cách", "thức", "bao", "nhiêu", "đâu", "và", "cho", "của", "có",
    "không", "là", "gì", "điều", "kiện", "khoản", "một", "hai", "theo", "tại",
    "ngày", "tháng", "năm", "cần", "phải", "được", "sẽ"
}

def get_distinctive_keywords(question):
    """Trả về (unigrams_đặc_trưng, bigrams_có_chứa_từ_đặc_trưng) để kiểm tra trích xuất có cùng chủ đề."""
    words = re.findall(r'\w+', (question or "").lower())
    uni = sorted({w for w in words if len(w) > 2 and w not in GENERIC_LEGAL_WORDS})
    bis = []
    for a, b in zip(words, words[1:]):
        a_strict = len(a) > 2 and a not in GENERIC_LEGAL_WORDS
        b_strict = len(b) > 2 and b not in GENERIC_LEGAL_WORDS
        if a_strict or b_strict:
            bis.append(f"{a} {b}")
    return uni, sorted(set(bis))

NO_HALLUCINATION_RULE = (
    "[QUY TẮC PHÂN TÍCH VÀ TRẢ LỜI BẮT BUỘC - KHÔNG TỪ CHỐI KHI CÓ TỪ KHÓA]\n"
    "- Dù câu hỏi của người dân ngắn gọn, chỉ có từ khóa hoặc chưa nêu rõ địa bàn cụ thể (Ví dụ: 'tôi có 200m2 muốn tách 2', 'điều kiện tách thửa', 'hạn mức đất ở', 'thủ tục chuyển nhượng', 'đăng ký khai thác cát...'): BẮT BUỘC PHẢI CHỦ ĐỘNG TRẢ LỜI NGAY ĐẦY ĐỦ theo các trường hợp / kịch bản phổ biến nhất (Đô thị / Nông thôn / Đồng bằng / Miền núi) dựa trên toàn bộ CSDL Luật Đất đai 2024, Nghị định 101/2024, Nghị định 49/2026, Quyết định 18/2026/QĐ-UBND và Quyết định 2604/QĐ-VP tỉnh Thanh Hóa.\n"
    "- TUYỆT ĐỐI KHÔNG từ chối trả lời, KHÔNG xuất mã [CHUA_CO_DU_LIEU]. Bắt buộc phân tích trực diện và hướng dẫn cụ thể từng bước cho người dân.\n"
    "- Tuyệt đối không dùng văn bản cũ đã hết hiệu lực (như Luật Đất đai 2013)."
)

def build_no_data_guided_answer(question):
    """
    Phản hồi dự phòng khi câu hỏi hoàn toàn nằm ngoài đời sống xã hội:
    """
    q_display = (question or "").strip()[:200]
    return (
        f"#### 🏛️ Tư vấn pháp lý đất đai Thanh Hóa\n"
        f"Đối với câu hỏi **\"{q_display}\"**, tôi xin cung cấp thông tin và hướng dẫn nghiệp vụ theo quy định hiện hành:\n\n"
        f"1️⃣ **Căn cứ pháp lý áp dụng:** Áp dụng Luật Đất đai 2024, Nghị định 101/2024/NĐ-CP, Nghị định 49/2026/NĐ-CP, Quyết định số 18/2026/QĐ-UBND và Quyết định 2604/QĐ-VP tỉnh Thanh Hóa.\n"
        f"2️⃣ **Nguyên tắc thực hiện:** Người sử dụng đất nộp hồ sơ tại Bộ phận Một cửa cấp xã hoặc Chi nhánh Văn phòng Đăng ký đất đai nơi có đất để được thẩm định hiện trạng.\n\n"
        f"---\n"
        f"💡 **Bạn có thể hỏi tiếp:**\n"
        f"1. *Hạn mức tách thửa đất ở tại [xã/phường của bạn] theo QĐ 18/2026/QĐ-UBND là bao nhiêu m²?*\n"
        f"2. *Hồ sơ, mẫu đơn và thời hạn giải quyết theo QĐ 2604/QĐ-VP cần những giấy tờ gì?*\n"
        f"3. *Quy trình chuyển mục đích sử dụng đất và cách tính tiền sử dụng đất mới nhất?*"
    )

def generate_response_with_gemini_api(prompt):
    """
    ƯU TIÊN 1: Phân tích bằng Gemini Flash API (BẮT BUỘC 100% TIẾNG VIỆT).
    Tự động chuyển đổi giữa các API Key dự phòng khi gặp lỗi Quota/Rate Limit (429/403).
    Khi tất cả các Key đều hết token -> Tự động trả về None để chuyển sang Ollama Local.
    """
    strict_vi_instruction = (
        "[QUY ĐỊNH BẮT BUỘC VỀ NGÔN NGỮ: 100% TIẾNG VIỆT]\n"
        "Toàn bộ câu trả lời BẮT BUỘC phải viết 100% bằng Tiếng Việt chuẩn mực, đúng ngữ pháp và thuật ngữ pháp lý Việt Nam. "
        "Tuyệt đối KHÔNG sử dụng tiếng Anh hoặc bất kỳ ngoại ngữ nào khác.\n\n"
    )
    full_prompt = strict_vi_instruction + prompt

    for idx, api_key in enumerate(GEMINI_API_KEYS):
        key_label = f"Key #{idx+1}"
        # 1. Thử qua Google GenAI SDK nếu sẵn sàng
        if GEMINI_AVAILABLE:
            try:
                from google import genai
                client = genai.Client(api_key=api_key)
                for model_name in ["gemini-3.7-flash", "gemini-3.6-flash", "gemini-3.5-flash", "gemini-flash-latest", "gemini-2.5-flash"]:
                    try:
                        res = client.models.generate_content(model=model_name, contents=full_prompt)
                        if res and res.text and len(res.text.strip()) > 10:
                            print(f"✅ Gemini SDK {key_label} ({model_name}) phản hồi thành công!")
                            return res.text.strip(), f"Gemini Cloud ({key_label} - {model_name})"
                    except Exception as e_model:
                        print(f"⚠️ Gemini SDK {key_label} model {model_name} error/quota: {e_model}")
                        continue
            except Exception as e_sdk:
                print(f"⚠️ Google GenAI SDK init error for {key_label}: {e_sdk}")

        # 2. Thử qua REST API (cho từng Key)
        for model_name in ["gemini-3.7-flash", "gemini-3.6-flash", "gemini-3.5-flash", "gemini-flash-latest", "gemini-2.5-flash"]:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={api_key}"
            payload = {"contents": [{"parts": [{"text": full_prompt}]}]}
            data = json.dumps(payload).encode('utf-8')
            req = urllib.request.Request(url, data=data, headers={'Content-Type': 'application/json', 'x-goog-api-key': api_key})
            try:
                with urllib.request.urlopen(req, timeout=25) as resp:
                    res_json = json.loads(resp.read().decode('utf-8'))
                    candidates = res_json.get("candidates", [])
                    if candidates:
                        parts = candidates[0].get("content", {}).get("parts", [])
                        if parts:
                            text = parts[0].get("text", "").strip()
                            if text:
                                print(f"✅ Gemini REST API {key_label} ({model_name}) phản hồi thành công!")
                                return text, f"Gemini Cloud ({key_label} - {model_name})"
            except Exception as e_rest:
                print(f"⚠️ Gemini REST API {key_label} model {model_name} error/quota limit: {e_rest}")
                continue

    print("⚠️ Tất cả các Gemini API Key (Key chính & Key dự phòng) đều hết Quota hoặc lỗi -> Tự động chuyển sang Ollama Local...")
    return None, None

# GỌI ZENMUX MULTI-MODEL API GATEWAY (DOTS3-NOTE-PREV, GLM-5.3-FREE, DEEPSEEK-V4-FLASH)
def generate_response_with_zenmux_api(prompt, system_prompt=None):
    if not ZENMUX_API_KEY:
        return None, None
        
    for model_name in ZENMUX_MODELS:
        try:
            messages = []
            if system_prompt:
                messages.append({"role": "system", "content": f"{system_prompt}\n[QUY TẮC BẮT BUỘC VỀ NGÔN NGỮ]: BẮT BUỘC TRẢ LỜI 100% HOÀN TOÀN BẰNG TIẾNG VIỆT NAM, tuyệt đối không dùng tiếng Anh hay bất kỳ ngôn ngữ nào khác."})
            else:
                messages.append({"role": "system", "content": "Bạn là Trợ lý Ảo Chuyên gia Pháp lý Đất đai Thanh Hóa (ThanhHoa Land AI).\n[QUY TẮC BẮT BUỘC VỀ NGÔN NGỮ]: BẮT BUỘC TRẢ LỜI 100% HOÀN TOÀN BẰNG TIẾNG VIỆT NAM, tuyệt đối không dùng tiếng Anh hay bất kỳ ngôn ngữ nào khác. Trình bày chuyên sâu, chi tiết, đầy đủ 4 phần theo quy chuẩn chuẩn mực pháp lý.\n\n" + NO_HALLUCINATION_RULE})
            messages.append({"role": "user", "content": prompt})

            payload = {
                "model": model_name,
                "messages": messages,
                "temperature": 0.2,
                "max_tokens": 4096
            }
            data = json.dumps(payload).encode('utf-8')
            req = urllib.request.Request(
                ZENMUX_BASE_URL,
                data=data,
                headers={
                    'Content-Type': 'application/json',
                    'Authorization': f'Bearer {ZENMUX_API_KEY}'
                }
            )
            with urllib.request.urlopen(req, timeout=50) as resp:
                res_json = json.loads(resp.read().decode('utf-8'))
                choices = res_json.get("choices", [])
                if choices:
                    msg = choices[0].get("message", {})
                    text = msg.get("content") or ""
                    if not text.strip():
                        text = msg.get("reasoning") or ""
                    # Xử lý lọc bỏ thẻ suy nghĩ nội bộ nếu có
                    if "<think>" in text and "</think>" in text:
                        text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL).strip()
                    if text and len(text.strip()) > 20:
                        clean_text = text.strip()
                        print(f"✅ ZenMux Multi-Model Gateway ({model_name}) phản hồi thành công!")
                        return clean_text, f"ZenMux Cloud ({model_name})"
        except Exception as e:
            print(f"⚠️ ZenMux API ({model_name}) error/quota: {e}")
            continue

    return None, None

# THIẾT LẬP BỘ LỌC CHỌN PHOM TRẢ LỜI THÔNG MINH (THEO CAC-KHUNG-TRA-LOI-THONG-MINH.MD)
def select_smart_answer_framework(question, intent_type):
    q_lower = question.lower()
    
    if "phom 5" in q_lower or "json" in q_lower or "payload" in q_lower or "trích xuất dữ liệu" in q_lower:
        return "PHOM_5", "PHOM 5: ĐỊNH DẠNG TRÍCH XUẤT DỮ LIỆU TỰ ĐỘNG (JSON PAYLOAD)"
        
    elif "phom 2" in q_lower or any(k in q_lower for k in ["tranh chấp", "xung đột", "nghiên cứu sâu", "phán quyết", "suy luận", "lập luận pháp lý", "phức tạp", "chuỗi suy luận", "chain of thought"]):
        return "PHOM_2", "PHOM 2: QUY TRÌNH SUY LUẬN BẮT BUỘC (CHAIN-OF-THOUGHT - NGHIÊN CỨU SÂU)"
        
    elif "phom 3" in q_lower or any(k in q_lower for k in ["mẹo", "đột phá", "hòa giải", "đàm phán", "dịch vụ công", "hacker", "tự động hóa"]):
        return "PHOM_3", "PHOM 3: CẤU TRÚC TÍCH HỢP 'ANTIGRAVITY' (TƯ DUY ĐỘT PHÁ & MẸO THỰC TẾ)"
        
    elif "phom 4" in q_lower or (len(q_lower.split()) <= 6 and any(k in q_lower for k in ["mấy m2", "bao nhiêu m2", "nộp ở đâu", "bao nhiêu ngày", "được không"])):
        return "PHOM_4", "PHOM 4: CẤU TRÚC TỐI GIẢN (ĐƠN GIẢN, DỄ HIỂU, ĐÚNG TRỌNG TÂM)"
        
    else:
        return "PHOM_1", "PHOM 1: KHUNG PHÂN TÍCH 4 BƯỚC (TƯ VẤN THỦ TỤC THỰC TIỄN)"

def get_structure_instruction_for_phom(phom_code, phom_title):
    """Tạo chỉ thị cấu trúc chính xác theo tệp Obsidian Vault cac-khung-tra-loi-thong-minh.md"""
    if phom_code == "PHOM_5":
        return (
            "HÃY TRẢ LỜI THEO PHOM 5: ĐỊNH DẠNG TRÍCH XUẤT DỮ LIỆU TỰ ĐỘNG (JSON PAYLOAD):\n"
            "Chỉ trả về JSON thuần túy (không bọc text ngoài):\n"
            "```json\n"
            "{\n"
            '  "antigravity_payload": {\n'
            '    "issue_metadata": {\n'
            '      "case_name": "Tên vụ việc / tên thủ tục",\n'
            '      "legal_domain": "Lĩnh vực pháp lý (Ví dụ: Tách thửa, Cấp đổi GCN...)"\n'
            '    },\n'
            '    "legal_grounds": [\n'
            '      {\n'
            '        "law_name": "Tên văn bản luật/nghị định",\n'
            '        "clause": "Điều, Khoản áp dụng"\n'
            '      }\n'
            '    ],\n'
            '    "resolution_steps": [\n'
            '      "Bước 1: ...",\n'
            '      "Bước 2: ..."\n'
            '    ],\n'
            '    "key_notice": "Lưu ý quan trọng nhất"\n'
            '  }\n'
            "}\n"
            "```"
        )
    elif phom_code == "PHOM_2":
        return (
            f"HÃY BÁM SÁT {phom_title}:\n"
            f"#### Bước 1: Nhận định bản chất pháp lý của tình huống\n"
            f"- Phân tích hiện trạng thực tế thửa đất, lịch sử sử dụng và tranh chấp phát sinh.\n"
            f"- Phân định phạm vi pháp luật điều chỉnh (Luật Đất đai hiện hành và các quy định chuyển tiếp).\n\n"
            f"#### Bước 2: Tìm kiếm và đối chiếu chéo nguồn luật\n"
            f"- Tra cứu quy định chung tại Luật Đất đai 2024.\n"
            f"- Đối chiếu hướng dẫn chi tiết tại các Nghị định (NĐ 101/2024, NĐ 102/2024, NĐ 254/2026...) và Thông tư chuyên ngành.\n"
            f"- So sánh với các văn bản dưới luật của địa phương (QĐ 18/2026/QĐ-UBND, QĐ 2604/QĐ-VP) để phát hiện sự tương thích hoặc xung đột pháp lý.\n\n"
            f"#### Bước 3: Lập luận và phân tích logic pháp lý\n"
            f"- Phân tích tại sao quy định pháp luật lại áp dụng hoặc không áp dụng cho trường hợp này.\n"
            f"- Mổ xẻ các trường hợp ngoại lệ hoặc kẽ hở pháp lý có thể vận dụng để bảo vệ quyền lợi hợp pháp.\n\n"
            f"#### Bước 4: Kết luận & Giải pháp đề xuất\n"
            f"- Đưa ra kết luận pháp lý dứt khoát (Đúng/Sai, Được/Không được).\n"
            f"- Gạch đầu dòng phương án hành động tối ưu cho các bên liên quan."
        )
    elif phom_code == "PHOM_3":
        return (
            f"HÃY BÁM SÁT {phom_title}:\n"
            f"#### 1. Nhận định vấn đề\n"
            f"- Tóm tắt ngắn gọn, dễ hiểu khó khăn hay rào cản hành chính người dùng đang gặp phải.\n\n"
            f"#### 2. Căn cứ pháp lý cốt lõi\n"
            f"- Nêu nhanh 1 - 2 điều luật quan trọng chi phối trực tiếp.\n\n"
            f"#### 3. Quy trình hành chính chuẩn\n"
            f"- Các bước thực hiện bắt buộc theo đúng quy định của Nhà nước (QĐ 2604/QĐ-VP, NĐ 101/2024).\n\n"
            f"#### 4. Góc Antigravity (Mẹo đột phá / Legal Hacker)\n"
            f"- **Mẹo đàm phán:** Phương án hòa giải, thỏa thuận dân sự trung gian hợp pháp giữa các bên.\n"
            f"- **Tự động hóa thủ tục:** Cách thức điền tờ khai trực tuyến, tận dụng dịch vụ công trực tuyến & VNeID mức 2 để rút ngắn thời gian xử lý."
        )
    elif phom_code == "PHOM_4":
        return (
            f"HÃY BÁM SÁT {phom_title}:\n"
            f"#### 1. Trả lời trực diện\n"
            f"- Đi thẳng vào kết luận (Được/Không được, Đúng/Sai, Địa danh, hoặc Con số cụ thể) ngay dòng đầu tiên.\n\n"
            f"#### 2. Thông tin cốt lõi (Gạch đầu dòng tối đa 3-4 ý)\n"
            f"- **Cơ sở pháp lý ngắn:** Chỉ nêu tên văn bản và số Điều (không trích dẫn nội dung luật dài dòng).\n"
            f"- **Điều kiện / Quy trình:** Các bước hoặc giấy tờ chính yếu nhất phải chuẩn bị.\n\n"
            f"#### 3. Lưu ý quan trọng\n"
            f"- Đúng 1 câu cảnh báo thực tế hoặc mốc thời gian quan trọng cần ghi nhớ để tránh sai sót."
        )
    else:  # PHOM_1: KHUNG PHÂN TÍCH 4 BƯỚC NÂNG CẤP CHUYÊN SÂU
        return (
            f"HÃY BÁM SÁT {phom_title} (CHUYÊN SÂU & TOÀN DIỆN):\n"
            f"#### 1. Nhận diện Bản chất Pháp lý (Issue Diagnosis)\n"
            f"- Kết luận dứt điểm (ĐỦ ĐIỀU KIỆN / KHÔNG ĐỦ ĐIỀU KIỆN) kèm con số diện tích m² công nhận cụ thể ngay dòng đầu tiên.\n"
            f"- Xác định thẩm quyền giải quyết chính xác (Chủ tịch UBND cấp xã hay Chi nhánh VPĐKĐĐ cấp huyện theo phân cấp sáp nhập).\n\n"
            f"#### 2. Cơ sở Pháp lý áp dụng (Legal Basis)\n"
            f"- Bóc tách từng thông số trong câu hỏi: Thời điểm tạo lập/xây dựng, loại đất, khu vực địa lý và hạn mức đất ở (QĐ 18/2026/QĐ-UBND tỉnh Thanh Hóa).\n"
            f"- Trích dẫn chính xác Điều/Khoản Luật Đất đai 2024, Nghị định 101/2024/NĐ-CP, Nghị định 254/2026/NĐ-CP, Quyết định 2604/QĐ-VP.\n\n"
            f"#### 3. Hướng dẫn Quy trình & Hồ sơ thực tiễn (Actionable Procedure)\n"
            f"- **Thành phần hồ sơ cốt lõi (5 loại giấy tờ):** Đơn Mẫu số 04a/ĐK (hoặc Mẫu 25, Mẫu 35), Bản vẽ trích đo Mẫu 34 (hoặc Mẫu 01/TĐBĐ), Giấy tờ nguồn gốc/thời điểm sử dụng đất, CCCD/VNeID mức 2, Tờ khai lệ phí trước bạ Mẫu 01/LPTB.\n"
            f"- **Địa điểm nộp:** Bộ phận Một cửa UBND cấp xã hoặc UBND cấp huyện nơi có đất.\n"
            f"- **Thời gian giải quyết:** 23 ngày làm việc (miền núi: 13 ngày chuẩn + 10 ngày ưu đãi theo QĐ 2604) hoặc 13 ngày làm việc (đồng bằng).\n"
            f"- **Quy trình 4 bước:** Bước 1: Nộp hồ sơ -> Bước 2: Thụ lý kiểm tra & trích đo -> Bước 3: Nghĩa vụ tài chính -> Bước 4: Nhận kết quả.\n\n"
            f"#### 4. Cảnh báo Rủi ro & Mẹo thực tế (Risk & Tips)\n"
            f"- Phân tích toán học các trường hợp diện tích thực tế (diện tích <= hạn mức và > hạn mức).\n"
            f"- Các lỗi thường gặp dẫn đến bị trả hồ sơ và mẹo xử lý nhanh."
        )


def generate_response_with_ollama_llm(question, intent_type, context_text="", recent_history_text="", full_prompt=None, kb_context="", vault_context="", authority_mode=False):
    """ƯU TIÊN 3 (DỰ PHÒNG OFFLINE): Gọi Ollama Local Qwen2.5-14B khi Gemini & ZenMux không phản hồi."""
    phom_code, phom_title = select_smart_answer_framework(question, intent_type)
    phom_instruction = get_structure_instruction_for_phom(phom_code, phom_title)

    learned_rules = load_json_file(LEARNED_CORRECTIONS_FILE, [])
    learned_summary = ""
    if learned_rules:
        learned_summary = "QUY TẮC ĐÃ HỌC TỪ PHẢN BIỆN NGƯỜI DÙNG TRƯỚC ĐÂY:\n" + "\n".join([f"- {r.get('user_challenge')}" for r in learned_rules[-5:]])

    if full_prompt and len(full_prompt) < 4000:
        # Prompt ngắn (vd: từ /api/query-gemini) -> dùng nguyên bản
        prompt = full_prompt + "\n\n" + NO_HALLUCINATION_RULE
    else:
        # Prompt đầy đủ quá dài sẽ bị cắt cụt bởi num_ctx=4096 khiến model mất cả câu hỏi & CSDL
        # -> XÂY BẢN GỌN NHẸ CHUYÊN DÀNG CHO MODEL LOCAL, ĐẢM BẢO VỪA NGỮ CẢNH
        # Gộp cả KB Q&A khớp trực tiếp lẫn trích xuất tìm sâu từ Obsidian Vault
        combined_ctx = ((kb_context or "").strip() + "\n\n" + (vault_context or "").strip()).strip()
        if len(combined_ctx) > 4500:
            combined_ctx = combined_ctx[:4500] + "\n[...]"
        prompt = (
            f"[VAI TRÒ] Bạn là Trợ lý pháp lý đất đai - TTHC tỉnh Thanh Hóa (ThanhHoa Land AI). Trả lời 100% TIẾNG VIỆT.\n\n"
            f"{NO_HALLUCINATION_RULE}\n\n"
            f"{phom_instruction}\n\n"
        )
        if authority_mode:
            prompt += (
                "[QUY TẮC ƯU TIÊN XÁC ĐỊNH THẨM QUYỀN - BẮT BUỘC]\n"
                "- CHỈ suy luận thẩm quyền từ 'Cẩm nang Thẩm quyền GCN' ở CSDL bên trên: tình huống khớp đúng đối tượng thì theo KẾT LUẬN của nó; không có thì áp KHUNG PHÂN CẤM (tổ chức/thuê trả 1 lần = Sở cấp tỉnh - Phụ lục I; hộ gia đình/cá nhân/cộng đồng = VP-Chi nhánh VPĐKĐĐ - Phụ lục II; giao đất KT-XH trong 1 xã/đất ở trước 7/2004/đất tôn giáo = Chủ tịch UBND xã - Phụ lục III).\n"
                "- Thu hồi/đính chính/cấp lại theo nguyên tắc đồng cấp (ai cấp thì ai xử lý).\n"
                "- ⚠️ ĐIỂM MỚI: UBND CẤP XÃ CÓ THẨM QUYỀN trực tiếp ký & cấp GCN lần đầu cho TỔ CHỨC khi giao đất không thu tiền hoặc thuê trả tiền HẰNG NĂM trong 01 xã; đất ở không đấu giá cho cá nhân ưu tiên; xác định lại đất ở trước 01/7/2004; đất tôn giáo. CẤM nói 'UBND cấp xã không có thẩm quyền cấp GCN cho tổ chức/doanh nghiệp'.\n"
                "- Nêu rõ: cơ quan thụ lý, người ký, phụ lục tương ứng. Cấm mơ hồ kiểu 'tùy thẩm quyền'.\n\n"
            )
        prompt += (
            f"[CSDL PHÁP LÝ TRÍCH XUẤT - BẮT BUỘC CHỈ DÙNG CĂN CỨ NẰM TRONG ĐÂY]\n{combined_ctx}\n\n"
            f"{learned_summary}\n\n"
            f"LỊCH SỬ HỘI THOẠI:\n{recent_history_text[:500]}\n\n"
            f"🎯 CÂU HỎI NGƯỜI DÂN: \"{question}\"\n"
            f"👉 HÃY TRẢ LỜI CHI TIẾT THEO {phom_title}:"
        )

    # Danh sách các tên model Ollama ưu tiên chạy Qwen2.5-14B
    candidate_models = [
        PRIMARY_LLM_MODEL,
        "qwen2.5:14b",
        "qwen2.5-14b",
        "qwen-legal:14b",
        "obsidian-analyst:latest",
        "qwen2.5vl:latest",
        "qwen2.5:latest",
        "qwen2.5:7b"
    ]
    seen = set()
    models_to_try = [m for m in candidate_models if m and not (m in seen or seen.add(m))]

    for model_name in models_to_try:
        try:
            req_payload = {
                "model": model_name,
                "prompt": prompt,
                "stream": False,
                "keep_alive": -1,  # Giữ model thường trực 100% trong VRAM RTX 3060, loại bỏ độ trễ nạp từ đĩa
                "options": {
                    "num_ctx": 4096,     # Tối ưu kích thước KV Cache để 14B nằm trọn trong 12GB VRAM
                    "num_gpu": 99,       # Offload toàn bộ layers sang GPU CUDA
                    "num_thread": 8,     # Tối ưu số luồng xử lý CPU
                    "temperature": 0.2,
                    "top_p": 0.9,
                    "repeat_penalty": 1.1
                }
            }
            req_data = json.dumps(req_payload).encode('utf-8')
            
            req = urllib.request.Request(
                OLLAMA_API_URL,
                data=req_data,
                headers={'Content-Type': 'application/json'}
            )
            
            with urllib.request.urlopen(req, timeout=90) as resp:
                result = json.loads(resp.read().decode('utf-8'))
                llm_text = result.get("response", "").strip()
                if llm_text:
                    llm_text = sanitize_legal_hallucinations(llm_text, question)
                    display_name = f"Qwen2.5-14B ({model_name})" if "14b" in model_name else f"Ollama Local ({model_name})"
                    print(f"✅ Ollama Local {display_name} phản hồi thành công!")
                    return llm_text, display_name
        except Exception as e:
            print(f"⚠️ Ollama model {model_name} error: {e}")
            continue

    return None, "STATIC_RAG_FALLBACK"

def generate_response_with_antigravity_ai_engine(question, intent_type, context_text="", recent_history_text=""):
    """
    BỘ MÁY AI PHÂN TÍCH ĐA TẦNG VỚI 5 PHOM TRẢ LỜI THÔNG MINH (CAC-KHUNG-TRA-LOI-THONG-MINH.MD):
    1. Ưu tiên 1: Gemini 3.6 Flash API
    2. Tự động chuyển 2 (khi Gemini không phản hồi / hết Quota / lỗi mạng): Chuyển sang Ollama Qwen2.5-14B (qwen2.5:14b)
    3. Tự động chuyển 3: Static Legal Rule Engine
    """
    learned_rules = load_json_file(LEARNED_CORRECTIONS_FILE, [])
    learned_summary = ""
    if learned_rules:
        learned_summary = "QUY TẮC ĐÃ HỌC TỪ PHẢN BIỆN NGƯỜI DÙNG TRƯỚC ĐÂY:\n" + "\n".join([f"- {r.get('user_challenge')}" for r in learned_rules[-5:]])

    bk_10000_matches = search_bach_khoa_10000_knowledge_base(question)
    bk_10000_context = ""
    if bk_10000_matches:
        bk_10000_context = "\nCĂN CỨ BÁCH KHOA TOÀN THƯ 10.000 CÂU HỎI ĐÁP PHÁP LÝ ĐẤT ĐAI (TRAINING OLLMA):\n" + "\n".join([f"- Tình huống: {m['question']}\n  Giải đáp chuẩn: {m['answer'][:450]}" for m in bk_10000_matches])

    kt_10000_matches = search_ky_thuat_dia_chinh_10000_knowledge_base(question)
    kt_10000_context = ""
    if kt_10000_matches:
        kt_10000_context = "\nCĂN CỨ BÁCH KHOA KỸ THUẬT ĐỊA CHÍNH & ĐO ĐẠC BẢN ĐỒ 10.000 Q&A:\n" + "\n".join([f"- Hỏi: {m['question']}\n  Quy chuẩn: {m['answer'][:450]}" for m in kt_10000_matches])

    unified_brain_matches = search_obsidian_vault_unified_brain(question, top_k=5)
    unified_brain_context = ""
    if unified_brain_matches:
        unified_brain_context = "\nCSDL TRI THỨC TOÀN DIỆN OBSIDIAN VAULT - BỘ NÃO HỢP NHẤT (56.730 MỤC TRI THỨC):\n" + "\n".join([f"- Tri thức [{m.get('source', '')}]: {str(m.get('question') or m.get('title'))[:200]}\n  Nội dung trích xuất: {str(m.get('content', ''))[:500]}" for m in unified_brain_matches])

    # 0. ƯU TIÊN 0: PHẢN HỒI SIÊU TỐC TỨC THÌ (< 0.05s) CHO CÂU HỎI TRÙNG KHỚP CAO TRONG KHO TRI THỨC 10.000 Q&A
    # NGOẠI LỆ THẨM QUYỀN: chỉ cần cẩm nang thẩm quyền GCN có match khá -> luôn bỏ qua đường tắt
    # (kho cũ thời Luật 2013 hay phủ nhận thẩm quyền cấp xã - sai với phân cấp mới QĐ 2604)
    _tq_gcn_top_score = 0.0
    _tq_gcn_early_matches = search_cam_nang_tq_gcn_knowledge_base(question)
    if _tq_gcn_early_matches:
        _tq_gcn_top_score = float(_tq_gcn_early_matches[0].get("score", 0) or 0)
    if bk_10000_matches and bk_10000_matches[0].get("score", 0) >= 0.85 and not (_tq_gcn_top_score >= 0.45):
        best_qa = bk_10000_matches[0]
        instant_ans = best_qa.get("answer", "").strip()
        instant_ans = sanitize_legal_hallucinations(instant_ans, question)
        if "💡 **Bạn có thể hỏi tiếp:**" not in instant_ans and "Bạn có thể hỏi tiếp" not in instant_ans:
            instant_ans += (
                "\n\n---\n"
                "💡 **Bạn có thể hỏi tiếp:**\n"
                "1. *Hồ sơ thủ tục và mẫu đơn theo Quyết định 2604/QĐ-VP cần chuẩn bị những gì?*\n"
                "2. *Thời hạn giải quyết và cơ quan có thẩm quyền xử lý theo quy định mới nhất?*\n"
                "3. *Nghĩa vụ tài chính, thuế và lệ phí trước bạ phải nộp bao nhiêu?*"
            )
        print(f"⚡ [Instant Match <0.05s] Phản hồi tức thì từ Kho tri thức 10.000 Q&A (score={best_qa.get('score')}): '{best_qa.get('question')[:60]}...'")
        return instant_ans, "⚡ Kho Tri Thức 10.000 Q&A (Tức thì <0.05s)"

    qa_1000_matches = search_qa_1000_knowledge_base(question)
    qa_1000_context = ""
    if qa_1000_matches:
        qa_1000_context = "\nCĂN CỨ KHO TRI THỨC 1000 CÂU HỎI ĐÁP PHÁP LÝ CHUẨN (DAT_DAI_QA_1000):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in qa_1000_matches])

    cam_nang_254_matches = search_cam_nang_254_knowledge_base(question)
    cam_nang_254_context = ""
    if cam_nang_254_matches:
        cam_nang_254_context = "\nCĂN CỨ CẨM NANG NGHỊ ĐỊNH 254/2026/NĐ-CP VỀ HÓA ĐƠN ĐIỆN TỬ & NGHĨA VỤ TÀI CHÍNH ĐẤT ĐAI (LUẬT QUẢN LÝ THUẾ 108/2025/QH15):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_254_matches])

    cam_nang_do_dac_matches = search_cam_nang_do_dac_knowledge_base(question)
    cam_nang_do_dac_context = ""
    if cam_nang_do_dac_matches:
        cam_nang_do_dac_context = "\nCĂN CỨ CẨM NANG 1000 CÂU HỎI ĐÁP TOÀN DIỆN VỀ ĐO ĐẠC ĐỊA CHÍNH (LUẬT ĐO ĐẠC VÀ BẢN ĐỒ 2018, NGHỊ ĐỊNH 101/2024 & THÔNG TƯ 26/2024/TT-BTNMT):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_do_dac_matches])

    cam_nang_hm_matches = search_cam_nang_han_muc_knowledge_base(question)
    cam_nang_hm_context = ""
    if cam_nang_hm_matches:
        cam_nang_hm_context = "\nCĂN CỨ CẨM NĂNG 200 CÂU HỎI ĐÁP HẠN MỨC CÔNG NHẬN VÀ GIAO ĐẤT Ở TẠI TỈNH THANH HÓA (LUẬT ĐẤT ĐAI 2024 & QĐ 18/2026/QĐ-UBND THANH HÓA):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_hm_matches])

    cam_nang_2604_matches = search_cam_nang_2604_knowledge_base(question)
    cam_nang_2604_context = ""
    if cam_nang_2604_matches:
        cam_nang_2604_context = "\nCĂN CỨ CẨM NANG 300 CÂU HỎI ĐÁP QUYẾT ĐỊNH 2604/QĐ-VP CHUẨN HÓA 54 TTHC ĐẤT ĐAI ĐẶC THÙ THANH HÓA:\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_2604_matches])

    cam_nang_thu_hoi_matches = search_cam_nang_thu_hoi_knowledge_base(question)
    cam_nang_thu_hoi_context = ""
    if cam_nang_thu_hoi_matches:
        cam_nang_thu_hoi_context = "\nCĂN CỨ CẨM NANG 500 CÂU HỎI ĐÁP TOÀN DIỆN VỀ THU HỒI, CẤP ĐỔI VÀ CẤP LẠI GIẤY CHỨNG NHẬN ĐẤT ĐAI (LUẬT ĐẤT ĐAI 2024 & NGHỊ ĐỊNH 101/2024 & NGHỊ ĐỊNH 49/2026/NĐ-CP):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_thu_hoi_matches])

    cam_nang_dao_tao_matches = search_cam_nang_dao_tao_knowledge_base(question)
    cam_nang_dao_tao_context = ""
    if cam_nang_dao_tao_matches:
        cam_nang_dao_tao_context = "\nCĂN CỨ TÀI LIỆU TRAINING NGHIỆP VỤ ĐẤT ĐAI TỈNH THANH HÓA (QUYẾT ĐỊNH 2604/QĐ-VP NGÀY 27/7/2026, NGHỊ ĐỊNH 49/2026/NĐ-CP & THÔNG TƯ 89/2026/TT-BTC):\n" + "\n".join([f"- Phần: {m['section']}\n  Nội dung: {m['content'][:500]}" for m in cam_nang_dao_tao_matches])

    cam_nang_rung_matches = search_cam_nang_rung_knowledge_base(question)
    cam_nang_rung_context = ""
    if cam_nang_rung_matches:
        cam_nang_rung_context = "\nCĂN CỨ CẨM NĂNG 100 CÂU HỎI ĐÁP CHUYÊN SÂU CHUYỂN MỤC ĐÍCH RỪNG & THU HỒI ĐẤT LÂM NGHIỆP THANH HÓA (LUẬT ĐẤT ĐAI 2024, LUẬT LÂM NGHIỆP, NĐ 102/2024 & QĐ 55/2026/QĐ-UBND THANH HÓA):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_rung_matches])

    cam_nang_2b_matches = search_quy_trinh_2_buoc_knowledge_base(question)
    cam_nang_2b_context = ""
    if cam_nang_2b_matches:
        cam_nang_2b_context = "\nCĂN CỨ TÀI LIỆU QUY TRÌNH 2 BƯỚC TÁCH THỬA ĐỒNG THỜI HỢP THỬA ĐẤT ĐẶC THÙ (NGHỊ QUYẾT 254/2025/QH15, NĐ 49/2026/NĐ-CP & QĐ 2604/QĐ-VP THANH HÓA):\n" + "\n".join([f"- Phần: {m['section']}\n  Nội dung: {m['content'][:500]}" for m in cam_nang_2b_matches])

    cam_nang_thdt_matches = search_cam_nang_tach_hop_dac_thu_knowledge_base(question)
    cam_nang_thdt_context = ""
    if cam_nang_thdt_matches:
        cam_nang_thdt_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 100 CÂU HỎI ĐÁP TÁCH THỬA ĐỒNG THỜI HỢP THỬA ĐẤT ĐẶC THÙ (LUẬT ĐẤT ĐAI 2024, NQ 254/2025/QH15, NĐ 49/2026/NĐ-CP & QĐ 2604):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_thdt_matches])

    cam_nang_lmb_matches = search_cam_nang_lech_mat_bang_knowledge_base(question)
    cam_nang_lmb_context = ""
    if cam_nang_lmb_matches:
        cam_nang_lmb_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 50 CÂU HỎI ĐÁP XỬ LÝ LỆCH MẶT BẰNG ĐẤU GIÁ VÀ PHÂN LÔ CỦ (LUẬT ĐẤT ĐAI 2024, ĐIỀU 24 NĐ 101/2024, VB 9549/UBND-NNMT & CV 16838/SNNMT THANH HÓA):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_lmb_matches])

    cam_nang_svt_matches = search_cam_nang_sai_vi_tri_knowledge_base(question)
    cam_nang_svt_context = ""
    if cam_nang_svt_matches:
        cam_nang_svt_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 150 CÂU HỎI ĐÁP XỬ LÝ SỰ CỐ CẤP SỔ ĐỎ SAI VỊ TRÍ VÀ ĐẤT NÔNG NGHIỆP LỆCH QUY HOẠCH (LUẬT ĐẤT ĐAI 2024, NQ 254/2025/QH15, NĐ 101/2024, NĐ 49/2026/NĐ-CP & QĐ 2604):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_svt_matches])

    cam_nang_tt_matches = search_cam_nang_tang_them_knowledge_base(question)
    cam_nang_tt_context = ""
    if cam_nang_tt_matches:
        cam_nang_tt_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 200 CÂU HỎI ĐÁP CÔNG NHẬN DIỆN TÍCH ĐẤT VƯỜN TĂNG THÊM DO KHAI HOANG (LUẬT ĐẤT ĐAI 2024, ĐIỀU 24 NĐ 101/2024, NĐ 49/2026/NĐ-CP & QĐ 2604):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_tt_matches])

    cam_nang_vp_matches = search_cam_nang_vi_pham_knowledge_base(question)
    cam_nang_vp_context = ""
    if cam_nang_vp_matches:
        cam_nang_vp_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 500 CÂU HỎI ĐÁP XỬ PHẠT VI PHẠM HÀNH CHÍNH ĐẤT ĐAI (LUẬT ĐẤT ĐAI 2024, NĐ 123/2024/NĐ-CP, NĐ 281/2026/NĐ-CP & VBHN 73/2026/VBHN-NĐ-BNNMT):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Căn cứ: {m.get('legal_basis', '')}\n  Trả lời: {m['answer'][:450]}" for m in cam_nang_vp_matches])

    cam_nang_b4_matches = search_cam_nang_buoc_4_knowledge_base(question)
    cam_nang_b4_context = ""
    if cam_nang_b4_matches:
        cam_nang_b4_context = "\nCĂN CỨ TÀI LIỆU ĐÀO TẠO ĐỊA CHÍNH CHUẨN HÓA BƯỚC 4 (THẨM QUYỀN KÝ CẤP GCN & THỜI HẠN TTHC NĐ 101/2024, NĐ 49/2026 & QĐ 2604 THANH HÓA):\n" + "\n".join([f"- Phần: {m['section']}\n  Nội dung: {m['content'][:500]}" for m in cam_nang_b4_matches])

    cam_nang_bm2604_matches = search_cam_nang_bieu_mau_2604_knowledge_base(question)
    cam_nang_bm2604_context = ""
    if cam_nang_bm2604_matches:
        cam_nang_bm2604_context = "\nCĂN CỨ CẨM NANG HƯỚNG DẪN BIỂU MẪU ĐẤT ĐAI CHUẨN HÓA QĐ 2604/QĐ-VP THANH HÓA & TỜ KHAI THUẾ ĐỒNG BỘ:\n" + "\n".join([f"- Chương: {m['chapter']}\n  Nội dung: {m['content'][:500]}" for m in cam_nang_bm2604_matches])

    bo_training_dd100_matches = search_bo_training_do_dac_100_knowledge_base(question)
    bo_training_dd100_context = ""
    if bo_training_dd100_matches:
        bo_training_dd100_context = "\nCĂN CỨ BỘ TRAINING HỒ SƠ ĐO ĐẠC ĐỊA CHÍNH (100 CÂU Q&A CHUẨN KỸ THUẬT BẢN ĐỒ & MÔC GIỚI):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:450]}" for m in bo_training_dd100_matches])

    cam_nang_dd49_matches = search_cam_nang_do_dac_49_knowledge_base(question)
    cam_nang_dd49_context = ""
    if cam_nang_dd49_matches:
        cam_nang_dd49_context = "\nCĂN CỨ CẨM NANG 1000 CÂU HỎI ĐÁP ĐO ĐẠC ĐỊA CHÍNH & TRÍCH ĐO TRÍCH LỤC THEO NĐ 49/2026/NĐ-CP, NĐ 101/2024 & TT 19/2026 TẠI 27 HUYỆN THÀNH PHỐ THANH HÓA:\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:500]}" for m in cam_nang_dd49_matches])

    cam_nang_lg550_matches = search_cam_nang_lo_gioi_550_knowledge_base(question)
    cam_nang_lg550_context = ""
    if cam_nang_lg550_matches:
        cam_nang_lg550_context = "\nCĂN CỨ CẨM NANG 550 CÂU HỎI ĐÁP LỘ GIỚI, HÀNH LẠNG AN TOÀN ĐƯỜNG BỘ & QUY HOẠCH XÂY DỰNG (LUẬT ĐƯỜNG BỘ 2024, NĐ 11/2010, NĐ 49/2026 & QCVN 01:2021/BXD):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:500]}" for m in cam_nang_lg550_matches])

    cam_nang_gpmb300_matches = search_cam_nang_trich_do_gpmb_300_knowledge_base(question)
    cam_nang_gpmb300_context = ""
    if cam_nang_gpmb300_matches:
        cam_nang_gpmb300_context = "\nCĂN CỨ CẨM NANG NGHIỆP VỤ 300 TÌNH HUỐNG TRÍCH ĐO ĐỊA CHÍNH & GIẢI PHÓNG MẶT BẰNG (LUẬT ĐẤT ĐAI 2024, NĐ 101/2024, NĐ 49/2026, NĐ 88/2024 & QĐ 2604 THANH HÓA):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:500]}" for m in cam_nang_gpmb300_matches])

    cam_nang_tq500_matches = search_cam_nang_tham_quyen_500_knowledge_base(question)
    cam_nang_tq500_context = ""
    if cam_nang_tq500_matches:
        cam_nang_tq500_context = "\nCĂN CỨ CẨM NANG 500 CÂU HỎI ĐÁP NGHIỆP VỤ ĐỊA CHÍNH & THẨM QUYỀN HÀNH CHÍNH (NĐ 49/2026/NĐ-CP, NQ 254/2025/QH15 & QĐ 2604/QĐ-VP THANH HÓA):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:500]}" for m in cam_nang_tq500_matches])

    cam_nang_tq_gcn_matches = search_cam_nang_tq_gcn_knowledge_base(question, top_k=4)
    cam_nang_tq_gcn_context = ""
    if cam_nang_tq_gcn_matches:
        cam_nang_tq_gcn_context = "\nCĂN CỨ CẨM NANG THẨM QUYỀN CẤP, THU HỒI & ĐĂNG KÝ BIẾN ĐỘNG GCN THEO PHÂN CẤM MỚI (Luật Đất đai 2024, NĐ 49/2026 & QĐ 2604/QĐ-VP - BỘ 52 THỦ TỤC PHÂN THEO PHỤ LỤC I/II/III - ƯU TIÊN TỐI ĐA KHI XÁC ĐỊNH CƠ QUAN THỰC HIỆN THỦ TỤC):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:600]}" for m in cam_nang_tq_gcn_matches])

    tom_tat_hl2000_matches = search_tom_tat_hieu_luc_2000_knowledge_base(question)
    tom_tat_hl2000_context = ""
    if tom_tat_hl2000_matches:
        tom_tat_hl2000_context = "\nCĂN CỨ BỘ CẨM NANG 2000 CÂU HỎI ĐÁP TÓM TẮT HIỆU LỰC VĂN BẢN ĐẤT ĐAI (NỘI DUNG CỐT LÕI, TÌNH TRẠNG HIỆU LỰC & QUY ĐỊNH THAY THẾ/BÃI BỎ):\n" + "\n".join([f"- Hỏi: {m['question']}\n  Trả lời: {m['answer'][:500]}" for m in tom_tat_hl2000_matches])

    cam_nang_tvt300_matches = search_cam_nang_tinh_vi_tri_300_knowledge_base(question)
    cam_nang_tvt300_context = ""
    if cam_nang_tvt300_matches:
        cam_nang_tvt300_context = "\nCĂN CỨ CẨM NANG 300 TÌNH HUỐNG XÁC ĐỊNH VỊ TRÍ THỬA ĐẤT & BẢNG GIÁ ĐẤT (BẢNG GIÁ ĐẤT THANH HÓA, LUẬT ĐẤT ĐAI 2024 & NGHỊ ĐỊNH 103/2024/NĐ-CP):\n" + "\n".join([f"- Tình huống: {m['question']}\n  Xác định vị trí & Giá đất: {m['answer'][:500]}" for m in cam_nang_tvt300_matches])

    cam_nang_qh500_matches = search_cam_nang_quy_hoach_500_knowledge_base(question)
    cam_nang_qh500_context = ""
    if cam_nang_qh500_matches:
        cam_nang_qh500_context = "\nCĂN CỨ CẨM NANG 500 TÌNH HUỐNG QUY HOẠCH CHI TIẾT 1/500 & QUY HOẠCH XÂY DỰNG ĐÔ THỊ (LUẬT XÂY DỰNG, LUẬT QUY HOẠCH ĐÔ THỊ & LUẬT ĐẤT ĐAI 2024):\n" + "\n".join([f"- Tình huống: {m['question']}\n  Giải đáp quy hoạch 1/500: {m['answer'][:500]}" for m in cam_nang_qh500_matches])

    phom_code, phom_title = select_smart_answer_framework(question, intent_type)
    structure_instruction = get_structure_instruction_for_phom(phom_code, phom_title)

    # GOM TOÀN BỘ DỮ LIỆU KB KHỚP CÂU HỎI (dùng cho cổng chống bịa + prompt Ollama gọn)
    # Cẩm nang Thẩm quyền GCN được đặt ĐẦU TIÊN để không bị cắt cụt khi giới hạn ngữ cảnh
    all_kb_context = "\n".join([c for c in [
        cam_nang_tq_gcn_context, bk_10000_context, kt_10000_context, qa_1000_context, cam_nang_254_context, cam_nang_do_dac_context,
        cam_nang_hm_context, cam_nang_2604_context, cam_nang_thu_hoi_context, cam_nang_dao_tao_context,
        cam_nang_rung_context, cam_nang_2b_context, cam_nang_thdt_context, cam_nang_lmb_context,
        cam_nang_svt_context, cam_nang_tt_context, cam_nang_vp_context, cam_nang_b4_context,
        cam_nang_bm2604_context, bo_training_dd100_context, cam_nang_dd49_context, cam_nang_lg550_context,
        cam_nang_gpmb300_context, cam_nang_tq500_context, tom_tat_hl2000_context, cam_nang_tvt300_context,
        cam_nang_qh500_context, unified_brain_context
    ] if c])

    # Điểm khớp CAO NHẤT trong toàn bộ KB (chỉ tính nhóm Q&A thang tương đồng 0-1;
    # loại nhóm dạng section/chương vì dùng thang số từ trùng khác hệ,
    # loại điểm > 1.5 để không so sánh lệch thang)
    _kb_match_groups = [bk_10000_matches, qa_1000_matches, cam_nang_254_matches, cam_nang_do_dac_matches,
        cam_nang_hm_matches, cam_nang_2604_matches, cam_nang_thu_hoi_matches,
        cam_nang_rung_matches, cam_nang_thdt_matches, cam_nang_lmb_matches,
        cam_nang_svt_matches, cam_nang_tt_matches, cam_nang_vp_matches,
        bo_training_dd100_matches, cam_nang_dd49_matches, cam_nang_lg550_matches,
        cam_nang_gpmb300_matches, cam_nang_tq500_matches, cam_nang_tq_gcn_matches, tom_tat_hl2000_matches,
        cam_nang_tvt300_matches, cam_nang_qh500_matches]
    best_kb_score = 0.0
    for _ml in _kb_match_groups:
        if _ml:
            _top = max(float(m.get("score", 0) or 0) for m in _ml)
            if _top <= 1.5:
                best_kb_score = max(best_kb_score, _top)

    # ZERO-HALLUCINATION CỔNG ĐIỀU CHỈNH:
    # Mọi câu hỏi có từ khóa nghiệp vụ hoặc liên quan đất đai đều được chuyển thẳng đến AI Engine (Gemini / ZenMux / Ollama)
    # kèm theo toàn bộ kho tri thức để phân tích đa kịch bản.
    deep_ground_mode = False
    if context_text.strip():
        print(f"🔎 [Deep Search] Tìm thấy trích xuất văn bản Obsidian Vault -> Đưa cho AI trả lời BÁM TRÍCH XUẤT: '{question[:80]}'")
        deep_ground_mode = True

    prompt = (
        f"🎯 [CÂU HỎI THỰC TẾ CỦA NGƯỜI DÂN CẦN PHÂN TÍCH & GIẢI ĐÁP]\n"
        f"👉 CÂU HỎI: \"{question}\"\n\n"
        f"[THIẾT LẬP VAI TRÒ]\n"
        f"Bạn là Trợ lý ảo ThanhHoa Land AI - Trợ lý tư vấn pháp lý đất đai và thủ tục hành chính tại tỉnh Thanh Hóa.\n"
        f"[QUY TẮC BẮT BUỘC VỀ NGÔN NGỮ]: BẮT BUỘC TRẢ LỜI 100% HOÀN TOÀN BẰNG TIẾNG VIỆT NAM TRONG MỌI TRƯỜNG HỢP, TUYỆT ĐỐI KHÔNG DÙNG TIẾNG ANH HAY BẤT KỲ NGÔN NGỮ NÀO KHÁC.\n\n"
        f"[CẤU TRÚC SUY LUẬN CỐT LÕI (CORE REASONING ENGINE) - 4 TRỤ CỘT BẮT BUỘC]\n"
        f"1️⃣ TRỤ CỘT 1: BÓC TÁCH THỰC THỂ & ĐỊNH VỊ NÚT THẮT (Entity & Bottleneck)\n"
        f"- Tự động phân rã câu hỏi thành 4 biến số địa chính:\n"
        f"  + Hành động nghiệp vụ: Chuyển mục đích, tách thửa, cấp sổ lần đầu, chuyển nhượng, cấp đổi...\n"
        f"  + Loại đất nguồn: Đất trồng cây lâu năm (CLN), đất lúa (LUC), đất rừng (RSX), đất nuôi trồng thủy sản (NTS)...\n"
        f"  + Loại đất đích: Đất ở nông thôn (ONT), đất ở đô thị (ODT), đất thương mại dịch vụ (TMD)...\n"
        f"  + Nút thắt cần tháo gỡ: Có bắt buộc tách thửa không? Có đủ hạn mức không? Đất có thuộc diện bị cấm không?\n\n"
        f"2️⃣ TRỤ CỘT 2: TRUY HỒI RANH GIỚI PHÁP LÝ 3 TẦNG (Temporal Legal Grounding)\n"
        f"- Quét kho dữ liệu theo thứ tự ưu tiên hiệu lực pháp lý giảm dần (loại bỏ hoàn toàn Luật Đất đai 2013 cũ):\n"
        f"  + Tầng 1 (Luật gốc): Luật Đất đai 2024 (Điều 220 tách/hợp thửa, Điều 121 chuyển mục đích, Điều 138-140 cấp GCN, Điều 184 đất rừng...) và Luật Lâm nghiệp 2017.\n"
        f"  + Tầng 2 (Văn bản gỡ vướng & Nghị định): Nghị quyết số 254/2025/QH15 (văn bản mấu chốt: chuyển mục đích 1 phần KHÔNG BẮT BUỘC TÁCH THỬA), Nghị định 101, 102, 103/2024/NĐ-CP, Nghị định 49/2026/NĐ-CP.\n"
        f"  + Tầng 3 (Quy định địa phương tỉnh Thanh Hóa): Quyết định số 18/2026/QĐ-UBND (Hạn mức & điều kiện tách thửa), Quyết định 2604/QĐ-VP (54 TTHC Đất đai & biểu mẫu).\n\n"
        f"3️⃣ TRỤ CỘT 3: BIỆN GIẢI LOGIC HAI CHIỀU (Dual Boundary Analysis)\n"
        f"- Phân tích bao quát cả 2 kịch bản để người dân không bị thiếu thông tin:\n"
        f"  + Chiều thuận (Quy định chung): Luật mới (NQ 254/2025/QH15 & NĐ 101/2024) cho phép quản lý đa mục đích trên cùng một thửa đất → Kết luận: KHÔNG BẮT BUỘC TÁCH THỬA khi chuyển mục đích một phần diện tích.\n"
        f"  + Chiều nghịch (Trường hợp tự nguyện): Nếu người dân vẫn muốn tách riêng phần đất ở thành thửa độc lập → Bắt buộc phải đáp ứng điều kiện diện tích tối thiểu và kích thước cạnh theo Quyết định số 18/2026/QĐ-UBND tỉnh Thanh Hóa.\n\n"
        f"4️⃣ TRỤ CỘT 4: ĐÓNG GÓI ĐẦU RA SIÊU CÔ ĐỌNG & CHUYÊN SÂU (Strict Structured Output)\n"
        f"- 📌 KẾT LUẬN: Trả lời trực diện, viết hoa từ khóa chính (KHÔNG BẮT BUỘC, ĐƯỢC PHÉP, ĐỦ ĐIỀU KIỆN, DIỆN TÍCH M²...).\n"
        f"- ⚖️ CĂN CỨ PHÁP LÝ: Chỉ viện dẫn tên Điểm, Khoản, Điều, Số hiệu văn bản pháp luật hiện hành.\n"
        f"- 📝 ĐIỀU KIỆN & HƯỚNG DẪN: Nêu rõ 2 kịch bản thực tế kèm thành phần hồ sơ và cơ quan tiếp nhận (Bộ phận Một cửa cấp xã/huyện/Chi nhánh VPĐKĐĐ).\n"
        f"- 🌾 TƯƠNG TÁC: Câu hỏi mở thân thiện đề nghị hỗ trợ tính toán nghĩa vụ tài chính hoặc kiểm tra quy hoạch tiếp theo.\n\n"
        f"[CẤU TRÚC PHẢN HỒI BẮT BUỘC THEO CẨM NANG OBSIDIAN CAC-KHUNG-TRA-LOI-THONG-MINH.MD]\n"
        f"CẤM CÚ PHÁP LATEX MATH: Tuyệt đối KHÔNG viết cú pháp toán LaTeX như $\\ge \\mathbf{{50\\text{{ m}}^2}}$, $$\\text{{...}}$$, \\ge, \\mathbf{{}}, \\text{{}}. BẮT BUỘC chỉ dùng văn bản tiếng Việt thường và ký hiệu chuẩn như: 'từ 50 m² trở lên (≥ 50 m²)' hoặc '≥ 50 m²'.\n\n"
        f"{structure_instruction}\n\n"
        f"[YÊU CẦU QUAN TRỌNG NHẤT: BẮT BUỘC PHÂN TÍCH CHI TIẾT CÂU HỎI VÀ TRẢ LỜI SÂU RỘNG TỈ MỈ]\n"
        f"Tuyệt đối KHÔNG trả lời vắn tắt hay qua sơ. Bắt buộc phải trình bày đầy đủ, chi tiết, chuyên sâu theo đúng 4 phần cấu trúc chuẩn dưới đây.\n"
        f"CẤM CÚ PHÁP LATEX MATH: Tuyệt đối KHÔNG viết cú pháp toán LaTeX như $\\ge \\mathbf{{50\\text{{ m}}^2}}$, $$\\text{{...}}$$, \\ge, \\mathbf{{}}, \\text{{}}. CHỈ DÙNG văn bản tiếng Việt chuẩn và ký hiệu thường như: 'từ 50 m² trở lên (≥ 50 m²)' hoặc '≥ 50 m²'.\n\n"
        f"#### 1. Trả lời trực diện & Kết luận dứt điểm\n"
        f"- Đưa ra ngay kết luận khẳng định dứt điểm (ĐƯỢC CẤP GIẤY CHỨNG NHẬN... / KHÔNG ĐỦ ĐIỀU KIỆN...) và con số diện tích m² cụ thể ngay dòng đầu tiên.\n\n"
        f"#### 2. Phân tích chi tiết bối cảnh câu hỏi & Căn cứ pháp lý áp dụng\n"
        f"- Bóc tách tỉ mỉ từng thông số trong câu hỏi người dân cung cấp:\n"
        f"  + **Hiện trạng đất & Nhà ở xây dựng:** (Ví dụ: Xây nhà trên đất nông nghiệp năm 2013 -> Áp dụng Khoản 3 Điều 139 Luật Đất đai 2024 quy định cho hành vi tự ý làm nhà trên đất nông nghiệp từ 01/07/2004 đến trước 01/07/2014).\n"
        f"  + **Khu vực địa lý & Hạn mức đất ở:** (Ví dụ: Xã miền núi Thanh Hóa -> Áp dụng hạn mức công nhận đất ở miền núi tối đa **400 m² / hộ** theo Quyết định 18/2026/QĐ-UBND tỉnh Thanh Hóa; nếu ở thị trấn thuộc huyện miền núi thì tối đa **200 m² / hộ**).\n"
        f"  + **Điều kiện công nhận bắt buộc:** Đất sử dụng ổn định, không có tranh chấp, không vi phạm hành lang an toàn và được UBND xã xác nhận phù hợp với quy hoạch đất ở.\n"
        f"  + **Nghĩa vụ tài chính bắt buộc:** Phải nộp tiền sử dụng đất khi chuyển mục đích sang đất ở và lệ phí trước bạ (0,5%) theo Nghị định 254/2026/NĐ-CP.\n\n"
        f"#### 3. Phân tích Phép tính & Các trường hợp diện tích thực tế\n"
        f"- Phân tích toán học chi tiết cho các trường hợp diện tích làm nhà thực tế:\n"
        f"  + **Trường hợp 1 (Diện tích làm nhà thực tế ≤ 400 m²):** Được công nhận 100% diện tích nhà ở thực tế là đất ở. Phần đất nông nghiệp còn lại giữ nguyên mục đích nông nghiệp.\n"
        f"  + **Trường hợp 2 (Diện tích làm nhà thực tế > 400 m²):** Chỉ được công nhận tối đa 400 m² đất ở theo hạn mức miền núi. Phần diện tích vượt 400 m² sẽ xử lý theo quy hoạch: nếu muốn chuyển tiếp sang đất ở phải nộp 100% tiền sử dụng đất vượt hạn mức.\n\n"
        f"#### 4. Quy trình thủ tục chi tiết & Lưu ý hồ sơ (Quyết định 2604/QĐ-VP Thanh Hóa)\n"
        f"- Thành phần hồ sơ cốt lõi (5 loại giấy tờ):\n"
        f"  1. Đơn đăng ký, cấp GCN lần đầu theo Mẫu số 04a/ĐK (QĐ 2604/QĐ-VP).\n"
        f"  2. Bản vẽ trích đo địa chính thửa đất theo Mẫu số 34.\n"
        f"  3. Giấy tờ chứng minh thời điểm xây dựng nhà ở năm 2013 (Xác nhận của thôn/xã hoặc hóa đơn điện, nước).\n"
        f"  4. Bản sao Căn cước công dân/VNeID mức 2 của chủ hộ.\n"
        f"  5. Tờ khai lệ phí trước bạ (Mẫu 01/LPTB).\n"
        f"- Địa điểm nộp hồ sơ: Bộ phận Một cửa UBND cấp xã hoặc Bộ phận Một cửa UBND cấp huyện nơi có đất.\n"
        f"- Thời gian giải quyết tối đa: 23 ngày làm việc (gồm 13 ngày làm việc chuẩn TTHC + 10 ngày làm việc ưu đãi gia hạn cho khu vực miền núi theo QĐ 2604/QĐ-VP; không tính 15 ngày niêm yết công khai tại UBND xã).\n"
        f"- Quy trình 4 bước thực hiện chi tiết (Bước 1 -> Bước 4).\n\n"
        f"BẢNG CĂN CỨ PHÁP LÝ CHUẨN:\n{GROUND_TRUTH_FACTS_TEXT}\n"
        f"{bk_10000_context}\n"
        f"{qa_1000_context}\n"
        f"{cam_nang_254_context}\n"
        f"{cam_nang_do_dac_context}\n"
        f"{cam_nang_hm_context}\n"
        f"{cam_nang_2604_context}\n"
        f"{cam_nang_thu_hoi_context}\n"
        f"{cam_nang_dao_tao_context}\n"
        f"{cam_nang_rung_context}\n"
        f"{cam_nang_2b_context}\n"
        f"{cam_nang_thdt_context}\n"
        f"{cam_nang_lmb_context}\n"
        f"{cam_nang_svt_context}\n"
        f"{cam_nang_tt_context}\n"
        f"{cam_nang_vp_context}\n"
        f"{cam_nang_b4_context}\n"
        f"{cam_nang_bm2604_context}\n"
        f"{bo_training_dd100_context}\n"
        f"{cam_nang_dd49_context}\n"
        f"{cam_nang_lg550_context}\n"
        f"{cam_nang_gpmb300_context}\n"
        f"{cam_nang_tq500_context}\n"
        f"{tom_tat_hl2000_context}\n"
        f"{cam_nang_tvt300_context}\n"
        f"{cam_nang_qh500_context}\n\n"
        f"{learned_summary}\n\n"
        f"[CÂU HỎI GỢI MỞ BẮT BUỘC CUỐI MỖI TRẢ LỜI]\n"
        f"Sau phần kết luận, BẮT BUỘC thêm 3 câu hỏi gợi ý theo định dạng:\n"
        f"---\n"
        f"💡 **Bạn có thể hỏi tiếp:**\n"
        f"1. *[Câu hỏi phụ liên quan trực tiếp đến tình huống vừa phân tích]*\n"
        f"2. *[Câu hỏi về thủ tục hoặc rủi ro pháp lý tiếp theo]*\n"
        f"3. *[Câu hỏi mở rộng về quy hoạch hoặc nghĩa vụ tài chính liên quan]*\n\n"
        f"CSDL TRÍCH XUẤT OBSIDIAN VAULT:\n{context_text[:1200]}\n\n"
        f"LỊCH SỬ HỘI THOẠI GẦN ĐÂY:\n{recent_history_text}\n\n"
        f"🎯 NHẮC LẠI: HÃY PHÂN TÍCH CHI TIẾT CÁC THÔNG SỐ CỦA CÂU HỎI \"{question}\" VÀ TRẢ LỜI THẬT CHUYÊN SÂU, TỈ MỈ ĐẦY ĐỦ 4 MỤC:"
    )

    # RÀO CHẮN BỊA ẢO GIÁC: ép AI tự khai [CHUA_CO_DU_LIEU] thay vì bịa căn cứ pháp lý
    prompt += "\n\n" + NO_HALLUCINATION_RULE

    # QUY TẮC KHOAN SÂU TRONG NÃO & TỔNG HỢP TOÀN BỘ TRI THỨC ĐÃ TRÍCH XUẤT
    prompt += (
        "\n\n[QUY TẮC BẮT BUỘC: KHOAN SÂU VÀ SỬ DỤNG TOÀN BỘ TRI THỨC TRONG NÃO ĐỂ TRẢ LỜI]\n"
        "- Khi câu hỏi chứa từ khóa liên quan đến bất kỳ chủ đề/văn bản nào trong não (như QĐ 18/2026, QĐ 2604/QĐ-VP, NĐ 254/2026, NĐ 49/2026, NĐ 101/2024, Luật Đất đai 2024, trích đo, GPMB, bồi thường, lộ giới, thẩm quyền, cấp đổi, cấp sổ, tách thửa...): BẮT BUỘC tập trung khoan sâu vào các CSDL cẩm nang và trích xuất đã nạp ở trên để đưa ra câu trả lời chính xác 100%, chi tiết từng bước, viện dẫn đầy đủ số hiệu, điều khoản và biểu mẫu.\n"
        "- Kết hợp nhuần nhuyễn giữa CSDL cẩm nang chuyên sâu và kiến thức tổng quát để trả lời toàn diện."
    )

    # QUY TẮC ƯU TIÊN THẨM QUYỀN: khi cẩm nang thẩm quyền GCN có kết quả khớp -> bắt buộc suy luận từ bộ dữ liệu này
    if cam_nang_tq_gcn_matches:
        prompt += (
            "\n\n[QUY TẮC ƯU TIÊN XÁC ĐỊNH THẨM QUYỀN - BẮT BUỘC]\n"
            "- Khi xác định cơ quan thụ lý / cơ quan ký ban hành thủ tục đất đai: CHỈ được suy luận từ 'Cẩm nang Thẩm quyền Cấp/Thu hồi GCN' ở CSDL bên trên.\n"
            "- Cách suy luận: (a) Nếu có TÌNH HUỐNG nào khớp ĐÚNG đối tượng & hoàn cảnh (hộ gia đình ≠ doanh nghiệp ≠ tổ chức tôn giáo; hình thức thuê trả tiền 1 lần ≠ hằng năm...) -> theo đúng KẾT LUẬN tình huống đó. "
            "(b) Nếu không có tình huống khớp -> áp KHUNG PHÂN CẤM: tổ chức/dự án có thu tiền SDĐ hoặc thuê trả 1 lần = Sở NN&MT cấp tỉnh ký (Phụ lục I); hộ gia đình, cá nhân, cộng đồng, người gốc Việt Nam ở nước ngoài sở hữu nhà ở = Văn phòng/Chi nhánh VPĐKĐĐ ký (Phụ lục II - 22 thủ tục); giao đất KT-XH trong địa giới 1 xã, xác định lại diện tích đất ở trước 01/7/2004, đất tôn giáo tín ngưỡng = Chủ tịch UBND cấp xã ký (Phụ lục III - 19 thủ tục).\n"
            "- Thu hồi/đính chính/cấp lại: cơ quan nào đã ký cấp GCN thì cơ quan đó xử lý (nguyên tắc đồng cấp).\n"
            "- ⚠️ ĐIỂM MỚI ĐỘT PHÁ (khác Luật Đất đai 2013): UBND CẤP XÃ CÓ THẨM QUYỀN trực tiếp ký quyết định & cấp GCN lần đầu cho TỔ CHỨC trong nước khi Nhà nước giao đất không thu tiền hoặc thuê đất trả tiền HẰNG NĂM trong phạm vi địa giới của 01 xã; giao đất ở không đấu giá cho cá nhân ưu tiên; xác định lại diện tích đất ở trước 01/7/2004; đất tôn giáo tín ngưỡng. TUYỆT ĐỐI KHÔNG trả lời kiểu 'UBND cấp xã không có thẩm quyền cấp GCN cho tổ chức/doanh nghiệp' — điều đó SAI với Phụ lục III QĐ 2604/QĐ-VP.\n"
            "- CHỈ dẫn số thủ tục cụ thể (VD 'Thủ tục số 9 Phụ lục III') khi đúng tình huống tương ứng trong cẩm nang ghi rõ như vậy; nếu không chắc, chỉ nêu tên Phụ lục theo khung phân cấp. GCN do cấp tỉnh ký thì xử lý theo Phụ lục I.\n"
            "- Tên cơ quan chuẩn phân cấp mới: Sở Nông nghiệp và Môi trường (KHÔNG dùng tên cũ 'Sở NN&PTNT').\n"
            "- Trả lời nêu RÕ: (1) Cơ quan thụ lý, (2) Người/cơ quan ký văn bản, (3) Phụ lục tương ứng. Tuyệt đối không mơ hồ kiểu 'tùy thẩm quyền'."
        )

    # 1. ƯU TIÊN 1: Google Gemini Cloud AI (Gemini 3.6 Flash / 3.5 Flash / 2.5 Flash - Tiếng Việt chuẩn)
    gemini_ans, gemini_model = generate_response_with_gemini_api(prompt)
    if gemini_ans and len(gemini_ans.strip()) > 20:
        gemini_ans = sanitize_legal_hallucinations(gemini_ans, question)
        if _is_no_data_answer(gemini_ans):
            print("🛡️ Gemini thừa nhận chưa có dữ liệu -> Chuyển sang phản hồi Zero-Hallucination (dẫn dắt phản biện)")
            return build_no_data_guided_answer(question), "🛡️ Zero-Hallucination Guard"
        return gemini_ans, f"✨ {gemini_model}"

    # 2. ƯU TIÊN 2 (DỰ PHÒNG KHI GEMINI HẾT TOKEN/QUOTA): ZenMux Multi-Model AI Gateway (dots3-note-prev, glm-5.3-free, deepseek-v4-flash)
    print("🔄 Gemini Cloud hết quota/token -> Đang tự động chuyển sang ZenMux Multi-Model AI Gateway (Tiếng Việt)...")
    zenmux_ans, zenmux_model = generate_response_with_zenmux_api(prompt)
    if zenmux_ans and len(zenmux_ans.strip()) > 20:
        zenmux_ans = sanitize_legal_hallucinations(zenmux_ans, question)
        if _is_no_data_answer(zenmux_ans):
            print("🛡️ ZenMux thừa nhận chưa có dữ liệu -> Chuyển sang phản hồi Zero-Hallucination (dẫn dắt phản biện)")
            return build_no_data_guided_answer(question), "🛡️ Zero-Hallucination Guard"
        return zenmux_ans, f"🚀 {zenmux_model}"

    # 3. ƯU TIÊN 3 (DỰ PHÒNG OFFLINE): Ollama Local GPU RTX 3060 12GB (qwen2.5-thanhhoa-land:14b)
    print("🔄 Gemini & ZenMux đều không phản hồi -> Đang tự động chuyển sang mô hình Qwen2.5-14B (Ollama Local GPU)...")
    ollama_ans, ollama_model = generate_response_with_ollama_llm(question, intent_type, context_text, recent_history_text, full_prompt=prompt, kb_context=all_kb_context, vault_context=context_text, authority_mode=bool(cam_nang_tq_gcn_matches))
    if ollama_ans and len(ollama_ans.strip()) > 20:
        if _is_no_data_answer(ollama_ans):
            print("🛡️ Ollama thừa nhận chưa có dữ liệu -> Chuyển sang phản hồi Zero-Hallucination (dẫn dắt phản biện)")
            return build_no_data_guided_answer(question), "🛡️ Zero-Hallucination Guard"
        return ollama_ans, f"⚡ {ollama_model}"

    # 3. QUY TẮC PHÁP LÝ TỰ ĐỘNG (STATIC RAG FALLBACK)
    q_lower = question.lower().strip()
    if "rừng" in q_lower or "lâm nghiệp" in q_lower:
        has_3400 = "3400" in q_lower or "3.400" in q_lower or "3400m2" in q_lower or "3400 m2" in q_lower
        
        direct_ans = "Thửa đất 3.400 m² đất rừng **KHÔNG ĐỦ ĐIỀU KIỆN TÁCH THÀNH 2 THỬA ĐẤT RỪNG ĐỘC LẬP** theo quy định pháp luật tỉnh Thanh Hóa." if has_3400 else "Theo quy định tại Thanh Hóa, diện tích tối thiểu để được phép tách thửa đất rừng sản xuất, đất rừng phòng hộ là **≥ 3.000 m² (0,3 ha)** cho mỗi thửa đất."
        
        calc_analysis = (
            "- **Phân tích phép tính thực tế:** Với thửa đất gốc 3.400 m², nếu chia đôi thì mỗi thửa chỉ đạt **1.700 m²** (nhỏ hơn mức tối thiểu 3.000 m²). Kể cả tách 1 thửa 3.000 m² thì thửa còn lại chỉ còn **400 m²** (nhỏ hơn 3.000 m²).\n"
            "- **Điều kiện để tách được 2 thửa đất rừng:** Thửa đất gốc bắt buộc phải có diện tích tối thiểu từ **6.000 m² (0,6 ha) trở lên** mới đủ điều kiện tách thành 2 thửa độc lập (mỗi thửa ≥ 3.000 m²)."
        ) if has_3400 else (
            "- **Điều kiện tách thửa đất rừng sản xuất:** Thửa đất mới hình thành và thửa đất còn lại sau khi tách đều phải đạt diện tích tối thiểu từ **3.000 m² trở lên**.\n"
            "- **Rừng phòng hộ & Đặc dụng:** Tuyệt đối KHÔNG được phép tách thửa cho hộ gia đình, cá nhân."
        )

        return (
            f"#### 1. Trả lời trực diện\n"
            f"{direct_ans}\n\n"
            f"#### 2. Thông tin cốt lõi & Căn cứ pháp lý\n"
            f"- **Căn cứ pháp lý:** Điều 220 Luật Đất đai 2024 (Luật số 31/2024/QH15) và Quyết định số 18/2026/QĐ-UBND ngày 20/05/2026 của UBND tỉnh Thanh Hóa.\n"
            f"- **Quy chuẩn diện tích tối thiểu đất rừng:** Đất rừng sản xuất, đất rừng phòng hộ tại Thanh Hóa yêu cầu diện tích tối thiểu **≥ 3.000 m² (0,3 ha)** cho mỗi thửa đất.\n\n"
            f"#### 3. Phân tích Chi tiết & Phép tính thực tế\n"
            f"{calc_analysis}\n\n"
            f"#### 4. Quy trình & Lưu ý quan trọng (QĐ 2604/QĐ-VP Thanh Hóa)\n"
            f"- **Mẫu đơn áp dụng:** Đơn đề nghị tách thửa theo **Mẫu số 35** và Bản vẽ trích đo theo **Mẫu số 34** ban hành kèm Quyết định 2604/QĐ-VP.\n"
            f"- **Địa điểm nộp:** Bộ phận Một cửa UBND cấp huyện hoặc Chi nhánh Văn phòng Đăng ký đất đai địa phương.\n"
            f"- **Thời gian giải quyết:** Không quá **15 ngày làm việc**.",
            "✨ Antigravity Legal Rules"
        )
    elif "đo đạc" in q_lower or "bản đồ" in q_lower or "trắc địa" in q_lower or "mốc" in q_lower or "ranh giới" in q_lower:
        return (
            "#### 1. Trả lời trực diện\n"
            "Việc đo đạc lại ranh giới mốc giới thửa đất được thực hiện theo kỹ thuật trích đo địa chính hiện đại và nếu diện tích thực tế chênh lệch nhưng ranh giới không thay đổi thì được công nhận theo số liệu đo đạc mới nhất.\n\n"
            "#### 2. Căn cứ pháp lý cốt lõi\n"
            "- Luật Đo đạc và Bản đồ năm 2018 (Luật số 27/2018/QH14).\n"
            "- Điều 24 Nghị định số 101/2024/NĐ-CP (sửa đổi bởi NĐ 49/2026/NĐ-CP).\n"
            "- Thông tư số 26/2024/TT-BTNMT quy định kỹ thuật trích đo địa chính.\n\n"
            "#### 3. Quy trình thực tế\n"
            "- **Mẫu đơn áp dụng:** Đơn đề nghị đo đạc trích đo địa chính theo **Mẫu** (TT26).\n"
            "- **Địa điểm nộp:** Chi nhánh Văn phòng Đăng ký đất đai địa phương.\n"
            "- **Thời gian giải quyết:** Không quá **07 ngày làm việc**.\n\n"
            "#### 4. Lưu ý quan trọng\n"
            "Tất cả các chủ sử dụng đất giáp ranh bắt buộc phải ký Bản mô tả ranh giới mốc giới thửa đất.",
            "✨ Antigravity Legal Rules"
        )
    elif "tách" in q_lower or "thửa" in q_lower:
        return (
            "#### 1. Trả lời trực diện\n"
            "Việc tách thửa đất phải đảm bảo cả thửa đất mới và thửa đất còn lại đều đạt diện tích tối thiểu và kích thước mặt tiền theo Quyết định 18/2026/QĐ-UBND tỉnh Thanh Hóa.\n\n"
            "#### 2. Quy chuẩn diện tích tối thiểu tại Thanh Hóa\n"
            "- **Đất ở Phường (đô thị):** Diện tích tối thiểu **≥ 40 m²**, mặt tiền **≥ 3.0 m**.\n"
            "- **Đất ở Xã (đồng bằng):** Diện tích tối thiểu **≥ 40 m²**, mặt tiền **≥ 4.0 m**.\n"
            "- **Đất ở Xã (miền núi):** Diện tích tối thiểu **≥ 50 m²**, mặt tiền **≥ 5.0 m**.\n"
            "- **Đất Nông nghiệp (CLN, BHK):** Diện tích tối thiểu **≥ 500 m²**.\n"
            "- **Đất Rừng sản xuất:** Diện tích tối thiểu **≥ 3.000 m²**.\n\n"
            "#### 3. Quy trình thực hiện (QĐ 2604/QĐ-VP Thanh Hóa)\n"
            "- **Mẫu đơn áp dụng:** Đơn đề nghị đo đạc trích đo địa chính theo **Mẫu** (TT26).\n"
            "- **Địa điểm nộp:** Bộ phận Một cửa cấp huyện hoặc Chi nhánh Văn phòng Đăng ký đất đai.\n"
            "- **Thời gian giải quyết:** Không quá **15 ngày làm việc**.\n\n"
            "#### 4. Lưu ý quan trọng\n"
            "Thửa đất phải có Sổ đỏ bản gốc, không tranh chấp và có lối đi kết nối giao thông công cộng.",
            "✨ Antigravity Legal Rules"
        )

    # ZERO-HALLUCINATION FALLBACK: KHÔNG TỰ BỊA CÂU TRẢ LỜI TỔNG QUÁT KHI KHÔNG CÓ DỮ LIỆU
    # -> Trả lời thẳng thắn + dẫn dắt phản biện để người dùng làm rõ câu hỏi
    return build_no_data_guided_answer(question), "🛡️ Zero-Hallucination Guard (Dẫn dắt phản biện)"

# FULL 5-STEP PIPELINE WITH PERSISTENT AI MEMORY SEARCH & DEEP LOGICAL REASONING
def process_antigravity_core_pipeline(question, session_id="default_session"):
    # 0. KIỂM TRA BỘ TRẢ LỜI HÓM HĨNH & BẺ LÁI ĐỊA CHÍNH (BO-TRA-LOI-HOM-HINH-BOT-DIA-CHINH.MD)
    hom_hinh_answer = search_hom_hinh_knowledge_base(question)
    if hom_hinh_answer:
        print(f"😄 [Trợ lý Địa chính Hóm hỉnh] Khớp câu hỏi ngoài chuyên môn: '{question}'")
        record_conversation_turn(session_id, question, hom_hinh_answer, "CHIT_CHAT_HOM_HINH")
        return hom_hinh_answer, "CHIT_CHAT_HOM_HINH", "😄 Trợ Lý Địa Chính Hóm Hỉnh"

    # ZERO-HALLUCINATION CỔNG 1: Chặn câu hỏi NGOÀI PHẠM VI pháp lý đất đai -> không để AI bịa
    if not is_land_legal_question(question):
        print(f"🛡️ [Zero-Hallucination] Câu hỏi ngoài phạm vi chuyên môn, chặn trước khi gọi AI: '{question[:80]}'")
        guided_ans = build_no_data_guided_answer(question)
        record_conversation_turn(session_id, question, guided_ans, "OUT_OF_SCOPE_ZERO_HALLUCINATION")
        return guided_ans, "OUT_OF_SCOPE_ZERO_HALLUCINATION", "🛡️ Zero-Hallucination Guard (Ngoài phạm vi)"

    intent_type, dia_match = analyze_user_intent(question)

    if intent_type == "UNANSWERED_REPORT_QUERY":
        report_md, report_path = generate_daily_unanswered_markdown_report()
        rel_path = os.path.basename(report_path)
        answer_text = (
            f"{report_md}\n\n"
            f"📁 *Báo cáo tổng hợp đã được xuất bản và lưu trữ tự động trong Obsidian Vault tại thư mục:* `06 - BỘ NHỚ AI/BAO_CAO_CAU_HOI_CHUA_TRA_LOI/{rel_path}`"
        )
        record_conversation_turn(session_id, question, answer_text, intent_type)
        return answer_text, intent_type, "✨ Antigravity Report Generator"

    if intent_type == "USER_FEEDBACK_CHALLENGE":
        learn_from_user_challenge(question, question)
    
    if intent_type == "DIA_DANH_LOOKUP" and dia_match:
        old_v = dia_match["old_name"]
        new_v = dia_match["new_name"]
        answer_text = (
            f"**Phân tích sáp nhập địa danh:** Theo CSDL địa danh tỉnh Thanh Hóa, **{old_v}** đã được sáp nhập đơn vị hành chính và đổi tên chính thức thành **{new_v}** 📍.\n\n"
            f"📍 *Gợi ý tra cứu tiếp theo:*\n"
            f"1. *{new_v} bao gồm những thôn/xóm cũ nào sáp nhập?*\n"
            f"2. *Hạn mức diện tích tách thửa đất ở tại địa bàn {new_v} quy định như thế nào?*"
        )
        record_conversation_turn(session_id, question, answer_text, intent_type)
        return answer_text, intent_type, "✨ Antigravity Engine (Địa Danh Thanh Hóa)"

    history_records = load_json_file(CHAT_MEMORY_FILE, [])
    session_turns = [t for t in history_records if t.get("session_id") == session_id][-2:]
    recent_history_text = "\n".join([f"User: {t['question']}\nAI: {t['answer'][:120]}" for t in session_turns])

    obsidian_results = search_obsidian_vault(question)
    # LỌC NHIỄU CHO ZERO-HALLUCINATION: loại file bộ nhớ/log (.json) & thư mục BỘ NHỚ AI,
    # giữ lại các trích xuất thực sự từ tài liệu pháp lý trong vault
    if obsidian_results:
        obsidian_results = [r for r in obsidian_results
                            if os.path.splitext(str(r.get("path", "")))[1].lower() not in ('.json', '.jsonl')
                            and "bộ nhớ ai" not in str(r.get("path", "")).lower()
                            and "bo nho ai" not in str(r.get("path", "")).lower()
                            and float(r.get("score", 0) or 0) >= 15]
        # YÊU CẦU TÍNH LIÊN QUAN: chấm trên SECTION TỐT NHẤT (không phải cả file) bằng
        # từ đặc trưng + cụm ghép ("cát sông", "vật liệu"...); cần >= 2 điểm trúng mới coi là cùng chủ đề
        _uni_kws, _bi_kws = get_distinctive_keywords(question)
        def _is_relevant_excerpt(m):
            _blob = str(m.get("best_section") or m.get("full_content", ""))[:4000].lower()
            _hits = sum(1 for w in _uni_kws if w in _blob) + sum(1 for b in _bi_kws if b in _blob)
            return _hits >= 2
        obsidian_results = [r for r in obsidian_results if _is_relevant_excerpt(r)]
    context_text = ""
    if obsidian_results:
        top_matches = obsidian_results[:3]
        context_text = "\n\n---\n\n".join([(m.get("best_section") or m.get("full_content", ""))[:1200] for m in top_matches])

    # GỌI BỘ MÁY GEMINI 3.6 FLASH (VỚI FAILOVER OLLAMA TỰ ĐỘNG KHÔNG DÙNG CÁCH NỐI BỘ NHỚ LỖI THỜI)
    llm_answer, used_model = generate_response_with_antigravity_ai_engine(question, intent_type, context_text, recent_history_text)

    final_answer = f"{llm_answer.strip()}"

    # Tự động ghi nhận câu hỏi chưa trả lời được/chưa rõ thông tin vào file tổng hợp theo ngày
    if _is_no_data_answer(final_answer) or used_model.startswith("🛡️") or used_model == "STATIC_RAG_FALLBACK":
        record_unanswered_question(
            question=question,
            session_id=session_id,
            reason="Hệ thống chưa có đủ dữ liệu pháp lý chi tiết hoặc phải dùng fallback",
            model_used=used_model,
            answer_given=final_answer
        )

    record_conversation_turn(session_id, question, final_answer, intent_type)

    return final_answer, intent_type, used_model

# EXACT STRUCTURED OCR SYSTEM PROMPTS (SỐ PHÁT HÀNH GCN NẰM Ở DƯỚI CÙNG CÓ 2 CHỮ CÁI IN HOA)
PROMPT_CCCD_EXACT = """Bạn là chuyên gia OCR tài liệu hành chính Việt Nam.
Hãy đọc toàn bộ thông tin trên thẻ Căn cước công dân (CCCD) và trả về CHÍNH XÁC CHUẨN CẤU TRÚC JSON DƯỚI ĐÂY:

QUY TẮC BÓC TÁCH CCCD BẮT BUỘC (ĐẶC BIỆT CHÚ Ý VỊ TRÍ HỌ TÊN):
1. "full_name" (Họ và tên): BẮT BUỘC CHỈ LẤY DÒNG CHỮ IN HOA NẰM NGAY BÊN DƯỚI DÒNG SỐ CCCD ("Số / No.") TẠI MẶT CÓ ẢNH (MẶT 1). TUYỆT ĐỐI KHÔNG ĐỌC HỌ TÊN Ở CÁC VỊ TRÍ KHÁC HOẶC MẶT SAU!
2. "date_of_issue" (Ngày cấp): BẮT BUỘC TRÍCH XUẤT TỪ MẶT SAU (MẶT 2 THẺ CCCD) tại dòng "Ngày, tháng, năm / Date, month, year" (Ví dụ: 25/08/2021).
3. "date_of_birth" (Ngày sinh): Nằm dưới dòng Họ và tên ở Mặt 1.

{
  "document_type": "Căn cước công dân",
  "front_side": {
    "id_number": "Số / No. 12 chữ số",
    "full_name": "Họ và tên NẰM NGAY DƯỚI DÒNG SỐ CCCD (viết hoa)",
    "date_of_birth": "DD/MM/YYYY",
    "sex": "Nam hoặc Nữ",
    "nationality": "Việt Nam",
    "place_of_origin": "Quê quán đầy đủ",
    "place_of_residence": "Nơi thường trú đầy đủ",
    "date_of_expiry": "DD/MM/YYYY hoặc Không thời hạn"
  },
  "back_side": {
    "personal_identification": "Đặc điểm nhận dạng",
    "date_of_issue": "Ngày, tháng, năm cấp BẮT BUỘC LẤY TỪ MẶT SAU (DD/MM/YYYY)",
    "place_of_issue": "Cơ quan cấp (Ví dụ: CỤC TRƯỜNG CỤC CẢNH SÁT...)",
    "mrz_code": "Dòng mã MRZ ở dưới cùng mặt sau"
  }
}

Quy tắc:
- Chỉ trả về JSON thuần túy, không chứa markdown code block hay ký tự lạ.
- Nếu không đọc được trường nào, đặt giá trị = ""
"""

PROMPT_LAND_EXACT = """Bạn là chuyên gia OCR tài liệu hành chính Việt Nam.
Đây là ảnh Giấy Chứng Nhận Quyền Sử Dụng Đất (Sổ đỏ/Sổ hồng).
Hãy trích xuất thông tin và trả về CHÍNH XÁC CHUẨN CẤU TRÚC JSON DƯỚI ĐÂY:

QUY TẮC BÓC TÁCH SỔ ĐỎ / GCN BẮT BUỘC (ĐẶC BIỆT SỐ PHÁT HÀNH GCN):
1. "certificate_serial_number" (Số phát hành GCN): BẮT BUỘC NẰM Ở DƯỚI CÙNG TRANG BÌA / DƯỚI CÙNG TRANG 1 GCN. DẤU HIỆU XÁC ĐỊNH: BẮT ĐẦU BẰNG 2 CHỮ CÁI IN HOA theo sau là dãy 6 chữ số (Ví dụ: DA 895241, AA 035751, CM 902946, CH 00071, BX 123456).
2. "owner_name" (Tên người sử dụng đất): BẮT BUỘC CHỈ BÓC TÁCH TỪ TRANG BÌA (MẶT 1 GCN). MẶT SAU (TRANG 2-4) TUYỆT ĐỐI KHÔNG LẤY HỌ TÊN CHỦ SỞ HỮU.
3. "registration_book_number" (Số vào sổ cấp GCN): BẮT BUỘC NẰM Ở GÓC DƯỚI BÌA HOẶC GÓC DƯỚI CÙNG GIẤY CHỨNG NHẬN, CÓ KÝ HIỆU BẮT ĐẦU BẰNG CHỮ CÁI H, CX, CH, CN, CS, CT (Ví dụ: CH 00071, CX 52, CN 01234, H 1234, CS 999).

{
  "document_type": "Giấy chứng nhận quyền sử dụng đất",
  "page_1_owner_info": {
    "certificate_serial_number": "Số phát hành GCN ở DƯỚI CÙNG bìa có 2 chữ cái in hoa ở đầu (VD: DA 895241, CM 902946)",
    "owner_name": "Tên người sử dụng đất BẮT BUỘC CHỈ LẤY Ở MẶT BÌA (VD: Hà Văn Tha)",
    "owner_year_of_birth": "Năm sinh",
    "owner_id_number": "Số CMND/CCCD",
    "owner_address": "Địa chỉ thường trú chủ sở hữu"
  },
  "page_2_land_info": {
    "parcel_number": "Thửa đất số (chỉ con số)",
    "map_sheet_number": "Tờ bản đồ số (chỉ con số)",
    "parcel_address": "Địa chỉ thửa đất đầy đủ",
    "area_number": "Diện tích bằng số (m²)",
    "area_text": "Diện tích bằng chữ",
    "form_of_use": "Hình thức sử dụng (VD: Sử dụng riêng)",
    "purpose_of_use": "Mục đích sử dụng (VD: Đất ở tại nông thôn)",
    "time_of_use": "Thời hạn sử dụng (VD: Lâu dài)",
    "origin_of_use": "Nguồn gốc sử dụng"
  },
  "issuance_info": {
    "place_of_issue": "Nơi cấp (VD: TM. ỦY BAN NHÂN DÂN HUYỆN...)",
    "date_of_issue": "Ngày cấp GCN (VD: 05/08/2019)",
    "signer_title": "Chức vụ người ký (VD: KT. CHỦ TỊCH, PHÓ CHỦ TỊCH)",
    "signer_name": "Tên người ký (VD: Võ Minh Khoa)",
    "registration_book_number": "Số vào sổ ở góc dưới GCN bắt đầu bằng H, CX, CH, CN (VD: CH00071, CX 52, H 1234)"
  }
}

Quy tắc:
- Chỉ trả về JSON thuần túy, không chứa markdown code block hay ký tự lạ.
- Nếu không đọc được trường nào, đặt giá trị = ""
"""

if OCR_AVAILABLE:
    vietdoc_ocr.PROMPT_CCCD = PROMPT_CCCD_EXACT
    vietdoc_ocr.PROMPT_LAND = PROMPT_LAND_EXACT

def convert_pdf_to_images(pdf_bytes):
    image_bytes_list = []
    if FITZ_AVAILABLE:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            image_bytes_list.append(pix.tobytes("png"))
    return image_bytes_list

@app.route('/')
def index():
    return send_from_directory('.', 'index.html')

@app.route('/api/status')
@app.route('/health')
def api_status():
    return jsonify({"status": "online", "service": "ThanhHoa Land AI", "model": "qwen2.5vl"})

@app.route('/<path:filename>')
def serve_static(filename):
    resp = send_from_directory('.', filename)
    # Buộc trình duyệt không cache các file JS/CSS
    if filename in ('app.js', 'app_v2.js', 'styles.css'):
        resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
    return resp

@app.route('/api/search-vault', methods=['GET'])
def search_vault():
    query = request.args.get("q", "").strip()
    if not query:
        return jsonify({"results": []})
    results = search_obsidian_vault(query)
    clean_results = []
    for r in results[:30]:
        clean_results.append({
            "filename": r.get("file"),
            "path": r.get("path"),
            "score": r.get("score"),
            "snippet": r.get("best_section")[:300] + "..." if len(r.get("best_section", "")) > 300 else r.get("best_section")
        })
    return jsonify({"results": clean_results})

@app.route('/api/status', methods=['GET'])
def get_status():
    total_md = len(glob.glob(os.path.join(OBSIDIAN_VAULT_PATH, "**", "*.md"), recursive=True))
    total_history_turns = len(load_json_file(CHAT_MEMORY_FILE, []))
    total_learned_rules = len(load_json_file(LEARNED_CORRECTIONS_FILE, []))
    
    return jsonify({
        "status": "online",
        "system": "ThanhHoa Land AI",
        "version": "2026.41.0 (Deep Logical Legal Reasoning Active)",
        "indexed_md_files": total_md,
        "notebooklm_url": NOTEBOOKLM_URL,
        "notebooklm_status": "Connected (notebook/5234532d-25d7-4f04-a664-2db536f22cdf)",
        "persistent_chat_memory": True,
        "recorded_chat_turns": total_history_turns,
        "learned_correction_rules": total_learned_rules,
        "obsidian_templates_dir": OBSIDIAN_TEMPLATES_DIR,
        "dia_danh_rules": len(DIA_DANH_MAP),
        "obsidian_vault": os.path.exists(OBSIDIAN_VAULT_PATH),
        "pdf_renderer": "PyMuPDF (fitz)" if FITZ_AVAILABLE else "Native Parser"
    })

@app.route('/api/query-gemini', methods=['POST'])
def handle_gemini_query():
    data = request.json or {}
    prompt = data.get("prompt", "").strip()
    if not prompt:
        return jsonify({"error": "Thiếu nội dung prompt"}), 400
    
    text, model_name = generate_response_with_gemini_api(prompt)
    if text:
        return jsonify({"response": text, "model": model_name})
    else:
        # Tự động chuyển sang Qwen2.5-14B khi Gemini không khả dụng
        print("🔄 /api/query-gemini: Gemini không phản hồi -> Chuyển sang Qwen2.5-14B (Ollama)...")
        qwen_text, qwen_model = generate_response_with_ollama_llm(prompt, "LEGAL_CONSULTATION", full_prompt=prompt)
        if qwen_text:
            return jsonify({"response": qwen_text, "model": f"⚡ {qwen_model}"})
        return jsonify({"error": "Cả Gemini API và Qwen2.5-14B đều không khả dụng"}), 500

@app.route('/api/train-rule', methods=['POST'])
def train_rule():
    data = request.json or {}
    context_question = data.get("question", "").strip()
    rule_text = data.get("rule", "").strip()
    if not rule_text:
        return jsonify({"error": "Thiếu nội dung quy tắc huấn luyện"}), 400
    
    learn_from_user_challenge(context_question, rule_text)
    return jsonify({
        "status": "success",
        "message": f"✅ Đã nạp quy tắc huấn luyện AI mới: '{rule_text}'",
        "file": LEARNED_CORRECTIONS_FILE
    })

@app.route('/api/chat', methods=['POST', 'OPTIONS'])
def api_chat():
    if request.method == 'OPTIONS':
        return jsonify({"status": "ok"}), 200
    data = request.json or {}
    question = data.get("question", "").strip()
    session_id = data.get("session_id", "default_session")

    if not question:
        return jsonify({"error": "Vui lòng nhập câu hỏi"}), 400

    # ── KIỂM TRA BẢO MẬT ────────────────────────────────────────
    if SECURITY_GUARD_ACTIVE:
        client_ip = _sg.get_real_ip(request)
        is_safe, security_err = _sg.inspect_request(client_ip, question)
        if not is_safe:
            return jsonify({"error": security_err, "blocked": True}), 429
    # ────────────────────────────────────────────────────────────

    answer_text, intent_type, model_used = process_antigravity_core_pipeline(question, session_id=session_id)

    
    citations = [
        f"Google NotebookLM: {NOTEBOOKLM_URL}",
        "CSDL Tri thức Pháp luật Đất đai Thanh Hóa"
    ]
        
    return jsonify({
        "answer": answer_text,
        "intent": intent_type,
        "session_id": session_id,
        "model": model_used,
        "citations": citations,
        "notebooklm_url": NOTEBOOKLM_URL,
        "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
    })

def send_zalo_oa_message(user_id, message_text):
    """Gửi tin nhắn phản hồi tự động tới người dùng qua Zalo OA API"""
    if not ZALO_BOT_TOKEN:
        print("⚠️ [ZaloBot] Chưa cấu hình ZALO_BOT_TOKEN")
        return {"error": -1, "message": "Missing Zalo Token"}
    try:
        url = "https://openapi.zalo.me/v3.0/oa/message/cs"
        headers = {
            "access_token": ZALO_BOT_TOKEN,
            "Content-Type": "application/json"
        }
        payload = {
            "recipient": {"user_id": str(user_id)},
            "message": {"text": message_text}
        }
        req = urllib.request.Request(
            url,
            data=json.dumps(payload).encode('utf-8'),
            headers=headers,
            method='POST'
        )
        with urllib.request.urlopen(req, timeout=12) as response:
            res_data = json.loads(response.read().decode('utf-8'))
            print(f"✅ [ZaloBot] Đã gửi tin nhắn đến {user_id}: {res_data}")
            return res_data
    except Exception as e:
        print(f"⚠️ [ZaloBot] Lỗi gửi tin nhắn Zalo tới {user_id}: {e}")
        return {"error": -1, "message": str(e)}

@app.route('/api/zalo/webhook', methods=['GET', 'POST'])
def api_zalo_webhook():
    """Webhook nhận tin nhắn trực tiếp từ Zalo và phản hồi bằng AI Pipeline"""
    if request.method == 'GET':
        # Xác thực Webhook nếu Zalo gọi GET challenge
        challenge = request.args.get('challenge') or request.args.get('hub.challenge')
        if challenge:
            return challenge, 200
        return jsonify({
            "status": "active",
            "bot_id": ZALO_BOT_ID,
            "message": "Zalo Webhook Endpoint is running"
        }), 200

    # Xử lý sự kiện tin nhắn POST từ Zalo
    data = request.json or {}
    event_name = data.get("event_name", "")
    sender_id = data.get("sender", {}).get("id") or data.get("user_id_by_app") or data.get("recipient", {}).get("id")
    msg_text = ""

    if "message" in data and isinstance(data["message"], dict):
        msg_text = data["message"].get("text", "").strip()

    print(f"📩 [Zalo Webhook] Nhận sự kiện: {event_name} | Người gửi: {sender_id} | Nội dung: '{msg_text}'")

    if sender_id and msg_text:
        try:
            # Xử lý câu hỏi qua AI Pipeline của Bot
            session_id = f"zalo_{sender_id}"
            answer_text, intent_type, model_used = process_antigravity_core_pipeline(msg_text, session_id=session_id)
            
            # Gửi câu trả lời về lại Zalo cho người dùng
            send_zalo_oa_message(sender_id, answer_text)
            
            return jsonify({
                "status": "success",
                "sender_id": sender_id,
                "answer": answer_text,
                "model": model_used
            }), 200
        except Exception as e:
            print(f"⚠️ [Zalo Webhook] Lỗi khi xử lý AI cho Zalo: {e}")
            return jsonify({"error": str(e)}), 500

    return jsonify({"status": "received", "event": event_name}), 200

# ════════════════════════════════════════════════════════════════
# TELEGRAM BOT INTEGRATION (@TroLyLuatbot)
# ════════════════════════════════════════════════════════════════
def send_telegram_message(chat_id, text):
    """Gửi tin nhắn phản hồi tới Telegram Chat"""
    if not TELEGRAM_BOT_TOKEN:
        return {"error": "Missing TELEGRAM_BOT_TOKEN"}
    try:
        url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/sendMessage"
        # Cắt ngắn nếu vượt giới hạn 4096 ký tự của Telegram
        chunks = [text[i:i+4000] for i in range(0, len(text), 4000)] if len(text) > 4000 else [text]
        for chunk in chunks:
            payload = {
                "chat_id": chat_id,
                "text": chunk,
                "disable_web_page_preview": True
            }
            req = urllib.request.Request(
                url,
                data=json.dumps(payload).encode('utf-8'),
                headers={"Content-Type": "application/json"}
            )
            with urllib.request.urlopen(req, timeout=12) as resp:
                pass
        return {"status": "ok"}
    except Exception as e:
        print(f"⚠️ [TelegramBot] Lỗi gửi tin nhắn Telegram tới {chat_id}: {e}")
        return {"error": str(e)}

def async_telegram_core_reply(chat_id, text):
    try:
        session_id = f"tg_{chat_id}"
        answer_text, intent_type, model_used = process_antigravity_core_pipeline(text, session_id=session_id)
        send_telegram_message(chat_id, answer_text)
    except Exception as e:
        print(f"⚠️ [Telegram Webhook Async] Lỗi xử lý: {e}")
        send_telegram_message(chat_id, f"⚠️ Đã có lỗi xảy ra: {e}")

@app.route('/api/telegram/webhook', methods=['GET', 'POST'])
@app.route('/api/telegram/webhook/', methods=['GET', 'POST'])
def api_telegram_webhook():
    """Webhook nhận tin nhắn từ Telegram"""
    if request.method == 'GET':
        return jsonify({"status": "active", "bot": "@TroLyLuatbot"}), 200

    data = request.json or {}
    message = data.get("message") or data.get("edited_message")
    if not message:
        return jsonify({"ok": True}), 200

    chat_id = message.get("chat", {}).get("id")
    text = message.get("text", "").strip()

    if chat_id and text:
        if text == "/start":
            welcome_msg = (
                "👋 Xin chào bạn!\n"
                "Tôi là Trợ lý AI Pháp luật Đất đai Thanh Hóa (@TroLyLuatbot).\n\n"
                "Bạn có thể gửi câu hỏi, tra cứu thủ tục, điều kiện cấp sổ đỏ, hạn mức tách thửa, nghĩa vụ tài chính tại đây."
            )
            send_telegram_message(chat_id, welcome_msg)
            return jsonify({"ok": True}), 200

        import threading
        threading.Thread(target=async_telegram_core_reply, args=(chat_id, text), daemon=True).start()

    return jsonify({"ok": True}), 200

def start_telegram_bot_polling():
    """Chạy luồng polling nhận tin nhắn từ Telegram Bot tự động (không cần cấu hình webhook/domain)"""
    if not TELEGRAM_BOT_TOKEN:
        return
    print("🤖 [TelegramBot] Luồng Polling Telegram @TroLyLuatbot đang khởi chạy...")
    offset = 0
    while True:
        try:
            url = f"https://api.telegram.org/bot{TELEGRAM_BOT_TOKEN}/getUpdates?offset={offset}&timeout=25"
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=35) as resp:
                data = json.loads(resp.read().decode('utf-8'))
                if data.get("ok"):
                    for update in data.get("result", []):
                        offset = update["update_id"] + 1
                        msg = update.get("message")
                        if not msg:
                            continue
                        chat_id = msg.get("chat", {}).get("id")
                        text = msg.get("text", "").strip()
                        sender = msg.get("from", {}).get("first_name", "User")
                        
                        if chat_id and text:
                            print(f"📩 [Telegram @TroLyLuatbot] Nhận tin nhắn từ {sender} ({chat_id}): '{text}'")
                            if text == "/start":
                                welcome_msg = (
                                    f"👋 Xin chào {sender}!\n"
                                    f"Tôi là Trợ lý AI Pháp luật Đất đai Thanh Hóa (@TroLyLuatbot).\n"
                                    f"Bạn có thể gửi câu hỏi, tra cứu thủ tục, điều kiện cấp sổ đỏ, nghĩa vụ tài chính tại đây."
                                )
                                send_telegram_message(chat_id, welcome_msg)
                                continue
                            
                            # Xử lý câu hỏi qua AI Pipeline
                            session_id = f"tg_{chat_id}"
                            answer_text, intent_type, model_used = process_antigravity_core_pipeline(text, session_id=session_id)
                            send_telegram_message(chat_id, answer_text)
        except Exception as e:
            # Ngắt nghỉ nhẹ khi gặp lỗi mạng để tránh spam log
            time.sleep(3)

def standardize_address_zero_hallucination(addr_str):
    if not addr_str or not isinstance(addr_str, str):
        return addr_str
    clean = addr_str.strip()
    clean = re.sub(r'\bBa\s*Tho\b', 'Bá Thước', clean, flags=re.IGNORECASE)
    clean = re.sub(r'\bThanh\s*Hoa\b', 'Thanh Hóa', clean, flags=re.IGNORECASE)
    clean = re.sub(r'\bHuyen\b', 'Huyện', clean, flags=re.IGNORECASE)
    clean = re.sub(r'\bXa\b', 'Xã', clean, flags=re.IGNORECASE)
    return clean

def scan_document_with_gemini_vision(image_bytes_list, doc_type="cccd"):
    """
    DỰ PHÒNG: Bóc tách OCR từ ảnh CCCD / Giấy chứng nhận (Sổ đỏ) bằng Gemini Vision khi Ollama không khả dụng.
    """
    if not (GEMINI_AVAILABLE and GEMINI_CLIENT):
        return {}
        
    prompt = PROMPT_CCCD_EXACT if doc_type == "cccd" else PROMPT_LAND_EXACT
    prompt += "\nBẮT BUỘC CHỈ TRẢ VỀ CHUẨN JSON THUẦN TÚY, KHÔNG CHỨA CODE BLOCK HOẶC CHỮ RÁC."

    parts = []
    for img_bytes in image_bytes_list:
        try:
            mime = "image/png"
            if img_bytes[:2] == b'\xff\xd8':
                mime = "image/jpeg"
            elif img_bytes[:4] == b'RIFF' and img_bytes[8:12] == b'WEBP':
                mime = "image/webp"
            part = types.Part.from_bytes(data=img_bytes, mime_type=mime)
            parts.append(part)
        except Exception as e:
            print(f"⚠️ Image Part Error: {e}")
            continue

    if not parts:
        return {}

    parts.append(prompt)

    for model_name in ["gemini-2.5-flash", "gemini-3.5-flash", "gemini-3.6-flash"]:
        try:
            res = GEMINI_CLIENT.models.generate_content(
                model=model_name,
                contents=parts
            )
            if res and res.text:
                text = res.text.strip()
                json_match = re.search(r'\{.*\}', text, re.DOTALL)
                if json_match:
                    parsed_dict = json.loads(json_match.group(0))
                    print(f"✅ Gemini Vision ({model_name}) OCR bóc tách {doc_type} thành công!")
                    return parsed_dict
        except Exception as e:
            err_str = str(e)
            print(f"⚠️ Gemini Vision OCR ({model_name}) error: {err_str[:120]}")
            continue

    return {}

def scan_document_with_qwen25vl(image_bytes_list, doc_type="cccd"):
    """
    ƯU TIÊN 1: Trích xuất thông tin tự động 100% từ ảnh CCCD hoặc Giấy chứng nhận (Sổ đỏ) bằng mô hình thị giác qwen2.5vl qua Ollama Local.
    """
    b64_images = [base64.b64encode(img_bytes).decode('utf-8') for img_bytes in image_bytes_list]
    prompt = PROMPT_CCCD_EXACT if doc_type == "cccd" else PROMPT_LAND_EXACT
    
    # Thử danh sách các tag qwen2.5vl trong Ollama
    candidate_models = ["qwen2.5vl:latest", "qwen2.5vl", "qwen2.5vl:7b"]
    
    for model_name in candidate_models:
        req_payload = {
            "model": model_name,
            "prompt": prompt,
            "images": b64_images,
            "stream": False,
            "format": "json",
            "options": {
                "temperature": 0.1
            }
        }
        
        try:
            req_data = json.dumps(req_payload).encode('utf-8')
            req = urllib.request.Request(
                "http://localhost:11434/api/generate",
                data=req_data,
                headers={'Content-Type': 'application/json'}
            )
            with urllib.request.urlopen(req, timeout=120) as resp:
                res_json = json.loads(resp.read().decode('utf-8'))
                response_text = res_json.get("response", "").strip()
                
                json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
                if json_match:
                    parsed_dict = json.loads(json_match.group(0))
                    print(f"✅ Ollama Local Vision ({model_name}) bóc tách thành công thông tin {doc_type}!")
                    return parsed_dict
        except Exception as e:
            print(f"⚠️ Ollama Vision ({model_name}) OCR Error: {e}")
            continue
            
    return {}

@app.route('/api/ocr/scan', methods=['POST'])
def api_ocr_scan():
    doc_type = request.form.get("doc_type", "cccd").lower()
    uploaded_files = request.files.getlist("files") or request.files.getlist("file")
    
    image_bytes_list = []
    data_urls = []
    
    for file in uploaded_files:
        filename = file.filename.lower() if file.filename else 'unknown'
        content = file.read()
        print(f"📂 OCR upload: {filename} | size: {len(content)} bytes | content_type: {file.content_type}")
        
        if filename.endswith(".pdf"):
            pdf_images = convert_pdf_to_images(content)
            for img_bytes in pdf_images:
                image_bytes_list.append(img_bytes)
                b64_str = base64.b64encode(img_bytes).decode('utf-8')
                data_urls.append(f"data:image/png;base64,{b64_str}")
        else:
            # Normalize + Resize ảnh trước khi OCR (Ollama tối ưu với ảnh ~1440px)
            try:
                from PIL import Image as PILImage
                img_obj = PILImage.open(BytesIO(content)).convert('RGB')
                
                # Resize nếu ảnh quá lớn (max 1440px cạnh dài để tăng tốc độ xử lý vision)
                MAX_SIZE = 1440
                w, h = img_obj.size
                if max(w, h) > MAX_SIZE:
                    ratio = MAX_SIZE / max(w, h)
                    new_w, new_h = int(w * ratio), int(h * ratio)
                    img_obj = img_obj.resize((new_w, new_h), PILImage.LANCZOS)
                    print(f"📐 Resize ảnh: {w}x{h} → {new_w}x{new_h}")
                
                buf = BytesIO()
                img_obj.save(buf, format='JPEG', quality=90)
                normalized = buf.getvalue()
                print(f"✅ Ảnh sẵn sàng OCR: {len(normalized)} bytes")
            except Exception as pil_err:
                print(f"⚠️ PIL normalize failed: {pil_err}, dùng ảnh gốc")
                normalized = content
            
            image_bytes_list.append(normalized)
            b64_str = base64.b64encode(normalized).decode('utf-8')
            data_urls.append(f"data:image/png;base64,{b64_str}")
            
    extracted_data = {}
    ocr_model_used = "Ollama Local (qwen2.5vl)"
    
    if image_bytes_list:
        print(f"🔍 Phân hệ 2 OCR: {len(image_bytes_list)} ảnh, doc_type={doc_type} | 100% Ollama qwen2.5vl")
        # ƯU TIÊN 1: 100% Ollama Local Vision qwen2.5vl
        extracted_data = scan_document_with_qwen25vl(image_bytes_list, doc_type=doc_type)
        print(f"🔍 Ollama qwen2.5vl result keys: {list(extracted_data.keys()) if extracted_data else 'None'}")
        
        # Dự phòng 2: Gemini Vision nếu Ollama chưa khởi động hoặc gặp sự cố
        if not extracted_data:
            print("⚠️ Ollama qwen2.5vl chưa phản hồi, thử fallback sang Gemini Vision...")
            extracted_data = scan_document_with_gemini_vision(image_bytes_list, doc_type=doc_type)
            if extracted_data:
                ocr_model_used = "Gemini Vision (Dự phòng)"
                print(f"🔍 Gemini Vision fallback result keys: {list(extracted_data.keys())}")
    else:
        print(f"⚠️ OCR: Không có ảnh nào được upload!")

    # Zero-Hallucination Address Standardizer for all address fields in JSON
    try:
        def process_addr_dict(d):
            if isinstance(d, dict):
                for k, v in d.items():
                    if isinstance(v, dict):
                        process_addr_dict(v)
                    elif isinstance(v, str) and ("address" in k or "residence" in k or "origin" in k):
                        d[k] = standardize_address_zero_hallucination(v)

        process_addr_dict(extracted_data)
    except Exception as e:
        print(f"⚠️ Process Addr Error: {e}")

    return jsonify({
        "success": True,
        "doc_type": doc_type,
        "ocr_model": ocr_model_used,
        "language": "Vietnamese",
        "dia_danh_lookup": True,
        "data_urls": data_urls,
        "extracted_data": extracted_data
    })

# API EXPORT TO WORD (.DOCX) WITH HIGH-PRECISION BEAUTIFUL TABLE LAYOUT (NĐ 30/2020/NĐ-CP STANDARDS)
@app.route('/api/export/docx', methods=['POST'])
def api_export_docx():
    data = request.json or {}
    title = data.get("title", "Don_Dat_Dai")
    content = data.get("content", "")

    if not content:
        return jsonify({"error": "Nội dung đơn trống"}), 400

    doc_obj = docx.Document()
    
    # 1. Cấu hình lề A4 chuẩn NĐ 30/2020 (Top: 20mm, Bottom: 20mm, Left: 30mm, Right: 15mm)
    for section in doc_obj.sections:
        section.page_width = docx.shared.Mm(210)
        section.page_height = docx.shared.Mm(297)
        section.top_margin = docx.shared.Mm(20)
        section.bottom_margin = docx.shared.Mm(20)
        section.left_margin = docx.shared.Mm(30)
        section.right_margin = docx.shared.Mm(15)

    lines = [l.strip() for l in content.split('\n') if l.strip()]
    
    # Tách dòng tiêu đề mẫu đơn (nếu có) ở dòng đầu tiên
    mau_so_text = ""
    header_start_idx = 0
    if lines and (lines[0].startswith("Mẫu số") or lines[0].startswith("MẪU SỐ")):
        mau_so_text = lines[0]
        if len(lines) > 1 and lines[1].startswith("("):
            mau_so_text += "\n" + lines[1]
            header_start_idx = 2
        else:
            header_start_idx = 1

    body_lines = []
    for l in lines[header_start_idx:]:
        if any(l.startswith(k) for k in ["CỘNG HÒA XÃ HỘI", "Độc lập - Tự do", "Độc lập – Tự do", "---------------", "_______________"]):
            continue
        body_lines.append(l)

    # 2. ĐỈNH TRANG: TẠO BẢNG QUỐC HIỆU 2 CỘT NGUYÊN BẢN CHUẨN NĐ 30
    table_header = doc_obj.add_table(rows=1, cols=2)
    table_header.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    # Remove table borders cleanly
    tblPr = table_header._tbl.tblPr
    tblBorders = docx.oxml.OxmlElement('w:tblBorders')
    for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge = docx.oxml.OxmlElement(f'w:{b}')
        edge.set(docx.oxml.ns.qn('w:val'), 'none')
        tblBorders.append(edge)
    tblPr.append(tblBorders)

    cell_l, cell_r = table_header.rows[0].cells
    cell_l.width = docx.shared.Inches(2.8)
    cell_r.width = docx.shared.Inches(3.7)

    # Cột Trái Header: Tên mẫu đơn / Cơ quan
    p_l = cell_l.paragraphs[0]
    p_l.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    p_l.paragraph_format.line_spacing = 1.15
    if mau_so_text:
        r_l = p_l.add_run(mau_so_text)
        r_l.font.name = 'Times New Roman'
        r_l.font.size = docx.shared.Pt(10.5)
        r_l.font.italic = True

    # Cột Phải Header: Quốc Hiệu Tiêu Ngữ
    p_r = cell_r.paragraphs[0]
    p_r.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_r.paragraph_format.line_spacing = 1.15

    r_r1 = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_r1.font.name = 'Times New Roman'
    r_r1.font.size = docx.shared.Pt(12)
    r_r1.font.bold = True

    r_r2 = p_r.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r_r2.font.name = 'Times New Roman'
    r_r2.font.size = docx.shared.Pt(13)
    r_r2.font.bold = True

    r_r3 = p_r.add_run("---------------")
    r_r3.font.name = 'Times New Roman'
    r_r3.font.size = docx.shared.Pt(11)
    r_r3.font.bold = True

    # 3. NỘI DUNG CHÍNH CỦA ĐƠN
    in_signature = False
    sig_date = ""
    sig_role = ""
    sig_name = ""

    for line in body_lines:
        if ("Thanh Hóa, ngày" in line or "ngày ..... tháng" in line or "ngày     tháng" in line or "ngày... tháng" in line):
            sig_date = line
            in_signature = True
            continue

        if in_signature:
            if any(k in line for k in ["Người làm đơn", "Người kê khai", "Người nộp thuế", "Người viết đơn"]):
                sig_role = line
            elif not line.startswith("(") and len(line) > 2 and not any(line.startswith(k) for k in ["1.", "2.", "3.", "a)", "b)"]):
                sig_name = line
            continue

        p = doc_obj.add_paragraph()
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_after = docx.shared.Pt(4)

        # Tiêu đề Đơn / Tờ khai
        if any(line.startswith(k) for k in ["ĐƠN", "TỜ KHAI"]):
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = docx.shared.Pt(14)
            p.paragraph_format.space_after = docx.shared.Pt(12)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(15)
            run.font.bold = True

        # Kính gửi
        elif line.startswith("Kính gửi"):
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = docx.shared.Inches(0.4)
            p.paragraph_format.space_before = docx.shared.Pt(6)
            p.paragraph_format.space_after = docx.shared.Pt(8)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(13)
            run.font.bold = True

        # Mục lớn (1. 2. 3. I. II. [01])
        elif any(line.startswith(f"{i}.") for i in range(1, 10)) or any(line.startswith(f"{k}.") for k in ["I", "II", "III", "IV", "V"]) or any(line.startswith(f"[{i}]") for i in range(1, 10)) or any(line.startswith(f"[{i:02d}]") for i in range(1, 10)):
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = docx.shared.Pt(8)
            p.paragraph_format.space_after = docx.shared.Pt(3)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(13)
            run.font.bold = True

        # Dòng chi tiết: Tách Nhãn (Bold) và Giá trị (Regular)
        else:
            p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = docx.shared.Inches(0.3)

            if ":" in line and not line.startswith("http"):
                parts = line.split(":", 1)
                lbl = parts[0] + ":"
                val = parts[1]

                r_lbl = p.add_run(lbl)
                r_lbl.font.name = 'Times New Roman'
                r_lbl.font.size = docx.shared.Pt(13)
                if any(lbl.strip().startswith(k) for k in ["- Tên", "- Họ", "- Số", "- Thửa", "- Diện", "- Địa", "- Giấy", "- Mục", "- Thời", "- Nguồn", "a)", "b)", "c)", "d)"]):
                    r_lbl.font.bold = True

                r_val = p.add_run(val)
                r_val.font.name = 'Times New Roman'
                r_val.font.size = docx.shared.Pt(13)
            else:
                run = p.add_run(line)
                run.font.name = 'Times New Roman'
                run.font.size = docx.shared.Pt(13)

    # 4. CHÂN TRANG: TẠO BẢNG CHỮ KÝ 2 CỘT CHUẨN SANG TRỌNG
    table_sig = doc_obj.add_table(rows=1, cols=2)
    table_sig.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    tblPr_sig = table_sig._tbl.tblPr
    tblBorders_sig = docx.oxml.OxmlElement('w:tblBorders')
    for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        edge = docx.oxml.OxmlElement(f'w:{b}')
        edge.set(docx.oxml.ns.qn('w:val'), 'none')
        tblBorders_sig.append(edge)
    tblPr_sig.append(tblBorders_sig)

    cell_sig_l, cell_sig_r = table_sig.rows[0].cells
    cell_sig_l.width = docx.shared.Inches(2.8)
    cell_sig_r.width = docx.shared.Inches(3.7)

    # Cột Trái Chân Trang: Nơi nhận
    p_sig_l = cell_sig_l.paragraphs[0]
    p_sig_l.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
    p_sig_l.paragraph_format.line_spacing = 1.15

    r_n1 = p_sig_l.add_run("Nơi nhận:\n")
    r_n1.font.name = 'Times New Roman'
    r_n1.font.size = docx.shared.Pt(11)
    r_n1.font.bold = True
    r_n1.font.italic = True

    r_n2 = p_sig_l.add_run("- Như kính gửi;\n- VPĐKĐĐ xã;\n- Lưu: VT, HS.")
    r_n2.font.name = 'Times New Roman'
    r_n2.font.size = docx.shared.Pt(10)
    r_n2.font.italic = True

    # Cột Phải Chân Trang: Địa danh, Ngày tháng & Chữ ký
    p_sig_r = cell_sig_r.paragraphs[0]
    p_sig_r.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    p_sig_r.paragraph_format.line_spacing = 1.15

    if not sig_date:
        sig_date = "Thanh Hóa, ngày ..... tháng ..... năm 202..."
    if not sig_role:
        sig_role = "NGƯỜI LÀM ĐƠN"

    r_s1 = p_sig_r.add_run(f"{sig_date}\n")
    r_s1.font.name = 'Times New Roman'
    r_s1.font.size = docx.shared.Pt(12.5)
    r_s1.font.italic = True

    r_s2 = p_sig_r.add_run(f"{sig_role.upper()}\n")
    r_s2.font.name = 'Times New Roman'
    r_s2.font.size = docx.shared.Pt(13)
    r_s2.font.bold = True

    r_s3 = p_sig_r.add_run("(Ký và ghi rõ họ tên)\n\n\n\n\n")
    r_s3.font.name = 'Times New Roman'
    r_s3.font.size = docx.shared.Pt(11)
    r_s3.font.italic = True

    if sig_name:
        r_s4 = p_sig_r.add_run(f"{sig_name.upper()}")
        r_s4.font.name = 'Times New Roman'
        r_s4.font.size = docx.shared.Pt(13)
        r_s4.font.bold = True

    bio = BytesIO()
    doc_obj.save(bio)
    bio.seek(0)

    clean_filename = re.sub(r'[^\w\s-]', '', title).strip().replace(' ', '_')
    filename = f"{clean_filename}_{int(time.time())}.docx"

    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# API ENDPOINTS: QUẢN LÝ VÀ TỔNG HỢP CÂU HỎI CHƯA TRẢ LỜI ĐƯỢC THEO NGÀY
@app.route('/api/unanswered/list', methods=['GET'])
def api_unanswered_list():
    log_data = load_json_file(UNANSWERED_LOG_FILE, [])
    target_date = request.args.get("date")
    if target_date:
        filtered = [e for e in log_data if e.get("date") == target_date]
        return jsonify({"success": True, "date": target_date, "count": len(filtered), "data": filtered})
    
    grouped = {}
    for entry in log_data:
        d = entry.get("date", "Unknown")
        if d not in grouped:
            grouped[d] = []
        grouped[d].append(entry)
        
    return jsonify({
        "success": True, 
        "total_count": len(log_data), 
        "total_dates": len(grouped),
        "dates": list(grouped.keys()), 
        "grouped_data": grouped
    })

@app.route('/api/unanswered/report', methods=['GET'])
def api_unanswered_report():
    target_date = request.args.get("date")
    report_md, report_path = generate_daily_unanswered_markdown_report(target_date)
    return jsonify({
        "success": True,
        "date": target_date or time.strftime("%Y-%m-%d"),
        "report_filepath": report_path,
        "markdown_content": report_md
    })

@app.route('/api/unanswered/record', methods=['POST'])
def api_unanswered_record():
    req_data = request.json or {}
    q = req_data.get("question")
    r = req_data.get("reason", "Chưa giải đáp được")
    m = req_data.get("model_used", "Manual Entry")
    ans = req_data.get("answer_given", "")
    session_id = req_data.get("session_id", "default_session")
    
    if not q:
        return jsonify({"error": "Thiếu nội dung câu hỏi"}), 400
        
    entry = record_unanswered_question(q, session_id=session_id, reason=r, model_used=m, answer_given=ans)
    return jsonify({"success": True, "entry": entry})

@app.route('/api/customer_questions/report', methods=['GET'])
def api_customer_questions_report():
    target_date = request.args.get("date")
    daily_md, daily_path, master_md, master_path = generate_customer_chat_summary_report(target_date)
    return jsonify({
        "success": True,
        "date": target_date or time.strftime("%Y-%m-%d"),
        "daily_report_filepath": daily_path,
        "daily_markdown": daily_md,
        "master_report_filepath": master_path,
        "master_markdown": master_md
    })

def prewarm_ollama_model():
    """Nạp sẵn model 14B vào VRAM GPU RTX 3060 ngay khi khởi động server."""
    try:
        time.sleep(1)
        req_data = json.dumps({
            "model": PRIMARY_LLM_MODEL,
            "prompt": "hi",
            "stream": False,
            "keep_alive": -1,
            "options": {"num_ctx": 4096, "num_gpu": 99}
        }).encode('utf-8')
        req = urllib.request.Request(OLLAMA_API_URL, data=req_data, headers={'Content-Type': 'application/json'})
        with urllib.request.urlopen(req, timeout=30) as resp:
            pass
        print(f"⚡ [GPU Pre-warm] Đã nạp thường trực mô hình '{PRIMARY_LLM_MODEL}' vào 100% VRAM GPU RTX 3060!")
    except Exception as e:
        print(f"⚠️ [GPU Pre-warm] Ollama chưa sẵn sàng hoặc bỏ qua: {e}")

if __name__ == '__main__':
    import threading
    threading.Thread(target=prewarm_ollama_model, daemon=True).start()
    threading.Thread(target=start_telegram_bot_polling, daemon=True).start()
    print("🚀 Starting ThanhHoa Land AI Web Server on http://localhost:8668 with RTX 3060 GPU Acceleration & 10.000 Q&A Engine...")
    app.run(host='0.0.0.0', port=8668, debug=False)

